
import os
import re
import time
import json
import csv
import hashlib
import html
import uuid
from io import BytesIO
import urllib.parse
import urllib.request
from typing import Dict, Optional, List, Tuple

import pdfplumber
import plotly.graph_objects as go
import streamlit as st
import streamlit.components.v1 as components
from dotenv import load_dotenv
from openai import OpenAI

try:
    import pycountry
except ImportError:
    pycountry = None

try:
    import geonamescache
except ImportError:
    geonamescache = None

try:
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
except ImportError:
    SimpleDocTemplate = None
    Paragraph = None
    ParagraphStyle = None
    Spacer = None
    getSampleStyleSheet = None
    A4 = None
    mm = None

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None
    Pt = None
    Inches = None
    WD_ALIGN_PARAGRAPH = None


def get_streamlit_secret(key: str, default=None):
    try:
        return st.secrets[key]
    except Exception:
        return default


# =========================================================
# ADVANCED COUNTRY CV TEMPLATE RULES (WorkZo v5.8)
# =========================================================
CV_COUNTRY_RULES = {
    "Germany": {
        "photo": True, "dob": True, "pages": 2,
        "avoid": "Avoid long paragraphs; photo and date of birth are optional and should not be invented.",
        "sections": ["Personal Information", "Professional Summary", "Work Experience", "Education", "Skills", "Languages", "Certifications"]
    },
    "Austria": {
        "photo": True, "dob": False, "pages": 2,
        "avoid": "Do not invent photo, nationality, marital status or visa details.",
        "sections": ["Personal Information", "Professional Summary", "Work Experience", "Education", "Skills", "Languages"]
    },
    "Switzerland": {
        "photo": True, "dob": False, "pages": 2,
        "avoid": "Keep formal and concise; do not invent personal details.",
        "sections": ["Profile", "Work Experience", "Education", "Skills", "Languages", "Certifications"]
    },
    "Canada": {
        "photo": False, "dob": False, "pages": 2,
        "avoid": "Never include photo, date of birth, marital status, nationality, religion, gender, full street address, or sensitive personal details.",
        "sections": ["Professional Summary", "Core Skills", "Professional Experience", "Projects", "Education", "Certifications", "Languages"]
    },
    "United States": {
        "photo": False, "dob": False, "pages": 1,
        "avoid": "Never include photo, date of birth, marital status, nationality, religion, gender, or full personal address.",
        "sections": ["Professional Summary", "Skills", "Professional Experience", "Projects", "Education"]
    },
    "United Kingdom": {
        "photo": False, "dob": False, "pages": 2,
        "avoid": "Do not include photo, date of birth, marital status, nationality, or sensitive personal details.",
        "sections": ["Professional Profile", "Key Skills", "Work Experience", "Education", "Certifications"]
    },
    "Australia": {
        "photo": False, "dob": False, "pages": 3,
        "avoid": "Do not include photo, date of birth, marital status, nationality, or sensitive personal details.",
        "sections": ["Professional Summary", "Key Skills", "Professional Experience", "Projects", "Education", "Certifications"]
    },
    "Netherlands": {
        "photo": False, "dob": False, "pages": 2,
        "avoid": "Keep it direct and concise; avoid unnecessary personal details.",
        "sections": ["Profile", "Work Experience", "Education", "Skills", "Languages"]
    },
    "France": {
        "photo": True, "dob": False, "pages": 2,
        "avoid": "Photo is optional; do not invent personal details.",
        "sections": ["Profil", "Expérience professionnelle", "Formation", "Compétences", "Langues"]
    },
    "India": {
        "photo": False, "dob": False, "pages": 2,
        "avoid": "Avoid unnecessary personal details; focus on skills, projects, tools, achievements, and education.",
        "sections": ["Professional Summary", "Skills", "Work Experience", "Projects", "Education", "Certifications"]
    },
    "Singapore": {
        "photo": False, "dob": False, "pages": 2,
        "avoid": "Keep it concise and ATS-friendly; avoid sensitive personal details.",
        "sections": ["Professional Summary", "Skills", "Experience", "Education", "Certifications", "Languages"]
    },
}

def get_country_cv_rules(country):
    return CV_COUNTRY_RULES.get(country,{
        "photo":False,
        "dob":False,
        "pages":2,
        "sections":["Professional Summary","Work Experience","Education","Skills"]
    })

# =========================================================
# PAGE CONFIG + LOGO PATH
# =========================================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = os.path.join(BASE_DIR, "logo.png")

st.set_page_config(
    page_title="WorkZo AI",
    page_icon=LOGO_PATH if os.path.exists(LOGO_PATH) else "🚀",
    layout="wide"
)

# =========================================================
# LOAD ENV
# =========================================================
load_dotenv()

api_key = os.getenv("OPENAI_API_KEY") or get_streamlit_secret("OPENAI_API_KEY")
if not api_key:
    st.error("OPENAI_API_KEY not found. Please check your .env file or Streamlit secrets.")
    st.stop()

client = OpenAI(api_key=api_key)

# =========================================================
# RATE LIMITING
# =========================================================
MAX_REQUESTS_PER_HOUR = 25
SCORING_VERSION = "v5.5"

if "request_count" not in st.session_state:
    st.session_state.request_count = 0

if "first_request_time" not in st.session_state:
    st.session_state.first_request_time = time.time()

if time.time() - st.session_state.first_request_time > 3600:
    st.session_state.request_count = 0
    st.session_state.first_request_time = time.time()

def can_make_request() -> bool:
    return st.session_state.request_count < MAX_REQUESTS_PER_HOUR
ISSUES_FILE = os.path.join(BASE_DIR, "workzo_beta_issues.csv")

# =========================================================
# NAVIGATION / SCROLL HELPERS
# =========================================================
def request_scroll_to_top() -> None:
    """Ask the next rerun to force the browser viewport to the top."""
    st.session_state["_workzo_scroll_to_top"] = True


def maybe_scroll_to_top() -> None:
    """Force Streamlit's parent page and scroll containers back to the top.
    This prevents pages from opening in the middle after a button click or upload.
    """
    if not st.session_state.pop("_workzo_scroll_to_top", False):
        return
    components.html(
        """
        <script>
        const forceTop = () => {
            const doc = window.parent.document;
            const targets = [
                window.parent,
                doc.documentElement,
                doc.body,
                doc.querySelector('section.main'),
                doc.querySelector('[data-testid="stAppViewContainer"]'),
                doc.querySelector('[data-testid="stMain"]'),
                doc.querySelector('[data-testid="stMainBlockContainer"]')
            ].filter(Boolean);
            for (const t of targets) {
                try {
                    if (typeof t.scrollTo === 'function') {
                        t.scrollTo({ top: 0, left: 0, behavior: 'auto' });
                    }
                    t.scrollTop = 0;
                } catch(e) {}
            }
        };
        forceTop();
        setTimeout(forceTop, 50);
        setTimeout(forceTop, 150);
        setTimeout(forceTop, 350);
        setTimeout(forceTop, 700);
        </script>
        """,
        height=0,
        width=0,
    )


def sync_navigation_state(page_key: str) -> None:
    """Store the selected page in app-owned state without touching widget keys."""
    st.session_state.page = page_key
    st.session_state.nav_page = page_key
    st.session_state.nav_change_nonce = st.session_state.get("nav_change_nonce", 0) + 1

def queue_navigation(page_key: str) -> None:
    """Button callback used by dashboard navigation cards."""
    st.session_state._workzo_pending_nav = page_key
    request_scroll_to_top()

def consume_pending_navigation() -> None:
    page_key = st.session_state.pop("_workzo_pending_nav", None)
    if page_key:
        sync_navigation_state(page_key)

def go_home() -> None:
    queue_navigation("dashboard")


def register_request() -> None:
    st.session_state.request_count += 1


# =========================================================
# BETA ANALYTICS - PRIVACY SAFE
# =========================================================
ANALYTICS_FILE = os.path.join(BASE_DIR, "workzo_beta_analytics.csv")
FEEDBACK_FILE = os.path.join(BASE_DIR, "workzo_beta_feedback.csv")

def get_or_create_anonymous_user_id() -> str:
    if "anonymous_user_id" not in st.session_state:
        st.session_state.anonymous_user_id = str(uuid.uuid4())
        st.session_state.is_repeat_user = False
    return st.session_state.anonymous_user_id

def get_or_create_session_id() -> str:
    if "session_id" not in st.session_state:
        st.session_state.session_id = str(uuid.uuid4())
        st.session_state.session_started_at = time.time()
    return st.session_state.session_id

def init_beta_analytics():
    get_or_create_anonymous_user_id()
    get_or_create_session_id()
    if "feature_usage_counts" not in st.session_state:
        st.session_state.feature_usage_counts = {}
    if "analytics_events" not in st.session_state:
        st.session_state.analytics_events = []
    if "last_tracked_page" not in st.session_state:
        st.session_state.last_tracked_page = ""

def safe_analytics_value(value):
    if value is None:
        return ""
    value = str(value)
    if len(value) > 120:
        value = value[:120] + "..."
    return value.replace("\n", " ").replace("\r", " ").strip()

def send_analytics_to_webhook(event: Dict):
    webhook_url = os.getenv("ANALYTICS_WEBHOOK_URL") or get_streamlit_secret("ANALYTICS_WEBHOOK_URL")
    if not webhook_url:
        return
    try:
        data = json.dumps(event).encode("utf-8")
        req = urllib.request.Request(
            webhook_url,
            data=data,
            headers={"Content-Type": "application/json", "User-Agent": "WorkZoAI-Beta"}
        )
        urllib.request.urlopen(req, timeout=4)
    except Exception:
        pass

def track_event(event_name: str, feature: str = "", metadata: Optional[Dict] = None):
    init_beta_analytics()
    metadata = metadata or {}
    session_duration_seconds = int(time.time() - st.session_state.get("session_started_at", time.time()))

    event = {
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
        "anonymous_user_id": st.session_state.get("anonymous_user_id", ""),
        "session_id": st.session_state.get("session_id", ""),
        "event_name": safe_analytics_value(event_name),
        "feature": safe_analytics_value(feature),
        "session_duration_seconds": session_duration_seconds,
        "repeat_user": safe_analytics_value(st.session_state.get("is_repeat_user", False)),
        "country": safe_analytics_value(st.session_state.get("country", "")),
        "migration_country": safe_analytics_value(st.session_state.get("migration_country", "")),
        "user_status": safe_analytics_value(st.session_state.get("user_status", "")),
        "preferred_language": safe_analytics_value(st.session_state.get("preferred_language", "")),
        "cv_uploaded": safe_analytics_value(bool(st.session_state.get("cv_text", ""))),
        "metadata": safe_analytics_value(json.dumps(metadata, ensure_ascii=False)),
    }

    st.session_state.analytics_events.append(event)

    if feature:
        st.session_state.feature_usage_counts[feature] = st.session_state.feature_usage_counts.get(feature, 0) + 1

    try:
        file_exists = os.path.exists(ANALYTICS_FILE)
        with open(ANALYTICS_FILE, "a", newline="", encoding="utf-8") as f:
            writer = csv.DictWriter(f, fieldnames=list(event.keys()))
            if not file_exists:
                writer.writeheader()
            writer.writerow(event)
    except Exception:
        pass

    send_analytics_to_webhook(event)

def track_feature_view(feature_name: str):
    if st.session_state.get("last_tracked_page") != feature_name:
        track_event("feature_view", feature_name)
        st.session_state.last_tracked_page = feature_name

def track_button_click(button_name: str, feature: str = "", metadata: Optional[Dict] = None):
    meta = {"button": button_name}
    if metadata:
        meta.update(metadata)
    track_event("button_click", feature or button_name, meta)


def read_csv_rows(file_path: str) -> List[Dict]:
    if not os.path.exists(file_path):
        return []
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            return list(reader)
    except Exception:
        return []

def render_founder_dashboard():
    st.markdown("### Founder Analytics Dashboard")
    st.caption("Private beta metrics. This dashboard is only shown after entering the founder PIN.")

    rows = read_csv_rows(ANALYTICS_FILE)

    if not rows:
        st.info("No analytics events recorded yet.")
        return

    total_events = len(rows)
    unique_users = len({r.get("anonymous_user_id", "") for r in rows if r.get("anonymous_user_id", "")})
    unique_sessions = len({r.get("session_id", "") for r in rows if r.get("session_id", "")})

    returning_users = 0
    user_sessions = {}
    for r in rows:
        uid = r.get("anonymous_user_id", "")
        sid = r.get("session_id", "")
        if uid and sid:
            user_sessions.setdefault(uid, set()).add(sid)
    returning_users = sum(1 for sessions in user_sessions.values() if len(sessions) > 1)

    durations = []
    for r in rows:
        try:
            durations.append(int(float(r.get("session_duration_seconds", 0))))
        except Exception:
            pass
    avg_session = int(sum(durations) / len(durations)) if durations else 0

    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.metric("Total events", total_events)
    with c2:
        st.metric("Unique users", unique_users)
    with c3:
        st.metric("Sessions", unique_sessions)
    with c4:
        st.metric("Avg event time", f"{avg_session // 60}m {avg_session % 60}s")

    st.markdown("#### Returning users")
    st.metric("Estimated returning users", returning_users)

    def count_by(key: str):
        counts = {}
        for r in rows:
            value = r.get(key, "") or "Unknown"
            counts[value] = counts.get(value, 0) + 1
        return sorted(counts.items(), key=lambda x: x[1], reverse=True)

    col_a, col_b = st.columns(2)
    with col_a:
        st.markdown("#### Most used features")
        for feature, count in count_by("feature")[:10]:
            st.write(f"- **{feature}**: {count}")

        st.markdown("#### Event types")
        for event, count in count_by("event_name")[:10]:
            st.write(f"- **{event}**: {count}")

    with col_b:
        st.markdown("#### Top countries")
        for country, count in count_by("country")[:10]:
            st.write(f"- **{country}**: {count}")

        st.markdown("#### User situations")
        for status, count in count_by("user_status")[:10]:
            st.write(f"- **{status}**: {count}")

    st.markdown("#### Feedback")
    feedback_rows = read_csv_rows(FEEDBACK_FILE)
    if feedback_rows:
        avg_rating_values = []
        for r in feedback_rows:
            try:
                avg_rating_values.append(int(float(r.get("rating", 0))))
            except Exception:
                pass
        avg_rating = round(sum(avg_rating_values) / len(avg_rating_values), 1) if avg_rating_values else "—"

        f1, f2 = st.columns(2)
        with f1:
            st.metric("Feedback responses", len(feedback_rows))
        with f2:
            st.metric("Average rating", avg_rating)

        with st.expander("View latest feedback", expanded=False):
            for r in feedback_rows[-10:][::-1]:
                st.markdown(f"""
**Rating:** {r.get('rating','')} / 5  
**Feature:** {r.get('feature','')}  
**Worked well:** {r.get('worked_well','')}  
**Needs improvement:** {r.get('needs_improvement','')}
---
""")
    else:
        st.info("No feedback submitted yet.")

    if os.path.exists(ANALYTICS_FILE):
        with open(ANALYTICS_FILE, "rb") as f:
            st.download_button(
                "Download analytics CSV",
                data=f.read(),
                file_name="workzo_beta_analytics.csv",
                mime="text/csv",
                key="founder_download_analytics"
            )

    if os.path.exists(FEEDBACK_FILE):
        with open(FEEDBACK_FILE, "rb") as f:
            st.download_button(
                "Download feedback CSV",
                data=f.read(),
                file_name="workzo_beta_feedback.csv",
                mime="text/csv",
                key="founder_download_feedback"
            )

    if os.path.exists(ISSUES_FILE):
        with open(ISSUES_FILE, "rb") as f:
            st.download_button(
                "Download issue reports CSV",
                data=f.read(),
                file_name="workzo_beta_issues.csv",
                mime="text/csv",
                key="founder_download_issues"
            )
def save_feedback(feature: str, rating: int, worked_well: str, needs_improvement: str):
    init_beta_analytics()

    feedback = {
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
        "anonymous_user_id": st.session_state.get("anonymous_user_id", ""),
        "session_id": st.session_state.get("session_id", ""),
        "feature": safe_analytics_value(feature),
        "rating": safe_analytics_value(rating),
        "worked_well": safe_analytics_value(worked_well),
        "needs_improvement": safe_analytics_value(needs_improvement),
        "country": safe_analytics_value(st.session_state.get("country", "")),
        "migration_country": safe_analytics_value(st.session_state.get("migration_country", "")),
        "user_status": safe_analytics_value(st.session_state.get("user_status", "")),
        "preferred_language": safe_analytics_value(st.session_state.get("preferred_language", "")),
    }

    try:
        file_exists = os.path.exists(FEEDBACK_FILE)
        with open(FEEDBACK_FILE, "a", newline="", encoding="utf-8") as f:
            writer = csv.DictWriter(f, fieldnames=list(feedback.keys()))
            if not file_exists:
                writer.writeheader()
            writer.writerow(feedback)
    except Exception:
        pass

    track_event("feedback_submitted", "Feedback", {"feature": feature, "rating": rating})

def render_feedback_collector(current_feature: str = ""):
    with st.expander("Share quick feedback", expanded=False):
        st.caption("Help improve WorkZo. Please do not paste CV text or personal details here.")
        feature = st.selectbox(
            "Which feature are you giving feedback on?",
            ["Overall", "Dashboard", "Job Assist", "Document Tools", "Work-O-Bot", "CV Template Builder", "Translation", "Other"],
            index=0,
            key=f"feedback_feature_{current_feature or 'general'}"
        )
        rating = st.slider("How useful was this?", 1, 5, 4, key=f"feedback_rating_{current_feature or 'general'}")
        worked_well = st.text_area("What worked well?", height=80, key=f"feedback_good_{current_feature or 'general'}")
        needs_improvement = st.text_area("What should improve?", height=80, key=f"feedback_bad_{current_feature or 'general'}")

        if st.button("Submit feedback", key=f"submit_feedback_{current_feature or 'general'}"):
            save_feedback(feature, rating, worked_well, needs_improvement)
            st.success("Thank you — your feedback was saved.")

def save_issue_report(area: str, issue_type: str, description: str):
    init_beta_analytics()
    issue = {
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
        "anonymous_user_id": st.session_state.get("anonymous_user_id", ""),
        "session_id": st.session_state.get("session_id", ""),
        "area": safe_analytics_value(area),
        "issue_type": safe_analytics_value(issue_type),
        "description": safe_analytics_value(description),
        "country": safe_analytics_value(st.session_state.get("country", "")),
        "user_status": safe_analytics_value(st.session_state.get("user_status", "")),
        "preferred_language": safe_analytics_value(st.session_state.get("preferred_language", "")),
    }
    try:
        file_exists = os.path.exists(ISSUES_FILE)
        with open(ISSUES_FILE, "a", newline="", encoding="utf-8") as f:
            writer = csv.DictWriter(f, fieldnames=list(issue.keys()))
            if not file_exists:
                writer.writeheader()
            writer.writerow(issue)
    except Exception:
        pass
    track_event("issue_reported", "Issue Reporting", {"area": area, "issue_type": issue_type})

def render_issue_reporter(current_feature: str = ""):
    with st.expander("🐞 Report a problem", expanded=False):
        st.caption("Tell us what went wrong. Please do not paste private CV text or personal details.")
        area = st.selectbox(
            "Where did this happen?",
            ["Overall", "Onboarding", "Dashboard", "Job Assist", "CV Builder", "Document Tools", "Work-O-Bot", "Downloads", "Other"],
            index=0,
            key=f"issue_area_{current_feature or 'general'}"
        )
        issue_type = st.selectbox(
            "Issue type",
            ["Bug / error", "Confusing output", "Wrong language", "Too much information", "Download problem", "Job search problem", "Feature suggestion"],
            index=0,
            key=f"issue_type_{current_feature or 'general'}"
        )
        description = st.text_area("What happened?", height=90, key=f"issue_description_{current_feature or 'general'}")
        if st.button("Submit issue", key=f"submit_issue_{current_feature or 'general'}"):
            if description.strip():
                save_issue_report(area, issue_type, description)
                st.success("Thanks — the issue report was saved.")
            else:
                st.warning("Please describe the issue briefly.")

def maybe_render_founder_access():
    founder_pin = os.getenv("FOUNDER_PIN") or get_streamlit_secret("FOUNDER_PIN")

    with st.sidebar.expander(txt("founder_access"), expanded=False):
        if founder_pin:
            entered_pin = st.text_input(txt("founder_pin"), type="password", key="founder_pin_input")
            if entered_pin == founder_pin:
                st.session_state.founder_unlocked = True
                st.success("Founder mode unlocked.")
        else:
            st.caption("FOUNDER_PIN is not configured. For local testing, create a temporary PIN below. Before sharing publicly, add FOUNDER_PIN in Streamlit Secrets.")
            temp_pin = st.text_input("Temporary founder PIN", type="password", key="founder_temp_pin_input")
            if temp_pin and len(temp_pin) >= 4:
                st.session_state.founder_unlocked = True
                st.success("Founder mode unlocked for this session.")

    if st.session_state.get("founder_unlocked"):
        with st.sidebar:
            if st.button(txt("founder_dashboard"), key="open_founder_dashboard"):
                st.session_state.page = txt("founder_dashboard")


def render_beta_analytics_summary():
    init_beta_analytics()
    session_duration_seconds = int(time.time() - st.session_state.get("session_started_at", time.time()))
    minutes = session_duration_seconds // 60
    seconds = session_duration_seconds % 60

    st.markdown("### Beta Analytics")
    st.caption("Privacy-safe usage tracking. CV text, email, phone number, and documents are not stored.")

    c1, c2, c3 = st.columns(3)
    with c1:
        st.metric("Session time", f"{minutes}m {seconds}s")
    with c2:
        st.metric("Events this session", len(st.session_state.get("analytics_events", [])))
    with c3:
        st.metric("Features used", len(st.session_state.get("feature_usage_counts", {})))

    if st.session_state.get("feature_usage_counts"):
        st.markdown("#### Feature usage this session")
        for feature, count in sorted(st.session_state.feature_usage_counts.items(), key=lambda x: x[1], reverse=True):
            st.write(f"- **{feature}**: {count}")

    if os.path.exists(ANALYTICS_FILE):
        try:
            with open(ANALYTICS_FILE, "rb") as f:
                st.download_button(
                    "Download beta analytics CSV",
                    data=f.read(),
                    file_name="workzo_beta_analytics.csv",
                    mime="text/csv",
                    key="download_beta_analytics_csv"
                )
        except Exception:
            pass


# =========================================================
# STYLES
# =========================================================
st.markdown("""
<style>
    .stApp {
        background: linear-gradient(180deg, #0b1020 0%, #111827 22%, #0f172a 100%);
    }
    .block-container {
        padding-top: 1.2rem;
        padding-bottom: 2rem;
        max-width: 1200px;
    }
    .beta-badge {
        background: linear-gradient(135deg, #14b8a6, #2563eb);
        color: white;
        padding: 4px 10px;
        border-radius: 999px;
        font-size: 12px;
        display: inline-block;
        font-weight: 600;
    }
    .hero-card {
        background: linear-gradient(135deg, rgba(37,99,235,0.22), rgba(20,184,166,0.16));
        border: 1px solid rgba(148,163,184,0.18);
        border-radius: 24px;
        padding: 20px 22px;
        box-shadow: 0 12px 35px rgba(2,6,23,0.25);
        margin-bottom: 14px;
    }
    .glass-card {
        border: 1px solid rgba(148,163,184,0.16);
        border-radius: 20px;
        padding: 16px;
        background: rgba(15,23,42,0.55);
        backdrop-filter: blur(8px);
        box-shadow: 0 8px 24px rgba(2,6,23,0.18);
        margin-bottom: 12px;
    }
    .metric-card {
        border: 1px solid rgba(148,163,184,0.16);
        border-radius: 20px;
        padding: 16px;
        background: linear-gradient(180deg, rgba(30,41,59,0.86), rgba(15,23,42,0.74));
        min-height: 116px;
        box-shadow: 0 8px 24px rgba(2,6,23,0.18);
    }
    .metric-label {
        color: #94a3b8;
        font-size: 0.92rem;
        margin-bottom: 8px;
    }
    .metric-value {
        color: white;
        font-size: 1.8rem;
        font-weight: 700;
        line-height: 1.1;
        margin-bottom: 6px;
    }
    .metric-foot {
        color: #cbd5e1;
        font-size: 0.9rem;
    }
    .card {
        border: 1px solid rgba(148,163,184,0.16);
        border-radius: 18px;
        padding: 16px;
        background: rgba(15,23,42,0.58);
        margin-bottom: 12px;
    }
    .section-title {
        font-size: 1.05rem;
        font-weight: 700;
        margin-bottom: 8px;
        color: #f8fafc;
    }
    .small-muted {
        color: #94a3b8;
        font-size: 0.95rem;
    }
    .pill {
        display: inline-block;
        padding: 6px 10px;
        border-radius: 999px;
        background: rgba(37,99,235,0.12);
        border: 1px solid rgba(96,165,250,0.18);
        color: #dbeafe;
        margin-right: 6px;
        margin-bottom: 6px;
        font-size: 0.9rem;
    }
    .feature-tile {
        border: 1px solid rgba(148,163,184,0.16);
        border-radius: 18px;
        padding: 16px;
        min-height: 140px;
        background: rgba(15,23,42,0.58);
        margin-bottom: 10px;
    }
    .feature-title {
        font-size: 1rem;
        font-weight: 700;
        color: #f8fafc;
        margin-bottom: 6px;
    }
    .feature-copy {
        color: #cbd5e1;
        font-size: 0.92rem;
        line-height: 1.45;
    }
    .workflow-step {
        border: 1px solid rgba(148,163,184,0.16);
        border-radius: 16px;
        padding: 12px 14px;
        background: rgba(30,41,59,0.5);
        margin-bottom: 8px;
    }
    .nav-help-card {
        border: 1px solid rgba(148,163,184,0.16);
        border-radius: 16px;
        padding: 12px 14px;
        background: rgba(15,23,42,0.55);
        margin-bottom: 8px;
        color: #cbd5e1;
        font-size: 0.92rem;
        line-height: 1.45;
    }
    .nav-help-card b { color: #f8fafc; }
    .next-action-card {
        border: 1px solid rgba(20,184,166,0.28);
        border-radius: 20px;
        padding: 18px 18px;
        background: linear-gradient(135deg, rgba(20,184,166,0.18), rgba(37,99,235,0.14));
        box-shadow: 0 10px 28px rgba(2,6,23,0.20);
        margin: 14px 0 16px 0;
    }
    .next-action-label {
        color: #99f6e4;
        font-size: 0.88rem;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 0.04em;
        margin-bottom: 6px;
    }
    .next-action-title {
        color: #f8fafc;
        font-size: 1.25rem;
        font-weight: 800;
        margin-bottom: 6px;
    }
    .next-action-copy {
        color: #dbeafe;
        font-size: 0.98rem;
        line-height: 1.45;
    }

    .choice-card {
        border: 1px solid rgba(148,163,184,0.16);
        border-radius: 18px;
        padding: 16px;
        min-height: 104px;
        background: rgba(15,23,42,0.58);
        margin-bottom: 10px;
    }
    .choice-card-title {
        font-size: 1rem;
        font-weight: 700;
        color: #f8fafc;
        margin-bottom: 6px;
    }
    .choice-card-copy {
        color: #cbd5e1;
        font-size: 0.92rem;
        line-height: 1.45;
    }
    div[data-testid="stFileUploader"] {
        margin-top: 0px;
    }
    div[data-testid="stButton"] > button {
        border-radius: 14px;
        border: 1px solid rgba(96,165,250,0.18);
    }

    .workzo-top-line {
        height: 1px;
        background: rgba(148,163,184,0.16);
        margin: 8px 0 18px 0;
    }
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background: rgba(15,23,42,0.35);
        border-radius: 16px;
        padding: 6px;
    }
    .stTabs [data-baseweb="tab"] {
        border-radius: 12px;
        padding: 8px 14px;
    }
    section[data-testid="stSidebar"] {
        border-right: 1px solid rgba(148,163,184,0.14);
    }
    section[data-testid="stSidebar"] div[data-testid="stVerticalBlock"] {
        gap: 0.6rem;
    }

    header[data-testid="stHeader"] {
        background: transparent !important;
        height: 0rem !important;
    }
    div[data-testid="stToolbar"] {
        display: none !important;
    }
    div[data-testid="stDecoration"] {
        display: none !important;
    }

    /* Cleaner sidebar scrolling and progress display */
    section[data-testid="stSidebar"] {
        border-right: 1px solid rgba(148,163,184,0.16);
    }
    section[data-testid="stSidebar"] ::-webkit-scrollbar {
        width: 7px;
    }
    section[data-testid="stSidebar"] ::-webkit-scrollbar-track {
        background: transparent;
    }
    section[data-testid="stSidebar"] ::-webkit-scrollbar-thumb {
        background: rgba(148,163,184,0.32);
        border-radius: 999px;
    }
    section[data-testid="stSidebar"] ::-webkit-scrollbar-thumb:hover {
        background: rgba(148,163,184,0.52);
    }
    section[data-testid="stSidebar"] div[data-testid="stProgress"] > div {
        height: 8px;
        border-radius: 999px;
    }
    .dashboard-equal-card {
        min-height: 260px;
    }

    /* WorkZo v8.9 dashboard polish */
    .workzo-section-title {
        font-size: 1.35rem;
        font-weight: 800;
        color: #f8fafc;
        margin: 22px 0 12px 0;
        letter-spacing: -0.01em;
    }
    .workzo-mini-title {
        font-size: 0.95rem;
        font-weight: 750;
        color: #e2e8f0;
        margin-bottom: 8px;
        line-height: 1.35;
    }
    .workzo-text {
        color: #f8fafc;
        font-size: 0.98rem;
        line-height: 1.55;
    }
    .workzo-muted-title {
        color: #94a3b8;
        font-size: 0.82rem;
        font-weight: 750;
        text-transform: uppercase;
        letter-spacing: 0.04em;
        margin: 12px 0 6px 0;
    }
    .resume-insights-line {
        border-top: 1px solid rgba(148,163,184,0.18);
        border-bottom: 1px solid rgba(148,163,184,0.12);
        padding: 14px 0 12px 0;
        margin: 8px 0 18px 0;
    }
    .resume-insights-line ul {
        margin-top: 4px;
        padding-left: 1.15rem;
    }
    .resume-insights-line li {
        margin-bottom: 6px;
        line-height: 1.45;
    }
    .sidebar-progress-label {
        color: #cbd5e1;
        font-size: 0.82rem;
        margin-top: -6px;
    }


    /* WorkZo v9.4 SaaS polish */
    .workzo-sidebar-section-label {
        color: #94a3b8;
        font-size: 0.72rem;
        font-weight: 800;
        text-transform: uppercase;
        letter-spacing: 0.08em;
        margin: 18px 0 6px 0;
    }
    .workzo-sidebar-chip {
        display: block;
        border: 1px solid rgba(148,163,184,0.16);
        background: rgba(15,23,42,0.46);
        border-radius: 12px;
        padding: 8px 10px;
        margin: 5px 0;
        color: #cbd5e1;
        font-size: 0.84rem;
        line-height: 1.25;
    }
    .workzo-sidebar-muted {
        color: #94a3b8;
        font-size: 0.82rem;
        line-height: 1.35;
        margin: 4px 0 8px 0;
    }
    .compact-job-link-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(150px, 1fr));
        gap: 8px;
        margin-top: 8px;
    }
    .compact-job-link {
        display: block;
        border: 1px solid rgba(148,163,184,0.16);
        background: rgba(15,23,42,0.48);
        border-radius: 12px;
        padding: 9px 11px;
        color: #e2e8f0 !important;
        text-decoration: none !important;
        font-size: 0.9rem;
        font-weight: 650;
        white-space: nowrap;
        overflow: hidden;
        text-overflow: ellipsis;
    }
    .compact-job-link:hover {
        border-color: rgba(20,184,166,0.45);
        background: rgba(20,184,166,0.10);
        color: #ffffff !important;
    }
    .workzo-saas-page-card {
        border: 1px solid rgba(148,163,184,0.12);
        background: rgba(15,23,42,0.38);
        border-radius: 26px;
        padding: 18px 20px;
        margin-bottom: 16px;
    }

</style>
""", unsafe_allow_html=True)

# =========================================================
# UI LANGUAGE TEXT
# =========================================================
UI_TEXT = {
    "English": {
        "title": "Your 360° AI Career Assistant",
        "subtitle": "AI-powered career support based on your country and CV.",
        "onboarding_title": "Welcome to WorkZo",
        "onboarding_subtitle": "Start with your country and CV. WorkZo will build the rest from there.",
        "country": "Select a Country",
        "ui_language": "App Language",
        "response_language": "AI Response Language",
        "resume_input": "Resume Input",
        "upload_cv": "Upload CV",
        "create_cv": "Create CV",
        "upload_resume": "Upload your CV",
        "create_resume": "Create your CV",
        "full_name": "Full Name",
        "email": "Email",
        "phone": "Phone Number",
        "summary": "Professional Summary",
        "skills": "Skills",
        "experience": "Work Experience",
        "education": "Education",
        "continue": "Continue to Dashboard",
        "edit_setup": "Edit Onboarding",
        "dashboard": "Dashboard",
        "actions": "Explore tools",
        "understand_job": "Understand Job",
        "improve_cv": "Improve CV",
        "career_communication": "Career Communication",
        "find_jobs": "Find Jobs",
        "mock_interview": "Mock Interview",
        "skill_gap": "Skill Gap",
        "career_roadmap": "Career Roadmap",
        "job_assist": "Job Assist",
        "document_tools": "Document Tools",
        "cv_translator": "CV Translator",
        "cover_letter_generator": "Cover Letter Generator",
        "cover_letter_translator": "Cover Letter Translator",
        "workobot": "Work-O-Bot",
        "workobot_sub": "AI career and language coach",
        "workobot_desc_short": "Mock interviews, career insights, skill gaps, communication practice, and Work-O-Bot chat.",
        "go_workobot": "Open Work-O-Bot",
        "cv_documents": "CV & Documents",
        "interview_practice": "Interview Practice",
        "career_insights": "Career Insights",
        "navigation_help": "Where to find what",
        "dashboard_desc_short": "Start here: scores, next steps, and recommended actions.",
        "job_assist_desc_short": "Find jobs or understand a job description before applying.",
        "cv_documents_desc_short": "Improve CV, make cover letters, translate documents, and use country CV templates.",
        "interview_practice_desc_short": "Practice interview answers, mock questions, and speaking preparation.",
        "career_insights_desc_short": "Check skill gaps, roadmap, country readiness, and next best step.",
        "preferred_language_help": "One language for the app, AI replies, and generated documents.",
        "start_actions_title": "What would you like to do today?",
        "go_job_assist": "Analyze a Job / Find Jobs",
        "go_cv_documents": "Improve CV / Documents",
        "go_interview": "Practice Interview",
        "go_career_insights": "See Career Insights",
        "start_here_intro": "Your dashboard shows a stable resume analysis, ATS readiness, and extracted profile. Use the menu on the left for deeper tools.",
        "detected_role": "Likely Current Role",
        "detected_summary": "Professional Summary",
        "detected_skills": "Key Skills",
        "suggested_roles": "Suggested Roles",
        "resume_loaded": "Resume Loaded",
        "yes": "Yes",
        "no": "No",
        "not_analyzed": "Not analyzed yet",
        "warn_country": "Please select a country.",
        "warn_upload_cv": "Please upload your CV before continuing.",
        "warn_full_name": "Please enter your full name.",
        "warn_summary": "Please enter your professional summary.",
        "warn_skills": "Please enter your skills.",
        "warn_experience": "Please enter your work experience.",
        "warn_education": "Please enter your education.",
        "unsupported_file": "Unsupported file type.",
        "ai_unavailable": "AI service unavailable",
        "job_desc": "Paste the job description",
        "your_cv": "Your CV",
        "target_job": "Paste target job description",
        "job_title": "Job title",
        "background": "Short background about yourself",
        "preferred_location": "Preferred city or region",
        "target_role": "Target role",
        "current_role": "Your current role",
        "current_profile": "Your current profile / CV",
        "interview_language": "Interview language",
        "language_support_task": "Support task",
        "enter_text": "Enter your text",
        "improve": "Improve",
        "analyze": "Analyze",
        "generate": "Generate",
        "translate": "Translate",
        "job_fit_analysis": "Job Fit Analysis",
        "career_comm_result": "Career Communication Feedback",
        "search_queries": "Suggested search links",
        "platform_hint": "These links open search pages. Later, you can replace them with local job portals.",
        "translator_source": "Source Language",
        "translator_target": "Target Language",
        "cover_letter_input": "Paste your cover letter",
        "cv_input": "Paste your CV",
        "copy_ready": "Copy-ready output",
        "resume_score": "Resume Score",
        "ats_score": "ATS Score",
        "strengths": "Top Strengths",
        "improvements": "Top Improvements",
        "job_fit_score": "Job Fit",
        "skill_gap_score": "Skill Gap",
        "interview_score": "Interview",
        "resume_insights": "Resume Insights",
        "update_resume": "Update My Resume",
        "top_countries_fit": "Country Fit Summary",
        "country_fit_note": "Likely strongest markets based on CV language, profile, and transferability.",
        "detected_country": "Detected Country",
        "refresh_resume": "Re-analyze Resume",
        "save_resume": "Save as My New Resume",
        "document_hub_caption": "Keep all writing and translation tools in one place.",
        "quick_prompts": "Quick prompts",
        "chat_mode": "Mode",
        "chat_placeholder": "Ask Work-O-Bot anything...",
        "clear_chat": "Clear Chat",
        "preferred_language": "Select a Language",
        "user_status": "Current Career Situation",
        "fresh_graduate": "Fresh graduate / entry level",
        "student_thesis_internship": "Student / Thesis / Internship seeker",
        "career_changer": "Career changer",
        "migrant": "Planning to migrate / applying abroad",
        "local_jobseeker": "Looking for jobs locally",
        "experienced": "Experienced professional",
        "returning": "Returning after a career break",
        "guided_cv_builder": "Guided CV Builder",
        "generate_cv": "Generate CV with AI",
        "extra_cv_info": "Would you like to add more information?",
        "country_readiness": "Country Readiness",
        "next_steps": "What should you do next?",
        "cv_template_builder": "Country-Specific Resume Builder",
        "migration_country": "Which country do you want to move/apply to?",
        "chosen_country_resume_score": "Resume readiness for chosen country",
    },
    "German": {
        "title": "Dein 360° KI-Karriereassistent",
        "subtitle": "KI-gestützte Karrierehilfe basierend auf deinem Land und deinem Lebenslauf.",
        "onboarding_title": "Willkommen bei WorkZo",
        "onboarding_subtitle": "Starte mit deinem Land und deinem Lebenslauf. WorkZo erstellt den Rest daraus.",
        "country": "Land auswählen",
        "ui_language": "App-Sprache",
        "response_language": "Antwortsprache der KI",
        "resume_input": "Lebenslauf-Eingabe",
        "upload_cv": "Lebenslauf hochladen",
        "create_cv": "Lebenslauf erstellen",
        "upload_resume": "Lebenslauf hochladen",
        "create_resume": "Lebenslauf erstellen",
        "full_name": "Vollständiger Name",
        "email": "E-Mail",
        "phone": "Telefonnummer",
        "summary": "Berufliche Zusammenfassung",
        "skills": "Kenntnisse",
        "experience": "Berufserfahrung",
        "education": "Ausbildung",
        "continue": "Zum Dashboard",
        "edit_setup": "Onboarding bearbeiten",
        "dashboard": "Dashboard",
        "actions": "Tools erkunden",
        "understand_job": "Stelle verstehen",
        "improve_cv": "Lebenslauf verbessern",
        "career_communication": "Karrierekommunikation",
        "find_jobs": "Jobs finden",
        "mock_interview": "Vorstellungsgespräch",
        "skill_gap": "Kompetenzlücke",
        "career_roadmap": "Karriereplan",
        "job_assist": "Job-Assistent",
        "document_tools": "Dokument-Tools",
        "cv_translator": "Lebenslauf übersetzen",
        "cover_letter_generator": "Anschreiben erstellen",
        "cover_letter_translator": "Anschreiben übersetzen",
        "workobot": "Work-O-Bot",
        "workobot_sub": "KI-Karriere- und Sprachcoach",
        "workobot_desc_short": "Mock-Interviews, Karriere-Einblicke, Kompetenzlücken, Kommunikationstraining und Work-O-Bot Chat.",
        "go_workobot": "Work-O-Bot öffnen",
        "cv_documents": "Lebenslauf & Dokumente",
        "interview_practice": "Interviewtraining",
        "career_insights": "Karriere-Einblicke",
        "navigation_help": "Wo finde ich was?",
        "dashboard_desc_short": "Starte hier: Scores, nächste Schritte und empfohlene Aktionen.",
        "job_assist_desc_short": "Finde Jobs oder verstehe eine Stellenbeschreibung vor der Bewerbung.",
        "cv_documents_desc_short": "Verbessere Lebenslauf, erstelle Anschreiben, übersetze Dokumente und nutze Länder-Vorlagen.",
        "interview_practice_desc_short": "Übe Interviewantworten, Mock-Fragen und mündliche Vorbereitung.",
        "career_insights_desc_short": "Prüfe Kompetenzlücken, Roadmap, Länder-Eignung und nächsten besten Schritt.",
        "preferred_language_help": "Eine Sprache für App, KI-Antworten und Dokumente.",
        "start_actions_title": "Was möchtest du heute tun?",
        "go_job_assist": "Job analysieren / Jobs finden",
        "go_cv_documents": "CV / Dokumente verbessern",
        "go_interview": "Interview üben",
        "go_career_insights": "Karriere-Einblicke ansehen",
        "start_here_intro": "Dein Dashboard zeigt eine stabile Lebenslaufanalyse, ATS-Bereitschaft und extrahierte Profildaten. Nutze das linke Menü für weitere Tools.",
        "detected_role": "Wahrscheinliche aktuelle Rolle",
        "detected_summary": "Berufliche Zusammenfassung",
        "detected_skills": "Wichtige Kenntnisse",
        "suggested_roles": "Vorgeschlagene Rollen",
        "resume_loaded": "Lebenslauf geladen",
        "yes": "Ja",
        "no": "Nein",
        "not_analyzed": "Noch nicht analysiert",
        "warn_country": "Bitte wähle ein Land.",
        "warn_upload_cv": "Bitte lade deinen Lebenslauf hoch.",
        "warn_full_name": "Bitte gib deinen vollständigen Namen ein.",
        "warn_summary": "Bitte gib deine berufliche Zusammenfassung ein.",
        "warn_skills": "Bitte gib deine Kenntnisse ein.",
        "warn_experience": "Bitte gib deine Berufserfahrung ein.",
        "warn_education": "Bitte gib deine Ausbildung ein.",
        "unsupported_file": "Nicht unterstützter Dateityp.",
        "ai_unavailable": "KI-Dienst nicht verfügbar",
        "job_desc": "Stellenbeschreibung einfügen",
        "your_cv": "Dein Lebenslauf",
        "target_job": "Ziel-Stellenbeschreibung einfügen",
        "job_title": "Berufsbezeichnung",
        "background": "Kurzer Hintergrund über dich",
        "preferred_location": "Bevorzugte Stadt oder Region",
        "target_role": "Zielrolle",
        "current_role": "Deine aktuelle Rolle",
        "current_profile": "Dein aktuelles Profil / Lebenslauf",
        "interview_language": "Sprache des Interviews",
        "enter_text": "Text eingeben",
        "improve": "Verbessern",
        "analyze": "Analysieren",
        "generate": "Erstellen",
        "translate": "Übersetzen",
        "job_fit_analysis": "Job-Fit-Analyse",
        "career_comm_result": "Feedback zur Karrierekommunikation",
        "search_queries": "Vorgeschlagene Suchlinks",
        "platform_hint": "Diese Links öffnen Suchseiten. Später kannst du lokale Jobportale nutzen.",
        "translator_source": "Ausgangssprache",
        "translator_target": "Zielsprache",
        "cover_letter_input": "Anschreiben einfügen",
        "cv_input": "Lebenslauf einfügen",
        "copy_ready": "Kopierfertige Ausgabe",
        "resume_score": "Lebenslauf-Score",
        "ats_score": "ATS-Score",
        "strengths": "Top-Stärken",
        "improvements": "Top-Verbesserungen",
        "job_fit_score": "Job-Fit",
        "skill_gap_score": "Kompetenzlücke",
        "interview_score": "Interview",
        "resume_insights": "Lebenslauf-Einblicke",
        "update_resume": "Lebenslauf aktualisieren",
        "top_countries_fit": "Top-Länder für diesen Lebenslauf",
        "country_fit_note": "Wahrscheinlich die stärksten Märkte basierend auf Sprache, Profil und Übertragbarkeit.",
        "detected_country": "Erkanntes Land",
        "refresh_resume": "Lebenslauf neu analysieren",
        "save_resume": "Als neuen Lebenslauf speichern",
        "document_hub_caption": "Alle Schreib- und Übersetzungstools an einem Ort.",
        "quick_prompts": "Schnellstarts",
        "chat_mode": "Modus",
        "chat_placeholder": "Frag Work-O-Bot etwas...",
        "clear_chat": "Chat löschen",
        "preferred_language": "Sprache auswählen",
        "user_status": "Aktuelle Karrieresituation",
        "fresh_graduate": "Absolvent/in / Berufseinsteiger/in",
        "student_thesis_internship": "Student/in / Abschlussarbeit / Praktikum",
        "career_changer": "Quereinsteiger/in",
        "migrant": "Migration / Bewerbung im Ausland geplant",
        "local_jobseeker": "Jobsuche im aktuellen Land",
        "experienced": "Erfahrene Fachkraft",
        "returning": "Rückkehr nach Karrierepause",
        "guided_cv_builder": "Geführter Lebenslauf-Builder",
        "generate_cv": "Lebenslauf mit KI erstellen",
        "extra_cv_info": "Möchtest du weitere Informationen hinzufügen?",
        "country_readiness": "Länder-Eignung",
        "next_steps": "Was solltest du als Nächstes tun?",
        "cv_template_builder": "Länderspezifischer Lebenslauf-Builder",
        "migration_country": "In welches Land möchtest du ziehen/dich bewerben?",
        "chosen_country_resume_score": "Lebenslauf-Eignung für das gewählte Land",
    },
    "Dutch": {
        "title": "Jouw 360° AI-carrièreassistent",
        "subtitle": "AI-ondersteuning gebaseerd op jouw land en cv.",
        "onboarding_title": "Welkom bij WorkZo",
        "onboarding_subtitle": "Begin met je land en cv. WorkZo bouwt de rest daarop.",
        "country": "Selecteer een land",
        "ui_language": "App-taal",
        "response_language": "AI-antwoordtaal",
        "resume_input": "CV-invoer",
        "upload_cv": "CV uploaden",
        "create_cv": "CV maken",
        "upload_resume": "Upload je CV",
        "create_resume": "Maak je CV",
        "full_name": "Volledige naam",
        "email": "E-mail",
        "phone": "Telefoonnummer",
        "summary": "Professionele samenvatting",
        "skills": "Vaardigheden",
        "experience": "Werkervaring",
        "education": "Opleiding",
        "continue": "Ga naar dashboard",
        "edit_setup": "Onboarding bewerken",
        "dashboard": "Dashboard",
        "actions": "Tools verkennen",
        "understand_job": "Vacature begrijpen",
        "improve_cv": "CV verbeteren",
        "career_communication": "Carrièrecommunicatie",
        "find_jobs": "Banen vinden",
        "mock_interview": "Proefinterview",
        "skill_gap": "Vaardigheidskloof",
        "career_roadmap": "Carrièreplan",
        "job_assist": "Jobhulp",
        "document_tools": "Documenttools",
        "cv_translator": "CV vertalen",
        "cover_letter_generator": "Motivatiebrief maken",
        "cover_letter_translator": "Motivatiebrief vertalen",
        "workobot": "Work-O-Bot",
        "workobot_sub": "AI carrière- en taalcoach",
        "workobot_desc_short": "Mock-interviews, carrière-inzichten, skill gaps, communicatie-oefening en Work-O-Bot chat.",
        "go_workobot": "Open Work-O-Bot",
        "cv_documents": "CV & Documenten",
        "interview_practice": "Interview oefenen",
        "career_insights": "Carrière-inzichten",
        "navigation_help": "Waar vind je wat",
        "dashboard_desc_short": "Begin hier: scores, volgende stappen en aanbevolen acties.",
        "job_assist_desc_short": "Vind banen of begrijp een vacaturetekst voordat je solliciteert.",
        "cv_documents_desc_short": "Verbeter je CV, maak brieven, vertaal documenten en gebruik landen-CV-templates.",
        "interview_practice_desc_short": "Oefen interviewantwoorden, mockvragen en spreekvoorbereiding.",
        "career_insights_desc_short": "Bekijk skill gaps, roadmap, landenfit en volgende beste stap.",
        "preferred_language_help": "Eén taal voor app, AI-antwoorden en documenten.",
        "start_actions_title": "Wat wil je vandaag doen?",
        "go_job_assist": "Vacature analyseren / banen vinden",
        "go_cv_documents": "CV / documenten verbeteren",
        "go_interview": "Interview oefenen",
        "go_career_insights": "Carrière-inzichten bekijken",
        "start_here_intro": "Je dashboard toont een stabiele cv-analyse, ATS-gereedheid en geëxtraheerde profielgegevens. Gebruik daarna het linkermenu voor meer tools.",
        "detected_role": "Waarschijnlijke huidige rol",
        "detected_summary": "Professionele samenvatting",
        "detected_skills": "Belangrijkste vaardigheden",
        "suggested_roles": "Voorgestelde rollen",
        "resume_loaded": "CV geladen",
        "yes": "Ja",
        "no": "Nee",
        "not_analyzed": "Nog niet geanalyseerd",
        "warn_country": "Selecteer een land.",
        "warn_upload_cv": "Upload je cv voordat je doorgaat.",
        "warn_full_name": "Voer je volledige naam in.",
        "warn_summary": "Voer je professionele samenvatting in.",
        "warn_skills": "Voer je vaardigheden in.",
        "warn_experience": "Voer je werkervaring in.",
        "warn_education": "Voer je opleiding in.",
        "unsupported_file": "Niet-ondersteund bestandstype.",
        "ai_unavailable": "AI-service niet beschikbaar",
        "job_desc": "Plak de vacaturetekst",
        "your_cv": "Jouw CV",
        "target_job": "Plak de doelvacature",
        "job_title": "Functietitel",
        "background": "Korte achtergrond over jezelf",
        "preferred_location": "Voorkeursstad of regio",
        "target_role": "Doelrol",
        "current_role": "Je huidige rol",
        "current_profile": "Je huidige profiel / CV",
        "interview_language": "Taal van het interview",
        "enter_text": "Voer je tekst in",
        "improve": "Verbeter",
        "analyze": "Analyseer",
        "generate": "Genereer",
        "translate": "Vertaal",
        "job_fit_analysis": "Vacaturematch-analyse",
        "career_comm_result": "Feedback op carrièrecommunicatie",
        "search_queries": "Voorgestelde zoeklinks",
        "platform_hint": "Deze links openen zoekpagina’s. Later kun je lokale vacatureplatforms gebruiken.",
        "translator_source": "Brontaal",
        "translator_target": "Doeltaal",
        "cover_letter_input": "Plak je motivatiebrief",
        "cv_input": "Plak je cv",
        "copy_ready": "Kopieerklare output",
        "resume_score": "CV-score",
        "ats_score": "ATS-score",
        "strengths": "Topsterktes",
        "improvements": "Topverbeteringen",
        "job_fit_score": "Vacaturematch",
        "skill_gap_score": "Vaardigheidskloof",
        "interview_score": "Interview",
        "resume_insights": "CV-inzichten",
        "update_resume": "CV bijwerken",
        "top_countries_fit": "Toplanden voor dit CV",
        "country_fit_note": "Waarschijnlijk de sterkste markten op basis van taal, profiel en overdraagbaarheid.",
        "detected_country": "Gedetecteerd land",
        "refresh_resume": "CV opnieuw analyseren",
        "save_resume": "Opslaan als nieuw CV",
        "document_hub_caption": "Houd al je schrijf- en vertaaltools op één plek.",
        "quick_prompts": "Snelle prompts",
        "chat_mode": "Modus",
        "chat_placeholder": "Vraag Work-O-Bot iets...",
        "clear_chat": "Chat wissen",
    }
}

def ui_lang() -> str:
    current = st.session_state.get("preferred_language", st.session_state.get("ui_language", "English"))
    return current if current in UI_TEXT else "English"

def txt(key: str) -> str:
    return UI_TEXT[ui_lang()].get(key, key)


# =========================================================
# STUDENT / THESIS / INTERNSHIP CAREER PATH
# =========================================================
STUDENT_STATUS_INTERNAL = "Student / Thesis / Internship seeker"

STUDENT_JOB_KEYWORDS_BY_COUNTRY = {
    "germany": ["Praktikum", "Werkstudent", "Abschlussarbeit", "Bachelorarbeit", "Masterarbeit", "Thesis", "Trainee"],
    "austria": ["Praktikum", "Werkstudent", "Abschlussarbeit", "Bachelorarbeit", "Masterarbeit", "Trainee"],
    "switzerland": ["Praktikum", "Werkstudent", "Internship", "Thesis", "Trainee"],
    "netherlands": ["Stage", "Afstudeerstage", "Werkstudent", "Internship", "Traineeship"],
    "the netherlands": ["Stage", "Afstudeerstage", "Werkstudent", "Internship", "Traineeship"],
    "belgium": ["Stage", "Internship", "Student job", "Thesis", "Traineeship"],
    "france": ["Stage", "Alternance", "Apprentissage", "Internship", "Trainee"],
    "spain": ["Prácticas", "Becario", "Internship", "Trainee"],
    "italy": ["Tirocinio", "Stage", "Internship", "Trainee"],
    "portugal": ["Estágio", "Internship", "Trainee"],
    "united kingdom": ["Placement Year", "Internship", "Graduate Intern", "Industrial Placement", "Sandwich Placement"],
    "uk": ["Placement Year", "Internship", "Graduate Intern", "Industrial Placement", "Sandwich Placement"],
    "ireland": ["Internship", "Graduate Intern", "Placement", "Trainee"],
    "united states": ["Internship", "Co-op", "Student Intern", "New Grad", "Campus"],
    "usa": ["Internship", "Co-op", "Student Intern", "New Grad", "Campus"],
    "canada": ["Internship", "Co-op", "Student Intern", "New Grad", "Campus"],
    "india": ["Internship", "Trainee", "Fresher Internship", "Graduate Trainee", "Campus"],
    "australia": ["Internship", "Vacation Program", "Graduate Program", "Student Intern"],
    "new zealand": ["Internship", "Graduate Program", "Student Intern"],
    "singapore": ["Internship", "Traineeship", "Graduate Intern", "Student Intern"],
    "default": ["Internship", "Student Intern", "Thesis", "Placement", "Working Student", "Trainee", "Graduate Intern"],
}

def is_student_thesis_status(user_status: str) -> bool:
    status = (user_status or "").lower()
    return any(x in status for x in ["student / thesis", "thesis", "internship seeker", "praktikum", "placement", "afstudeer", "abschlussarbeit", "werkstudent", "stagezoeker"])

def get_student_job_keywords(country_name: str) -> List[str]:
    country = (country_name or st.session_state.get("migration_country") or st.session_state.get("country") or "").strip().lower()
    return STUDENT_JOB_KEYWORDS_BY_COUNTRY.get(country, STUDENT_JOB_KEYWORDS_BY_COUNTRY["default"])

def render_student_opportunity_guidance(country_name: str):
    keywords = get_student_job_keywords(country_name)
    st.markdown(f"### {txt('student_opportunities')}")
    st.caption(txt("student_guidance_caption"))
    c1, c2 = st.columns(2)
    with c1:
        st.markdown(f"**{txt('student_search_keywords')}**")
        st.markdown(" ".join([f"<span class='pill'>{html.escape(k)}</span>" for k in keywords[:8]]), unsafe_allow_html=True)
    with c2:
        st.markdown(f"**{txt('student_cv_tips')}**")
        st.markdown("- Put education, thesis topic, and projects near the top.\n- Add tools, coursework, and university/research projects.\n- Use student-friendly role titles instead of only full-time junior roles.\n- Keep the CV focused on learning potential and practical proof.")

LANG_EXTRA = {
    "English": {
        "navigation": "Navigation", "workflow_progress": "Workflow progress", "suggested_workflow": "Suggested workflow",
        "workflow_upload_cv": "1. Upload CV", "workflow_understand_job": "2. Understand a job", "workflow_improve_cv": "3. Improve CV", "workflow_apply_smarter": "4. Apply smarter",
        "career_command_center": "Your Career Dashboard", "career_move_organized": "Everything you need to plan your next career step",
        "country_label": "Country", "status_label": "Status", "not_specified": "Not specified",
        "metric_resume_quality": "Overall quality and clarity", "metric_ats_friendly": "How scanner-friendly your resume looks", "metric_detected_resume": "Detected from your current resume",
        "target_roles": "Target roles", "metric_role_cluster": "Suggested role cluster options", "what_next": "What you should do next",
        "next_default_1": "Improve the weakest CV sections first.", "next_default_2": "Use Job Assist after your CV is ready.", "next_default_3": "Tailor the CV for each job description.",
        "next_migrate_1": "Adapt your CV to the target-country format.", "next_migrate_2": "Add local role keywords and remove country-inappropriate details.", "next_migrate_3": "Use Document Tools → Country CV Template before applying.",
        "next_graduate_1": "Add 1–2 portfolio projects.", "next_graduate_2": "Highlight tools, coursework, internships, and projects.", "next_graduate_3": "Target entry-level, trainee, or junior roles.",
        "next_changer_1": "Connect your previous experience to the target role.", "next_changer_2": "Add proof projects and bridge skills.", "next_changer_3": "Avoid applying directly to senior roles in the new field.",
        "country_cv_readiness": "Country CV readiness", "status_guidance": "Status guidance", "target_country_label": "Target country", "judge_market": "WorkZo will judge your CV against this market.", "use_country_template": "Use the Country CV Template tool if the format does not match the target country.",
        "job_assist_desc": "Understand a role, find live jobs, and see where you realistically fit before applying.", "document_tools_desc": "Improve your CV, generate a cover letter, translate documents, and update your resume faster.", "workobot_desc": "Country-aware coaching for interviews, communication, mock questions, and next steps.",
        "standout_next_best_step": "Standout feature: Next Best Step", "next_best_caption": "A personalized action center based on your resume, country, and role direction.", "generate_next_best_step": "Generate My Next Best Step", "building_action_center": "Building your action center...", "refreshing": "Refreshing...", "dashboard_updated": "Dashboard updated.",
        "next_best_empty": "Generate a personalized action center to see your best immediate goal, weekly priorities, role cluster, and the message you should send today.",
        "founder_dashboard": "Founder Dashboard", "founder_access": "Founder access", "founder_pin": "Founder PIN", "open_job": "Open job", "job_summary_first": "Top matches first. Click any job card to open the posting.", "details_later": "Show details", "country_fit_summary": "Country fit summary", "country_fit_details": "Country-fit details", "view_details": "View details", "no_founder_pin": "Add FOUNDER_PIN in secrets to unlock founder analytics.", "recommended_next_step": "Recommended Next Step", "next_action_upload_title": "Upload your CV first", "next_action_upload_desc": "WorkZo needs your CV to calculate scores and guide your next career step.", "next_action_upload_button": "Upload CV", "next_action_improve_title": "Improve your CV for ATS", "next_action_improve_desc": "Your ATS score can improve. Start by strengthening keywords, structure, and country-specific formatting.", "next_action_improve_button": "Improve My CV", "next_action_job_title": "Analyze a job before applying", "next_action_job_desc": "Your CV is ready. Paste a job description or find matching jobs to understand your fit.", "next_action_job_button": "Go to Job Assist", "next_action_interview_title": "Practice for your next interview", "next_action_interview_desc": "You have started job preparation. Now practice answers based on your profile and target role.", "next_action_interview_button": "Practice Interview"
    },
    "German": {
        "navigation": "Navigation", "workflow_progress": "Workflow-Fortschritt", "suggested_workflow": "Empfohlener Ablauf",
        "workflow_upload_cv": "1. Lebenslauf hochladen", "workflow_understand_job": "2. Stelle verstehen", "workflow_improve_cv": "3. Lebenslauf verbessern", "workflow_apply_smarter": "4. Gezielter bewerben",
        "career_command_center": "Karriere-Kommandozentrale", "career_move_organized": "Dein nächster Karriereschritt, klar organisiert",
        "country_label": "Land", "status_label": "Status", "not_specified": "Nicht angegeben",
        "metric_resume_quality": "Gesamtqualität und Klarheit", "metric_ats_friendly": "Wie ATS-freundlich dein Lebenslauf wirkt", "metric_detected_resume": "Aus deinem aktuellen Lebenslauf erkannt",
        "target_roles": "Zielrollen", "metric_role_cluster": "Vorgeschlagene Rollen-Cluster", "what_next": "Was du als Nächstes tun solltest",
        "next_default_1": "Verbessere zuerst die schwächsten Lebenslaufbereiche.", "next_default_2": "Nutze den Job-Assistenten, sobald dein Lebenslauf bereit ist.", "next_default_3": "Passe den Lebenslauf an jede Stellenbeschreibung an.",
        "next_migrate_1": "Passe deinen Lebenslauf an das Format des Ziellandes an.", "next_migrate_2": "Ergänze lokale Rollen-Keywords und entferne unpassende Angaben.", "next_migrate_3": "Nutze vor der Bewerbung Dokument-Tools → Lebenslauf-Vorlage nach Land.",
        "next_graduate_1": "Füge 1–2 Portfolio-Projekte hinzu.", "next_graduate_2": "Betone Tools, Kurse, Praktika und Projekte.", "next_graduate_3": "Bewirb dich gezielt auf Einstiegs-, Trainee- oder Junior-Rollen.",
        "next_changer_1": "Verbinde deine bisherige Erfahrung klar mit der Zielrolle.", "next_changer_2": "Füge Nachweisprojekte und Brückenkompetenzen hinzu.", "next_changer_3": "Bewirb dich im neuen Bereich nicht direkt auf Senior-Rollen.",
        "country_cv_readiness": "Lebenslauf-Eignung für das Zielland", "status_guidance": "Hinweise zu deiner Situation", "target_country_label": "Zielland", "judge_market": "WorkZo bewertet deinen Lebenslauf für diesen Arbeitsmarkt.", "use_country_template": "Nutze die Lebenslauf-Vorlage nach Land, wenn das Format nicht zum Zielland passt.",
        "job_assist_desc": "Verstehe eine Rolle, finde Live-Jobs und erkenne realistisch, wo du vor der Bewerbung passt.", "document_tools_desc": "Verbessere deinen Lebenslauf, erstelle ein Anschreiben, übersetze Dokumente und aktualisiere deine Bewerbung schneller.", "workobot_desc": "Länderspezifisches Coaching für Interviews, Kommunikation, Übungsfragen und nächste Schritte.",
        "standout_next_best_step": "Besondere Funktion: Nächster bester Schritt", "next_best_caption": "Ein personalisiertes Aktionszentrum basierend auf deinem Lebenslauf, Land und deiner Rollenrichtung.", "generate_next_best_step": "Meinen nächsten besten Schritt erstellen", "building_action_center": "Aktionszentrum wird erstellt...", "refreshing": "Aktualisiere...", "dashboard_updated": "Dashboard aktualisiert.",
        "next_best_empty": "Erstelle ein personalisiertes Aktionszentrum, um dein wichtigstes Sofortziel, Wochenprioritäten, Rollen-Cluster und die heutige Nachricht zu sehen.",
        "founder_dashboard": "Founder-Dashboard", "founder_access": "Founder-Zugang", "founder_pin": "Founder-PIN", "open_job": "Job öffnen", "job_summary_first": "Beste Treffer zuerst. Klicke auf eine Jobkarte, um die Anzeige zu öffnen.", "details_later": "Details anzeigen", "country_fit_summary": "Beste Länder für diesen Lebenslauf", "country_fit_details": "Details zur Länder-Eignung", "view_details": "Details anzeigen", "no_founder_pin": "Füge FOUNDER_PIN in Secrets hinzu, um Founder Analytics zu öffnen.", "recommended_next_step": "Empfohlener nächster Schritt", "next_action_upload_title": "Lade zuerst deinen Lebenslauf hoch", "next_action_upload_desc": "WorkZo braucht deinen Lebenslauf, um Scores zu berechnen und deinen nächsten Karriereschritt zu empfehlen.", "next_action_upload_button": "Lebenslauf hochladen", "next_action_improve_title": "Verbessere deinen Lebenslauf für ATS", "next_action_improve_desc": "Dein ATS-Score kann besser werden. Stärke zuerst Keywords, Struktur und länderspezifisches Format.", "next_action_improve_button": "Lebenslauf verbessern", "next_action_job_title": "Analysiere eine Stelle vor der Bewerbung", "next_action_job_desc": "Dein Lebenslauf ist bereit. Füge eine Stellenbeschreibung ein oder finde passende Jobs.", "next_action_job_button": "Zum Job-Assistenten", "next_action_interview_title": "Übe für dein nächstes Interview", "next_action_interview_desc": "Du hast mit der Bewerbungsvorbereitung begonnen. Übe jetzt Antworten passend zu Profil und Zielrolle.", "next_action_interview_button": "Interview üben"
    },
    "Dutch": {
        "navigation": "Navigatie", "workflow_progress": "Workflowvoortgang", "suggested_workflow": "Aanbevolen workflow",
        "workflow_upload_cv": "1. CV uploaden", "workflow_understand_job": "2. Vacature begrijpen", "workflow_improve_cv": "3. CV verbeteren", "workflow_apply_smarter": "4. Slimmer solliciteren",
        "career_command_center": "Carrière-commandocentrum", "career_move_organized": "Je volgende carrièremove, helder georganiseerd",
        "country_label": "Land", "status_label": "Status", "not_specified": "Niet opgegeven",
        "metric_resume_quality": "Algemene kwaliteit en duidelijkheid", "metric_ats_friendly": "Hoe ATS-vriendelijk je cv is", "metric_detected_resume": "Gedetecteerd uit je huidige cv",
        "target_roles": "Doelrollen", "metric_role_cluster": "Voorgestelde rolclusters", "what_next": "Wat je nu moet doen",
        "next_default_1": "Verbeter eerst de zwakste cv-onderdelen.", "next_default_2": "Gebruik Jobhulp zodra je cv klaar is.", "next_default_3": "Pas je cv aan voor elke vacaturetekst.",
        "next_migrate_1": "Pas je cv aan aan het format van het doelland.", "next_migrate_2": "Voeg lokale rolkeywords toe en verwijder ongepaste gegevens.", "next_migrate_3": "Gebruik Documenttools → Country CV Template voordat je solliciteert.",
        "next_graduate_1": "Voeg 1–2 portfolio-projecten toe.", "next_graduate_2": "Benadruk tools, cursussen, stages en projecten.", "next_graduate_3": "Richt je op starters-, trainee- of juniorrollen.",
        "next_changer_1": "Koppel je eerdere ervaring aan de doelrol.", "next_changer_2": "Voeg bewijsprojecten en brugvaardigheden toe.", "next_changer_3": "Solliciteer niet direct op seniorrollen in het nieuwe vakgebied.",
        "country_cv_readiness": "CV-gereedheid voor doelland", "status_guidance": "Advies voor je situatie", "target_country_label": "Doelland", "judge_market": "WorkZo beoordeelt je cv voor deze arbeidsmarkt.", "use_country_template": "Gebruik de Country CV Template-tool als het format niet past bij het doelland.",
        "job_assist_desc": "Begrijp een rol, vind live vacatures en zie realistisch waar je past voordat je solliciteert.", "document_tools_desc": "Verbeter je cv, maak een motivatiebrief, vertaal documenten en werk je sollicitatie sneller bij.", "workobot_desc": "Landbewuste coaching voor interviews, communicatie, oefenvragen en vervolgstappen.",
        "standout_next_best_step": "Sterke functie: Volgende beste stap", "next_best_caption": "Een persoonlijk actiecentrum gebaseerd op je cv, land en rolrichting.", "generate_next_best_step": "Maak mijn volgende beste stap", "building_action_center": "Actiecentrum wordt gemaakt...", "refreshing": "Vernieuwen...", "dashboard_updated": "Dashboard bijgewerkt.",
        "next_best_empty": "Maak een persoonlijk actiecentrum om je beste directe doel, weekprioriteiten, rolcluster en het bericht dat je vandaag moet sturen te zien.",
        "founder_dashboard": "Founder dashboard",
        "founder_access": "Founder access",
        "founder_pin": "Founder PIN",
        "open_job": "Open job",
        "job_summary_first": "Best matches first. Click a job card to open it.",
        "details_later": "Show details",
        "country_fit_summary": "Country fit",
        "country_fit_details": "Country fit details",
        "view_details": "View details",
        "no_founder_pin": "Add FOUNDER_PIN in secrets to unlock founder analytics.",
        "recommended_next_step": "Recommended next step",
        "next_action_upload_title": "Upload your CV first",
        "next_action_upload_desc": "WorkZo needs your CV to calculate scores and guide your next career step.",
        "next_action_upload_button": "Upload CV",
        "next_action_improve_title": "Improve your CV for ATS",
        "next_action_improve_desc": "Your ATS score can improve. Start with keywords, structure and country-specific formatting.",
        "next_action_improve_button": "Improve my CV",
        "next_action_job_title": "Analyze a job before applying",
        "next_action_job_desc": "Your CV is ready. Paste a job description or find matching jobs.",
        "next_action_job_button": "Go to Job Help",
        "next_action_interview_title": "Practice for your next interview",
        "next_action_interview_desc": "Practice answers based on your profile and target role.",
        "next_action_interview_button": "Practice interview"
    }
}
for _lang, _items in LANG_EXTRA.items():
    UI_TEXT.setdefault(_lang, {}).update(_items)

# =========================================================
# UI LOCALIZATION CLEANUP PATCH (WorkZo v8.1)
# Fixes mixed English/German/Dutch labels in onboarding and CV builder.
# =========================================================
UI_LOCALIZATION_PATCH = {
    "English": {
        "select_optional": "Select one option",
        "beta_privacy_note": "⚠️ WorkZo AI is currently in Beta. Some results may not be perfect yet. Your feedback helps improve the tool. Privacy note: WorkZo collects anonymous usage data such as features used, selected country/status, and session duration. It does not store your CV text, email, phone number, address, or personal documents.",
        "language_help": "This language is used for the app, AI replies, and generated documents.",
        "status_help": "Optional. Select this only if it matches your current situation.",
        "resume_choice_caption": "Choose one option to add your resume. You can upload an existing CV or create one with the guided builder.",
        "upload_cv_title": "Upload CV",
        "upload_cv_desc": "Best if you already have a PDF or text CV.",
        "create_cv_title": "Create CV",
        "create_cv_desc": "Best if you want WorkZo to build a CV from your details.",
        "choose_resume_continue": "Please choose Upload CV or Create CV to continue.",
        "file_too_large": "File too large. Please upload a CV under 5 MB.",
        "upload_cv_instruction": "Please upload a PDF or TXT CV using the Upload CV option above.",
        "guided_cv_caption": "Add what you know. WorkZo will generate a cleaner CV and ask what is missing.",
        "location": "Location",
        "languages_label": "Languages",
        "cert_courses": "Certifications / Courses",
        "projects_label": "Projects",
        "pdf_read_error": "PDF read error",
    },
    "German": {
        "select_optional": "Option auswählen",
        "beta_privacy_note": "⚠️ WorkZo AI ist aktuell in der Beta-Version. Manche Ergebnisse sind möglicherweise noch nicht perfekt. Dein Feedback hilft, das Tool zu verbessern. Datenschutzhinweis: WorkZo erfasst anonyme Nutzungsdaten wie verwendete Funktionen, ausgewähltes Land/Status und Sitzungsdauer. Dein Lebenslauftext, deine E-Mail, Telefonnummer, Adresse oder persönliche Dokumente werden nicht gespeichert.",
        "language_help": "Diese Sprache wird für die App, KI-Antworten und erstellte Dokumente verwendet.",
        "status_help": "Optional. Wähle dies nur aus, wenn es zu deiner aktuellen Situation passt.",
        "resume_choice_caption": "Wähle eine Option aus, um deinen Lebenslauf hinzuzufügen. Du kannst einen vorhandenen Lebenslauf hochladen oder mit dem geführten Builder einen neuen erstellen.",
        "upload_cv_title": "Lebenslauf hochladen",
        "upload_cv_desc": "Ideal, wenn du bereits einen PDF- oder Text-Lebenslauf hast.",
        "create_cv_title": "Lebenslauf erstellen",
        "create_cv_desc": "Ideal, wenn WorkZo aus deinen Angaben einen Lebenslauf erstellen soll.",
        "choose_resume_continue": "Bitte wähle Lebenslauf hochladen oder Lebenslauf erstellen, um fortzufahren.",
        "file_too_large": "Die Datei ist zu groß. Bitte lade einen Lebenslauf unter 5 MB hoch.",
        "upload_cv_instruction": "Bitte lade oben über die Option Lebenslauf hochladen eine PDF- oder TXT-Datei hoch.",
        "guided_cv_caption": "Füge hinzu, was du weißt. WorkZo erstellt daraus einen klareren Lebenslauf und ergänzt fehlende Punkte.",
        "location": "Standort",
        "languages_label": "Sprachen",
        "cert_courses": "Zertifikate / Kurse",
        "projects_label": "Projekte",
        "pdf_read_error": "PDF-Lesefehler",
        "career_command_center": "Dein Karriere-Dashboard",
        "career_move_organized": "Alles, was du für deinen nächsten Karriereschritt brauchst",
    },
    "Dutch": {
        "preferred_language": "Taal selecteren",
        "user_status": "Huidige carrièresituatie",
        "fresh_graduate": "Afgestudeerd / starter",
        "student_thesis_internship": "Student / scriptie / stagezoeker",
        "career_changer": "Carrièreswitcher",
        "migrant": "Verhuizen / solliciteren in het buitenland",
        "local_jobseeker": "Lokaal werk zoeken",
        "experienced": "Ervaren professional",
        "returning": "Terugkeer na loopbaanpauze",
        "guided_cv_builder": "Begeleide CV-builder",
        "generate_cv": "CV genereren met AI",
        "extra_cv_info": "Wil je extra informatie toevoegen?",
        "country_readiness": "Landgeschiktheid",
        "next_steps": "Wat moet je nu doen?",
        "cv_template_builder": "Landspecifieke CV-builder",
        "migration_country": "Naar welk land wil je verhuizen/solliciteren?",
        "chosen_country_resume_score": "CV-gereedheid voor gekozen land",
        "select_optional": "Selecteer een optie",
        "beta_privacy_note": "⚠️ WorkZo AI is momenteel in bèta. Sommige resultaten zijn mogelijk nog niet perfect. Jouw feedback helpt om de tool te verbeteren. Privacyverklaring: WorkZo verzamelt anonieme gebruiksgegevens zoals gebruikte functies, gekozen land/status en sessieduur. Je CV-tekst, e-mail, telefoonnummer, adres of persoonlijke documenten worden niet opgeslagen.",
        "language_help": "Deze taal wordt gebruikt voor de app, AI-antwoorden en gegenereerde documenten.",
        "status_help": "Optioneel. Selecteer dit alleen als het past bij je huidige situatie.",
        "resume_choice_caption": "Kies één optie om je CV toe te voegen. Je kunt een bestaand CV uploaden of er een maken met de begeleide builder.",
        "upload_cv_title": "CV uploaden",
        "upload_cv_desc": "Beste keuze als je al een PDF- of tekst-CV hebt.",
        "create_cv_title": "CV maken",
        "create_cv_desc": "Beste keuze als je wilt dat WorkZo een CV maakt op basis van jouw gegevens.",
        "choose_resume_continue": "Kies CV uploaden of CV maken om door te gaan.",
        "file_too_large": "Bestand is te groot. Upload een CV kleiner dan 5 MB.",
        "upload_cv_instruction": "Upload hierboven een PDF- of TXT-CV via de optie CV uploaden.",
        "guided_cv_caption": "Voeg toe wat je weet. WorkZo maakt er een duidelijker CV van en vult ontbrekende punten aan.",
        "location": "Locatie",
        "languages_label": "Talen",
        "cert_courses": "Certificaten / Cursussen",
        "projects_label": "Projecten",
        "pdf_read_error": "PDF-leesfout",
    },
}
for _lang, _items in UI_LOCALIZATION_PATCH.items():
    UI_TEXT.setdefault(_lang, {}).update(_items)

# =========================================================
# DYNAMIC COUNTRY + LANGUAGE DATA
# =========================================================
def get_country_options() -> List[str]:
    if pycountry:
        countries = sorted({c.name for c in pycountry.countries if getattr(c, "name", None)})
    else:
        countries = [
            "Austria", "Canada", "France", "Germany", "India", "Netherlands",
            "United Kingdom", "United States"
        ]
    return countries

def get_language_options() -> List[str]:
    if pycountry:
        languages = sorted({
            getattr(lang, "name", "").strip()
            for lang in pycountry.languages
            if getattr(lang, "name", None)
            and len(getattr(lang, "name", "")) > 1
            and "sign language" not in getattr(lang, "name", "").lower()
        })
    else:
        languages = ["Dutch", "English", "French", "German", "Hindi"]
    return languages

def get_geo_defaults() -> Tuple[str, str]:
    """
    Best-effort IP-based suggestion. Gracefully falls back.
    """
    if "geo_country_suggestion" in st.session_state and "geo_language_suggestion" in st.session_state:
        return st.session_state.geo_country_suggestion, st.session_state.geo_language_suggestion

    fallback_country = "Germany"
    fallback_language = "English"

    try:
        req = urllib.request.Request(
            "https://ipapi.co/json/",
            headers={"User-Agent": "Mozilla/5.0"}
        )
        with urllib.request.urlopen(req, timeout=3) as response:
            payload = json.loads(response.read().decode("utf-8"))
            country_name = payload.get("country_name") or fallback_country

            languages_raw = payload.get("languages", "")
            first_lang_code = languages_raw.split(",")[0].split("-")[0].strip() if languages_raw else ""

            language_name = fallback_language
            if pycountry and first_lang_code:
                lang_obj = pycountry.languages.get(alpha_2=first_lang_code)
                if lang_obj and getattr(lang_obj, "name", None):
                    language_name = lang_obj.name

            st.session_state.geo_country_suggestion = country_name
            st.session_state.geo_language_suggestion = language_name
            return country_name, language_name
    except Exception:
        st.session_state.geo_country_suggestion = fallback_country
        st.session_state.geo_language_suggestion = fallback_language
        return fallback_country, fallback_language



def get_country_code(country_name: str) -> str:
    if not country_name or not pycountry:
        return ""
    try:
        obj = pycountry.countries.get(name=country_name)
        if obj and getattr(obj, "alpha_2", None):
            return obj.alpha_2
    except Exception:
        pass

    try:
        matches = pycountry.countries.search_fuzzy(country_name)
        if matches and getattr(matches[0], "alpha_2", None):
            return matches[0].alpha_2
    except Exception:
        pass
    return ""

@st.cache_data(show_spinner=False)
def get_local_city_index() -> Dict[str, List[str]]:
    """
    Local city index for fast suggestions. Uses geonamescache when available.
    """
    index: Dict[str, List[Tuple[str, int]]] = {}
    if not geonamescache:
        return {}

    try:
        gc = geonamescache.GeonamesCache(min_city_population=5000)
    except TypeError:
        gc = geonamescache.GeonamesCache()

    cities = gc.get_cities()

    for city in cities.values():
        name = (city.get("name") or "").strip()
        country_code = (city.get("countrycode") or "").strip().upper()
        population = int(city.get("population") or 0)

        if not name or not country_code:
            continue
        if population < 5000:
            continue

        index.setdefault(country_code, []).append((name, population))

    final_index: Dict[str, List[str]] = {}
    for code, rows in index.items():
        rows = sorted(rows, key=lambda x: (-x[1], x[0].lower()))
        seen = set()
        names = []
        for name, _ in rows:
            k = name.casefold()
            if k not in seen:
                seen.add(k)
                names.append(name)
        final_index[code] = names
    return final_index

@st.cache_data(show_spinner=False, ttl=86400)
def fetch_country_cities(country_name: str) -> List[str]:
    """
    Country-level city list with two layers:
    1) local geonamescache data when installed
    2) CountriesNow public API fallback for global coverage
    """
    country_code = get_country_code(country_name)
    merged: List[str] = []

    local_index = get_local_city_index()
    if country_code and country_code in local_index:
        merged.extend(local_index[country_code])

    urls = [
        "https://countriesnow.space/api/v0.1/countries/cities/q?country=" + urllib.parse.quote(country_name),
        "https://countriesnow.space/api/v0.1/countries/cities?country=" + urllib.parse.quote(country_name),
    ]

    for url in urls:
        try:
            req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
            with urllib.request.urlopen(req, timeout=8) as response:
                payload = json.loads(response.read().decode("utf-8"))
                data = payload.get("data", [])
                if isinstance(data, list):
                    merged.extend([str(x).strip() for x in data if str(x).strip()])
                if merged:
                    break
        except Exception:
            continue

    deduped = []
    seen = set()
    for city in merged:
        k = city.casefold()
        if k not in seen:
            seen.add(k)
            deduped.append(city)

    return deduped

def fetch_city_suggestions(query: str, country_name: str = "", limit: int = 20) -> List[str]:
    query = (query or "").strip()
    cities = fetch_country_cities(country_name) if country_name else []

    if not cities:
        return []

    if not query:
        return cities[:limit]

    q = query.casefold()
    starts = [c for c in cities if c.casefold().startswith(q)]
    contains = [c for c in cities if q in c.casefold() and c not in starts]
    return (starts + contains)[:limit]

def http_get_json(url: str, headers: Optional[Dict[str, str]] = None, timeout: int = 12) -> Dict:
    req = urllib.request.Request(url, headers=headers or {"User-Agent": "Mozilla/5.0"})
    with urllib.request.urlopen(req, timeout=timeout) as response:
        return json.loads(response.read().decode("utf-8"))

@st.cache_data(show_spinner=False, ttl=900)
def fetch_arbeitnow_jobs(query: str, location: str = "", limit: int = 15) -> List[Dict]:
    url = "https://www.arbeitnow.com/api/job-board-api"
    try:
        payload = http_get_json(url)
    except Exception:
        return []

    items = payload.get("data", []) if isinstance(payload, dict) else []
    q = (query or "").casefold()
    loc = (location or "").casefold()
    if loc.startswith("anywhere in "):
        loc = ""
    results = []

    for item in items:
        title = str(item.get("title", "")).strip()
        company = str(item.get("company_name", "")).strip()
        job_location = str(item.get("location", "")).strip()
        description = str(item.get("description", "")).strip()
        tags = ", ".join(item.get("tags", [])[:5]) if isinstance(item.get("tags"), list) else ""
        remote = bool(item.get("remote", False))
        url_apply = item.get("url") or item.get("job_url") or ""

        haystack = " ".join([title, company, job_location, description, tags]).casefold()
        if q and q not in haystack:
            continue
        if loc and loc not in haystack and not remote:
            continue

        results.append({
            "source": "Arbeitnow",
            "title": title,
            "company": company,
            "location": job_location or ("Remote" if remote else "Germany"),
            "remote": remote,
            "url": url_apply,
            "summary": tags or (description[:220] + "..." if description else "")
        })
        if len(results) >= limit:
            break

    return results

@st.cache_data(show_spinner=False, ttl=900)
def fetch_arbeitsagentur_jobs(query: str, location: str = "", limit: int = 15) -> List[Dict]:
    client_id = (os.getenv("BA_JOBS_API_KEY") or get_streamlit_secret("BA_JOBS_API_KEY") or "c003a37f-024f-462a-b36d-b001be4cd24a")
    normalized_location = "" if str(location or "").lower().startswith("anywhere in ") else (location or "")
    params = {
        "was": query or "",
        "wo": normalized_location,
        "size": str(limit),
    }
    url = "https://jobsuche.api.bund.dev/pc/v4/app/jobs?" + urllib.parse.urlencode(params)

    try:
        payload = http_get_json(url, headers={"X-API-Key": client_id, "User-Agent": "Mozilla/5.0"}, timeout=12)
    except Exception:
        return []

    raw_items = []
    if isinstance(payload, dict):
        for key in ["stellenangebote", "jobOffers", "jobs", "data"]:
            value = payload.get(key)
            if isinstance(value, list):
                raw_items = value
                break

    results = []
    for item in raw_items:
        title = str(item.get("beruf") or item.get("titel") or item.get("title") or "").strip()
        employer = item.get("arbeitgeber") or item.get("company") or {}
        if isinstance(employer, dict):
            company = str(employer.get("name") or employer.get("firma") or "").strip()
        else:
            company = str(employer or "").strip()

        location_value = item.get("arbeitsort") or item.get("arbeitsorte") or item.get("location") or {}
        if isinstance(location_value, list) and location_value:
            first_loc = location_value[0]
            if isinstance(first_loc, dict):
                place = ", ".join([str(first_loc.get("ort") or "").strip(), str(first_loc.get("region") or "").strip()]).strip(", ")
            else:
                place = str(first_loc)
        elif isinstance(location_value, dict):
            place = ", ".join([str(location_value.get("ort") or "").strip(), str(location_value.get("region") or "").strip()]).strip(", ")
        else:
            place = str(location_value or "").strip()

        refnr = str(item.get("refnr") or item.get("referenznummer") or "").strip()
        detail_url = f"https://www.arbeitsagentur.de/jobsuche/jobdetail/{urllib.parse.quote(refnr)}" if refnr else "https://www.arbeitsagentur.de/jobsuche/"
        summary = str(item.get("eintrittsdatum") or item.get("aktuelleVeroeffentlichungsdatum") or item.get("modifikationsTimestamp") or "").strip()

        if title:
            results.append({
                "source": "Bundesagentur für Arbeit",
                "title": title,
                "company": company or "Employer not shown",
                "location": place or (location or "Germany"),
                "remote": False,
                "url": detail_url,
                "summary": summary
            })
        if len(results) >= limit:
            break

    return results

def get_status_job_modifiers(user_status: str, country_name: str = "") -> List[str]:
    status = (user_status or "").lower()
    if is_student_thesis_status(user_status):
        return get_student_job_keywords(country_name)
    if any(x in status for x in ["fresh", "graduate", "absolvent", "entry level"]):
        return ["junior", "entry level", "trainee", "graduate", "internship", "no experience"]
    if any(x in status for x in ["career changer", "quereinsteiger", "changer"]):
        return ["junior", "career changer", "entry level", "trainee", "quereinsteiger"]
    if any(x in status for x in ["apply online", "remote", "online"]):
        return ["remote", "online", "work from home", "hybrid", "junior", "entry level"]
    if any(x in status for x in ["returning", "break", "pause"]):
        return ["returnship", "part time", "junior", "entry level", "back to work"]
    if any(x in status for x in ["experienced", "senior"]):
        return ["experienced", "specialist", "senior"]
    return ["junior", "entry level", "specialist"]

def build_live_job_queries(roles: List[str], user_status: str, max_queries: int = 18, country_name: str = "") -> List[str]:
    base_roles = [r.strip() for r in roles if r and r.strip()]
    if not base_roles:
        base_roles = ["Data Analyst", "IT Support", "Customer Support", "Business Analyst"]
    modifiers = get_status_job_modifiers(user_status, country_name)
    queries: List[str] = []
    for role in base_roles[:6]:
        if role not in queries:
            queries.append(role)
        for modifier in modifiers[:5]:
            candidate = f"{modifier} {role}"
            if candidate not in queries:
                queries.append(candidate)
        for modifier in ["remote", "online", "hybrid"]:
            candidate = f"{role} {modifier}"
            if candidate not in queries:
                queries.append(candidate)
        if len(queries) >= max_queries:
            break
    return queries[:max_queries]

def score_job_for_user(job: Dict, user_status: str, roles: List[str], country_name: str = "") -> int:
    text = " ".join([str(job.get(k, "")) for k in ["title", "summary", "company", "location", "source"]]).lower()
    score = 0
    for role in roles:
        role_words = [w for w in re.findall(r"[a-zA-Z]+", role.lower()) if len(w) > 2]
        score += sum(6 for w in role_words if w in text)
    for mod in get_status_job_modifiers(user_status, country_name):
        if mod.lower() in text:
            score += 12
    status = (user_status or "").lower()
    if any(x in status for x in ["fresh", "graduate", "student", "career changer", "returning"]):
        if any(x in text for x in ["senior", "lead", "principal", "manager"]):
            score -= 20
        if any(x in text for x in ["junior", "entry", "trainee", "intern", "working student", "graduate"]):
            score += 20
    if job.get("remote"):
        score += 5
    return score

def sort_jobs_for_user(jobs: List[Dict], user_status: str, roles: List[str], country_name: str = "") -> List[Dict]:
    return sorted(jobs, key=lambda job: score_job_for_user(job, user_status, roles, country_name), reverse=True)

def fetch_live_jobs_for_germany(roles: List[str], location: str = "", user_status: str = "") -> List[Dict]:
    results: List[Dict] = []
    search_queries = build_live_job_queries(roles, user_status, max_queries=18, country_name="Germany")

    # German job boards often use German role keywords. Add broad local-language fallbacks
    # so live search does not return only 0-3 results for English role titles.
    german_fallback_queries = [
        "Junior", "Quereinsteiger", "Berufseinsteiger", "Trainee", "Praktikum",
        "Werkstudent", "Abschlussarbeit", "Bachelorarbeit", "Masterarbeit", "Thesis",
        "Datenanalyst", "Data Analyst", "Business Analyst", "IT Support",
        "IT Support Mitarbeiter", "Helpdesk", "Kundenservice", "Customer Support",
        "Sachbearbeiter", "Backoffice", "Service Desk"
    ]
    for q in german_fallback_queries:
        if q not in search_queries:
            search_queries.append(q)

    for query in search_queries[:32]:
        results.extend(fetch_arbeitsagentur_jobs(query, location, limit=25))
        results.extend(fetch_arbeitnow_jobs(query, location, limit=25))

    # If a city search is too narrow, broaden once to the whole country.
    if len(results) < 12 and location and location.lower() not in {"germany", "anywhere in germany"}:
        for query in search_queries[:18]:
            results.extend(fetch_arbeitsagentur_jobs(query, "", limit=25))
            results.extend(fetch_arbeitnow_jobs(query, "", limit=25))

    deduped = []
    seen = set()
    for item in results:
        key = (item.get("source", "") + "|" + item.get("title", "") + "|" + item.get("company", "") + "|" + item.get("location", "")).casefold()
        if key not in seen:
            seen.add(key)
            deduped.append(item)
    return sort_jobs_for_user(deduped, user_status, roles, "Germany")[:80]


def render_live_jobs(jobs: List[Dict], country_name: str = "", max_visible: int = 30):
    country_lower = (country_name or "").strip().lower()
    adzuna_ready = bool(os.getenv("ADZUNA_APP_ID") or get_streamlit_secret("ADZUNA_APP_ID")) and bool(os.getenv("ADZUNA_APP_KEY") or get_streamlit_secret("ADZUNA_APP_KEY"))
    if not jobs:
        if country_lower == "germany":
            st.info("No live jobs matched this Germany search right now. Try the whole country, a broader role title, or another status level.")
        elif not adzuna_ready:
            st.info("No live jobs could be loaded for this country because Adzuna is not configured. Add ADZUNA_APP_ID and ADZUNA_APP_KEY in Streamlit secrets to increase live job coverage globally.")
        else:
            st.info("No live jobs matched this exact search right now. Try a broader location, a simpler job title, or the whole country.")
        return

    visible_jobs = jobs[:max_visible]
    title = f"Live job openings in {country_name}" if country_name else "Live job openings"
    st.markdown(f"### {title}")
    st.caption(txt("job_summary_first"))
    st.caption(f"Showing {len(visible_jobs)} of {len(jobs)} matched openings. Results are prioritized by your career status and target role.")

    top_jobs = visible_jobs[:3]
    remaining_jobs = visible_jobs[3:]

    def render_clickable_job_card(job: Dict, compact: bool = False):
        summary = (job.get("summary", "") or "—")
        limit = 130 if compact else 170
        if len(summary) > limit:
            summary = summary[:limit].rsplit(" ", 1)[0] + "..."
        url = job.get("url") or "#"
        title_html = html.escape(str(job.get("title", "Role")))
        company_html = html.escape(str(job.get("company", "Employer not shown")))
        meta_html = html.escape(f"{job.get('location','')} • {job.get('source','')}")
        summary_html = html.escape(summary)
        job_html = f"""
<a href="{html.escape(url, quote=True)}" target="_blank" style="text-decoration:none; color:inherit;">
  <div class="card" style="cursor:pointer; transition:0.15s; border-color:rgba(96,165,250,0.32);">
    <div class="section-title">{title_html}</div>
    <div><strong>{company_html}</strong></div>
    <div class="small-muted">{meta_html}</div>
    <div style="margin-top:8px; color:#cbd5e1;">{summary_html}</div>
    <div style="margin-top:10px; font-size:0.9rem; color:#93c5fd;">↗ {txt('open_job')}</div>
  </div>
</a>
"""
        st.markdown(job_html, unsafe_allow_html=True)

    for job in top_jobs:
        render_clickable_job_card(job)

    if remaining_jobs:
        with st.expander(f"{txt('details_later')} — {len(remaining_jobs)} more jobs", expanded=False):
            for job in remaining_jobs:
                render_clickable_job_card(job, compact=True)


def country_to_indeed_domain(country_name: str) -> str:
    country = (country_name or "").strip().lower()
    domain_map = {
        "germany": "de.indeed.com", "netherlands": "nl.indeed.com", "the netherlands": "nl.indeed.com",
        "united kingdom": "uk.indeed.com", "uk": "uk.indeed.com", "ireland": "ie.indeed.com",
        "united states": "www.indeed.com", "usa": "www.indeed.com", "canada": "ca.indeed.com",
        "india": "in.indeed.com", "australia": "au.indeed.com", "new zealand": "nz.indeed.com",
        "france": "fr.indeed.com", "spain": "es.indeed.com", "italy": "it.indeed.com",
        "austria": "at.indeed.com", "switzerland": "ch.indeed.com", "belgium": "be.indeed.com",
        "sweden": "se.indeed.com", "denmark": "dk.indeed.com", "norway": "no.indeed.com",
        "finland": "fi.indeed.com", "poland": "pl.indeed.com", "singapore": "sg.indeed.com",
        "south africa": "za.indeed.com", "brazil": "br.indeed.com", "mexico": "mx.indeed.com",
        "japan": "jp.indeed.com", "united arab emirates": "ae.indeed.com"
    }
    return domain_map.get(country, "www.indeed.com")


def get_country_linkedin_geo(country_name: str) -> str:
    # LinkedIn works globally without geoId; country name in location is enough for broad matching.
    return urllib.parse.quote(country_name or "")


@st.cache_data(show_spinner=False, ttl=1800)
def fetch_remotive_jobs(query: str, location: str = "", limit: int = 20) -> List[Dict]:
    """Free remote-job API fallback. Useful for global/online job seekers."""
    q = (query or "").strip()
    if not q:
        return []
    url = "https://remotive.com/api/remote-jobs?" + urllib.parse.urlencode({"search": q, "limit": str(limit)})
    try:
        payload = http_get_json(url, timeout=12)
    except Exception:
        return []
    items = payload.get("jobs", []) if isinstance(payload, dict) else []
    results: List[Dict] = []
    for item in items[:limit]:
        title = str(item.get("title", "")).strip()
        company = str(item.get("company_name", "")).strip()
        url_apply = str(item.get("url", "")).strip()
        category = str(item.get("category", "")).strip()
        candidate_required_location = str(item.get("candidate_required_location", "Remote")).strip()
        description = re.sub(r"<[^>]+>", " ", str(item.get("description", "")))
        description = re.sub(r"\s+", " ", description).strip()
        if title:
            results.append({
                "source": "Remotive",
                "title": title,
                "company": company or "Employer not shown",
                "location": candidate_required_location or "Remote / Worldwide",
                "remote": True,
                "url": url_apply,
                "summary": category or (description[:220] + "..." if description else "Remote role")
            })
    return results


def get_global_job_search_query(role: str, country_name: str, location: str, user_status: str = "") -> str:
    role = (role or "jobs").strip()
    loc = (location or country_name or "").strip()
    status = (user_status or "").lower()
    modifiers = []
    if is_student_thesis_status(user_status):
        modifiers.extend(get_student_job_keywords(country_name)[:3])
    elif any(x in status for x in ["fresh", "graduate"]):
        modifiers.extend(["entry level", "junior", "graduate"])
    elif any(x in status for x in ["career changer", "quereinsteiger"]):
        modifiers.extend(["career changer", "junior"])
    elif any(x in status for x in ["online", "remote"]):
        modifiers.extend(["remote", "online"])
    elif any(x in status for x in ["experienced", "senior"]):
        modifiers.extend(["experienced"])
    prefix = " ".join(modifiers[:2])
    return " ".join([prefix, role, "jobs", loc]).strip()

def get_job_board_links(country_name: str, location: str, role: str = "", user_status: str = "") -> List[Tuple[str, str]]:

    country_name = (country_name or "").strip()
    location = (location or country_name or "").strip()
    role = (role or "jobs").strip()
    country = country_name.lower()
    q_text = get_global_job_search_query(role, country_name, location, user_status)
    q = urllib.parse.quote(q_text)
    loc_q = urllib.parse.quote(location or country_name)
    role_q = urllib.parse.quote(role)
    indeed_domain = country_to_indeed_domain(country_name)

    boards: List[Tuple[str, str]] = [
        ("LinkedIn Jobs", f"https://www.linkedin.com/jobs/search/?keywords={q}&location={loc_q}"),
        ("Indeed", f"https://{indeed_domain}/jobs?q={q}&l={loc_q}"),
        ("Google Jobs", f"https://www.google.com/search?q={q}"),
    ]

    if country == "germany":
        boards.extend([
            ("StepStone Germany", f"https://www.stepstone.de/jobs/{role_q}/in-{loc_q}"),
            ("XING Jobs", f"https://www.xing.com/jobs/search?keywords={role_q}&location={loc_q}"),
            ("Bundesagentur für Arbeit", f"https://www.arbeitsagentur.de/jobsuche/suche?was={role_q}&wo={loc_q}"),
        ])
    elif country in {"netherlands", "the netherlands"}:
        boards.extend([
            ("National Vacaturebank", f"https://www.nationalevacaturebank.nl/vacatures/zoekterm/{role_q}"),
            ("Werk.nl", f"https://www.werk.nl/werkzoekenden/vacatures/?q={role_q}"),
            ("Iamexpat Jobs", f"https://www.iamexpat.nl/career/jobs-netherlands?search={role_q}"),
        ])
    elif country in {"united kingdom", "uk"}:
        boards.extend([
            ("Reed", f"https://www.reed.co.uk/jobs/{role_q}-jobs-in-{loc_q}"),
            ("Totaljobs", f"https://www.totaljobs.com/jobs/{role_q}/in-{loc_q}"),
            ("CV-Library", f"https://www.cv-library.co.uk/{role_q}-jobs-in-{loc_q}"),
        ])
    elif country in {"united states", "usa"}:
        boards.extend([
            ("USAJobs", f"https://www.usajobs.gov/Search/Results?k={role_q}&l={loc_q}"),
            ("Dice Tech Jobs", f"https://www.dice.com/jobs?q={role_q}&location={loc_q}"),
            ("ZipRecruiter", f"https://www.ziprecruiter.com/jobs-search?search={role_q}&location={loc_q}"),
        ])
    elif country == "india":
        boards.extend([
            ("Naukri", f"https://www.naukri.com/{role_q}-jobs-in-{loc_q}"),
            ("Foundit", f"https://www.foundit.in/search/{role_q}-jobs-in-{loc_q}"),
            ("TimesJobs", f"https://www.timesjobs.com/candidate/job-search.html?searchType=personalizedSearch&txtKeywords={role_q}&txtLocation={loc_q}"),
        ])
    elif country == "canada":
        boards.extend([
            ("Job Bank Canada", f"https://www.jobbank.gc.ca/jobsearch/jobsearch?searchstring={role_q}&locationstring={loc_q}"),
            ("Workopolis", f"https://www.workopolis.com/jobsearch/{role_q}-jobs/{loc_q}"),
        ])
    elif country == "australia":
        boards.extend([
            ("Seek Australia", f"https://www.seek.com.au/{role_q}-jobs/in-{loc_q}"),
            ("Jora Australia", f"https://au.jora.com/{role_q}-jobs-in-{loc_q}"),
        ])
    else:
        boards.extend([
            ("Glassdoor", f"https://www.glassdoor.com/Job/jobs.htm?sc.keyword={role_q}&locT=N&locId=&locKeyword={loc_q}"),
            ("Remote OK", f"https://remoteok.com/remote-{role_q}-jobs"),
            ("Remotive", f"https://remotive.com/remote-jobs/search?search={role_q}"),
        ])

    # Always include remote/global sources because some users apply internationally.
    boards.extend([
        ("Remotive Remote Jobs", f"https://remotive.com/remote-jobs/search?search={role_q}"),
        ("Remote OK", f"https://remoteok.com/remote-{role_q}-jobs"),
    ])

    deduped = []
    seen = set()
    for label, url in boards:
        key = (label + url).lower()
        if key not in seen:
            seen.add(key)
            deduped.append((label, url))
    return deduped

def render_job_board_search_cards(country_name: str, location: str, roles: List[str], user_status: str = ""):
    st.markdown("### Live Job Board Searches")
    st.caption("Open current job results on major job boards when direct live results are limited.")
    for role in roles[:3]:
        with st.expander(role, expanded=True):
            links = get_job_board_links(country_name, location, role, user_status)[:5]
            compact_links = []
            for label, url in links:
                safe_label = html.escape(label)
                safe_url = html.escape(url, quote=True)
                compact_links.append(
                    f'<a class="compact-job-link" href="{safe_url}" target="_blank" rel="noopener noreferrer">↗ {safe_label}</a>'
                )
            st.markdown(
                "<div class='compact-job-link-grid'>" + "".join(compact_links) + "</div>",
                unsafe_allow_html=True
            )


def build_role_suggestions(user_titles: List[str], detected_roles_text: str, current_role: str) -> List[str]:
    roles: List[str] = []
    for role in user_titles:
        role = role.strip()
        if role and role not in roles:
            roles.append(role)

    if detected_roles_text:
        for line in detected_roles_text.splitlines():
            candidate = line.replace("-", "").strip()
            if candidate and candidate not in roles:
                roles.append(candidate)

    if current_role and current_role.strip() and current_role.strip() not in roles:
        roles.append(current_role.strip())

    fallback = [
        "Data Analyst",
        "Business Analyst",
        "Reporting Analyst",
        "Operations Analyst",
        "IT Support Specialist",
    ]
    for role in fallback:
        if len(roles) >= 5:
            break
        if role not in roles:
            roles.append(role)

    return roles[:5]

def get_country_market_hint(country_name: str) -> str:
    country = (country_name or "").strip().lower()
    if country == "germany":
        return "German language helps strongly for many local roles, but English-speaking jobs exist in tech, startups, analytics, product, and international companies."
    if country in {"netherlands", "the netherlands"}:
        return "English-speaking roles are more common than in many EU markets, especially in tech, operations, and international business functions."
    if country in {"austria", "switzerland"}:
        return "Local language is often important, especially for customer-facing and traditional companies."
    if country in {"canada", "united states", "united kingdom", "ireland", "australia", "new zealand"}:
        return "English-first roles are common, but local resume style and market positioning still matter."
    return "Target roles with transferable skills first, then adapt your CV and search keywords to local market expectations."


ADZUNA_COUNTRY_CODES = {
    "australia": "au",
    "austria": "at",
    "belgium": "be",
    "brazil": "br",
    "canada": "ca",
    "france": "fr",
    "germany": "de",
    "india": "in",
    "italy": "it",
    "mexico": "mx",
    "netherlands": "nl",
    "the netherlands": "nl",
    "new zealand": "nz",
    "poland": "pl",
    "singapore": "sg",
    "south africa": "za",
    "spain": "es",
    "switzerland": "ch",
    "united kingdom": "gb",
    "uk": "gb",
    "united states": "us",
    "usa": "us",
}

@st.cache_data(show_spinner=False, ttl=900)
def fetch_adzuna_jobs(query: str, country_name: str, location: str = "", limit: int = 24) -> List[Dict]:
    country_code = ADZUNA_COUNTRY_CODES.get((country_name or "").strip().lower())
    app_id = os.getenv("ADZUNA_APP_ID") or get_streamlit_secret("ADZUNA_APP_ID")
    app_key = os.getenv("ADZUNA_APP_KEY") or get_streamlit_secret("ADZUNA_APP_KEY")

    if not country_code or not app_id or not app_key:
        return []

    results: List[Dict] = []
    seen = set()

    for page_no in range(1, 6):
        params = {
            "app_id": app_id,
            "app_key": app_key,
            "results_per_page": "12",
            "what_phrase": query or "",
            "where": location or "",
            "sort_by": "date",
            "content-type": "application/json",
            "max_days_old": "30",
        }

        url = f"https://api.adzuna.com/v1/api/jobs/{country_code}/search/{page_no}?" + urllib.parse.urlencode(params)

        try:
            payload = http_get_json(url, timeout=12)
        except Exception:
            continue

        items = payload.get("results", []) if isinstance(payload, dict) else []
        if not items:
            continue

        for item in items:
            title = str(item.get("title", "")).strip()
            company_obj = item.get("company") or {}
            company = str(company_obj.get("display_name") if isinstance(company_obj, dict) else company_obj or "").strip()

            location_obj = item.get("location") or {}
            if isinstance(location_obj, dict):
                display_loc = str(location_obj.get("display_name") or "").strip()
            else:
                display_loc = str(location_obj or "").strip()

            redirect_url = str(item.get("redirect_url", "")).strip()
            description = str(item.get("description", "")).strip()
            contract = str(item.get("contract_type", "")).strip()
            salary_min = item.get("salary_min")
            salary_max = item.get("salary_max")
            salary_text = ""
            if salary_min or salary_max:
                salary_text = f"Salary: {salary_min or '—'} to {salary_max or '—'}"

            summary_parts = [x for x in [contract, salary_text, (description[:200] + "...") if description else ""] if x]
            summary = " • ".join(summary_parts)

            key = (title + "|" + company + "|" + display_loc).casefold()
            if title and key not in seen:
                seen.add(key)
                results.append({
                    "source": "Adzuna",
                    "title": title,
                    "company": company or "Employer not shown",
                    "location": display_loc or (location or country_name),
                    "remote": False,
                    "url": redirect_url,
                    "summary": summary
                })
            if len(results) >= limit:
                return results[:limit]

    return results[:limit]

def fetch_live_jobs_global(country_name: str, roles: List[str], location: str = "", user_status: str = "") -> List[Dict]:
    results: List[Dict] = []
    country_lower = (country_name or "").strip().lower()
    search_queries = build_live_job_queries(roles, user_status, max_queries=24, country_name=country_name)

    # Germany has two additional free sources.
    if country_lower == "germany":
        results.extend(fetch_live_jobs_for_germany(roles, location, user_status))

    # Adzuna supports many countries when API keys are configured.
    for query in search_queries:
        results.extend(fetch_adzuna_jobs(query, country_name, location, limit=30))

    # Free global remote source. This keeps worldwide job search useful even without Adzuna keys.
    for query in search_queries[:10]:
        results.extend(fetch_remotive_jobs(query, location, limit=20))

    # Broaden location if exact city/country query produces too few results.
    if len(results) < 15 and location and location.strip().lower() != (country_name or "").strip().lower():
        for query in search_queries[:12]:
            results.extend(fetch_adzuna_jobs(query, country_name, "", limit=30))
            results.extend(fetch_remotive_jobs(query, "", limit=20))

    deduped: List[Dict] = []
    seen = set()
    for item in results:
        key = (item.get("source", "") + "|" + item.get("title", "") + "|" + item.get("company", "") + "|" + item.get("location", "") + "|" + item.get("url", "")).casefold()
        if key not in seen:
            seen.add(key)
            deduped.append(item)

    return sort_jobs_for_user(deduped, user_status, roles, country_name)[:100]


def fallback_job_search_plan(country_name: str, location: str, roles: List[str], cv_text: str) -> Dict:
    best_titles = []
    for role in roles[:5]:
        best_titles.append({
            "title": role,
            "fit_reason": "This matches the transferable skills and positioning suggested by your resume.",
            "seniority": "Junior to Mid-level",
            "english_realistic": "Depends on country and company type"
        })

    search_terms = []
    for role in roles[:4]:
        search_terms.extend([
            f"{role} jobs in {location}",
            f"{role} {country_name} English speaking",
            f"{role} remote {country_name}",
        ])

    return {
        "best_match_titles": best_titles,
        "search_terms": search_terms[:12],
        "market_strategy": [
            get_country_market_hint(country_name),
            f"Start with roles closest to your current profile and apply broadly in {location or country_name}.",
            "Prioritize companies with international teams, clear English-language postings, and realistic entry requirements."
        ],
        "best_live_matches": [],
        "priority_plan": [
            "Apply first to the 10 closest-fit roles.",
            "Use 2 to 3 tailored CV versions for your top role clusters.",
            "Track applications and improve search keywords every 3 to 4 days."
        ],
        "resume_positioning": [
            "Keep your headline aligned to the target role.",
            "Highlight measurable achievements and relevant tools.",
            "Move the most market-relevant skills higher in the CV."
        ],
        "fastest_route": [
            "Apply to 10 to 15 realistic roles this week.",
            "Tailor your CV headline and summary for each role cluster.",
            "Use LinkedIn, Indeed, and the strongest local job board for your country.",
            "Message 3 recruiters or hiring managers where possible.",
            "Practice short interview answers through Work-O-Bot."
        ]
    }

def generate_job_search_plan(country_name: str, location: str, roles: List[str], cv_text: str, live_jobs: List[Dict]) -> Dict:
    prompt = f"""
Return ONLY valid JSON. Do not use markdown.

JSON format:
{{
  "best_match_titles": [
    {{
      "title": "string",
      "fit_reason": "string",
      "seniority": "string",
      "english_realistic": "string"
    }}
  ],
  "search_terms": ["string"],
  "market_strategy": ["string"],
  "best_live_matches": ["string"],
  "priority_plan": ["string"],
  "resume_positioning": ["string"],
  "fastest_route": ["string"]
}}

Country: {country_name}
Preferred location: {location}
Target roles: {roles}

Candidate CV:
{cv_text}

Live jobs:
{live_jobs[:8]}

Rules:
- Be realistic about seniority. Do not suggest mid-senior if the CV does not support it.
- Prefer entry-level or junior recommendations when direct experience is limited.
- Fill every array with useful content.
- best_match_titles must contain exactly 3 items.
- search_terms must contain 4 to 6 items and avoid repetitive variations.
- market_strategy must contain 3 short bullets.
- priority_plan must contain 3 short bullets.
- resume_positioning must contain 3 short bullets.
- fastest_route must contain 3 short bullets.
"""
    result = run_ai_prompt(prompt, force_language=st.session_state.get("preferred_language", "English"), json_mode=True)
    if result.startswith("ERROR:"):
        return fallback_job_search_plan(country_name, location, roles, cv_text)

    try:
        parsed = safe_json_loads(result)
        if not isinstance(parsed, dict) or not parsed.get("best_match_titles"):
            return fallback_job_search_plan(country_name, location, roles, cv_text)
        return parsed
    except Exception:
        return fallback_job_search_plan(country_name, location, roles, cv_text)

def render_job_plan(plan: Dict):
    st.markdown(f"### {txt('top_match_titles')}")
    titles = plan.get("best_match_titles", [])
    if isinstance(titles, list) and titles:
        for item in titles[:3]:
            if isinstance(item, dict):
                title = str(item.get("title", "Role"))
                seniority = str(item.get("seniority", "—"))
                english_realistic = str(item.get("english_realistic", "—"))
                fit_reason = str(item.get("fit_reason", "—"))
                if len(fit_reason) > 120:
                    fit_reason = fit_reason[:120].rsplit(" ", 1)[0] + "..."
                card_html = f"""
<div class="card">
  <div class="section-title">{html.escape(title)}</div>
  <div class="small-muted">{txt('likely_seniority')}: {html.escape(seniority)} • {txt('english_apply')}: {html.escape(english_realistic)}</div>
  <div style="margin-top:8px;">{html.escape(fit_reason)}</div>
</div>
"""
                st.markdown(card_html, unsafe_allow_html=True)
            else:
                st.markdown(f"- {item}")
    else:
        st.info(txt("no_roles_generated"))

    with st.expander(txt("more_job_details"), expanded=False):
        col1, col2 = st.columns(2)
        with col1:
            st.markdown(f"#### {txt('best_search_terms')}")
            for item in plan.get("search_terms", [])[:5]:
                st.markdown(f"- {item}")
            st.markdown(f"#### {txt('market_strategy')}")
            for item in plan.get("market_strategy", [])[:3]:
                st.markdown(f"- {item}")
            st.markdown(f"#### {txt('priority_plan')}")
            for item in plan.get("priority_plan", [])[:3]:
                st.markdown(f"- {item}")
        with col2:
            st.markdown(f"#### {txt('resume_positioning')}")
            for item in plan.get("resume_positioning", [])[:3]:
                st.markdown(f"- {item}")
            st.markdown(f"#### {txt('fastest_route')}")
            for item in plan.get("fastest_route", [])[:3]:
                st.markdown(f"- {item}")
            best_live = plan.get("best_live_matches", [])
            if best_live:
                st.markdown(f"#### {txt('best_live_matches')}")
                for item in best_live[:3]:
                    st.markdown(f"- {item}")

def analyze_cv_text_features(cv_text: str) -> Dict:
    text = cv_text or ""
    lower = text.lower()
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    words = re.findall(r"\S+", text)
    word_count = len(words)

    heading_patterns = {
        "professional_summary_present": r"(professional summary|summary|profile|about me|objective)",
        "work_experience_present": r"(work experience|experience|employment history|professional experience)",
        "skills_present": r"(^|\n)(skills|technical skills|core skills|competencies)\b",
        "education_present": r"(education|academic background|qualifications)",
    }

    email_present = bool(re.search(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", text, re.I))
    phone_present = bool(re.search(r"(\+?\d[\d\s\-()]{7,}\d)", text))
    linkedin_present = "linkedin.com" in lower
    contact_info_present = email_present or phone_present

    bullet_lines = [l for l in lines if l.startswith(("•", "-", "*"))]
    bullet_points_count = len(bullet_lines)
    bullet_points_present = bullet_points_count >= 2

    date_ranges_present = bool(re.search(r"(19\d{2}|20\d{2}).{0,10}(19\d{2}|20\d{2}|present|current|heute|till date)", lower, re.I))
    years_mentions = re.findall(r"(\d{1,2})\+?\s+(years|year|yrs)", lower)
    years_count = len(years_mentions)

    quantified_patterns = [
        r"\d+%",
        r"€\s?\d+",
        r"\$\s?\d+",
        r"\b\d+[+,]?\s+(users|clients|projects|tickets|years|months|team members|stakeholders)\b",
        r"\b(increased|reduced|improved|managed|supported|led)\b.{0,20}\d+",
    ]
    quantified_hits = 0
    for p in quantified_patterns:
        quantified_hits += len(re.findall(p, lower, re.I))
    quantified_achievements_present = quantified_hits > 0

    short_lines = sum(1 for l in lines if len(l) < 80)
    long_lines = sum(1 for l in lines if len(l) > 140)
    section_heading_hits = sum(1 for pattern in heading_patterns.values() if re.search(pattern, lower, re.I | re.M))
    clear_formatting = len(lines) >= 8 and short_lines >= 5 and long_lines <= max(2, len(lines) // 8)
    plain_text_readable = len(lines) >= 5 and word_count > 60 and long_lines <= max(3, len(lines) // 6)
    standard_headings_present = section_heading_hits >= 2

    action_verbs = [
        "developed", "analyzed", "managed", "led", "created", "built", "improved",
        "supported", "optimized", "designed", "implemented", "coordinated",
        "resolved", "delivered", "maintained", "automated"
    ]
    action_verb_hits = sum(1 for verb in action_verbs if verb in lower)

    skills_line_count = 0
    for l in lines:
        if "," in l and len(l.split(",")) >= 3:
            skills_line_count += 1

    contact_score = 1.0 if contact_info_present else 0.0
    summary_score = 1.0 if re.search(heading_patterns["professional_summary_present"], lower, re.I) else 0.0
    experience_score = 1.0 if re.search(heading_patterns["work_experience_present"], lower, re.I) else 0.0
    skills_score = min(1.0, 0.35 + 0.2 * skills_line_count) if re.search(heading_patterns["skills_present"], lower, re.I | re.M) else 0.0
    education_score = 1.0 if re.search(heading_patterns["education_present"], lower, re.I) else 0.0
    quantified_score = min(1.0, quantified_hits / 4) if quantified_hits else 0.0
    bullet_score = min(1.0, bullet_points_count / 8) if bullet_points_count else 0.0
    date_score = min(1.0, max(1, years_count) / 4) if date_ranges_present or years_count else 0.0
    formatting_score = 1.0 if clear_formatting else (0.55 if plain_text_readable else 0.2)

    if word_count < 120:
        grammar_quality = "weak"
    elif word_count < 220:
        grammar_quality = "average"
    else:
        grammar_quality = "good"

    features = {
        "contact_info_present": contact_info_present,
        "professional_summary_present": bool(re.search(heading_patterns["professional_summary_present"], lower, re.I)),
        "work_experience_present": bool(re.search(heading_patterns["work_experience_present"], lower, re.I)),
        "skills_present": bool(re.search(heading_patterns["skills_present"], lower, re.I | re.M)),
        "education_present": bool(re.search(heading_patterns["education_present"], lower, re.I)),
        "quantified_achievements_present": quantified_achievements_present,
        "clear_formatting": clear_formatting,
        "grammar_quality": grammar_quality,
        "standard_headings_present": standard_headings_present,
        "plain_text_readable": plain_text_readable,
        "bullet_points_present": bullet_points_present,
        "date_ranges_present": date_ranges_present,
        "word_count": word_count,
        "long_lines": long_lines,
        "bullet_points_count": bullet_points_count,
        "quantified_hits": quantified_hits,
        "action_verb_hits": action_verb_hits,
        "linkedin_present": linkedin_present,
        "section_heading_hits": section_heading_hits,
        "contact_score": contact_score,
        "summary_score": summary_score,
        "experience_score": experience_score,
        "skills_score": skills_score,
        "education_score": education_score,
        "quantified_score": quantified_score,
        "bullet_score": bullet_score,
        "date_score": date_score,
        "formatting_score": formatting_score,
    }
    return features


# =========================================================
# CV CLEANING + STRUCTURING
# =========================================================
def join_spaced_caps(match):
    return match.group(0).replace(" ", "")

def clean_cv_text(raw_text: str) -> str:
    """
    Fix common PDF extraction problems before WorkZo sends CV text to AI.
    """
    if not raw_text:
        return ""

    text = raw_text.replace("\x00", " ")
    text = re.sub(r"\r", "\n", text)
    text = re.sub(r"\b(?:[A-Z]\s){2,}[A-Z]\b", join_spaced_caps, text)

    replacements = {
        "VIZUALIZATION": "VISUALIZATION",
        "Scrapping": "Scraping",
        "Analisys": "Analysis",
        "suppoprt": "support",
        "Engince": "Engine",
        "knowlegde": "knowledge",
    }
    for wrong, right in replacements.items():
        text = re.sub(wrong, right, text, flags=re.I)

    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    return text.strip()

def build_created_cv_text(full_name, email, phone, location, summary, skills, experience, education, projects="", certifications="", languages="", extra_info="") -> str:
    return f"""
Name: {full_name}
Email: {email}
Phone: {phone}
Location: {location}

Professional Summary:
{summary}

Skills:
{skills}

Work Experience:
{experience}

Projects:
{projects}

Certifications:
{certifications}

Education:
{education}

Languages:
{languages}

Additional Information:
{extra_info}
""".strip()

def generate_cv_from_user_details(cv_details: str, target_country: str, user_status: str) -> str:
    prompt = build_quality_prompt(
        task="Create a complete professional CV from the user's raw details.",
        user_input=cv_details,
        expected_structure="""
Return:
1. Full CV
2. Missing Information to Improve It
3. Resume Score Estimate
4. ATS Improvement Suggestions
5. Questions to Ask the User Next
"""
    )
    prompt += f"""

Target country: {target_country}
User status: {user_status}

Rules:
- Build a clean, ATS-friendly CV.
- Use only information the user provided.
- Do not invent employers, dates, degrees, tools, or achievements.
- If details are missing, add a clear section called "Missing Information to Improve It".
- Make the CV suitable for the target country where possible.
"""
    return run_ai_prompt(prompt)

def organize_cv_for_display(cv_text: str) -> str:
    return clean_cv_text(cv_text)

# =========================================================
# SESSION STATE
# =========================================================
country_options = get_country_options()
language_options = get_language_options()
geo_country_default, geo_language_default = get_geo_defaults()

defaults = {
    "page": "onboarding",
    "onboarding_complete": False,
    "country": geo_country_default if geo_country_default in country_options else "Germany",
    "preferred_language": geo_language_default if geo_language_default in language_options else "English",
    "ui_language": geo_language_default if geo_language_default in language_options else "English",
    "response_language": geo_language_default if geo_language_default in language_options else "English",
    "user_status": "Career changer",
    "migration_country": geo_country_default if geo_country_default in country_options else "Germany",
    "career_goal": "",
    "cv_mode": "Upload CV",
    "cv_text": "",
    "dashboard_action": "Dashboard",
    "nav_page": "dashboard",

    # extracted dashboard profile
    "profile_summary": "",
    "current_role_detected": "",
    "key_skills_detected": "",
    "suggested_roles_detected": "",
    "best_fit_countries_detected": "",
    "next_actions_detected": "",
    "country_cv_readiness": "",
    "status_guidance": "",
    "cv_profile_raw": "",

    # dashboard analysis
    "cv_score_value": None,
    "ats_score_value": None,
    "country_readiness_score_value": None,
    "resume_strengths": "",
    "resume_improvements": "",
    "latest_resume_dashboard_analysis": "",
    "dashboard_cv_hash": "",

    # other scores
    "job_fit_score_value": None,
    "skill_gap_score_value": None,
    "interview_score_value": None,

    # outputs
    "latest_job_analysis": "",
    "latest_cv_analysis": "",
    "latest_interview": "",
    "latest_skill_gap": "",
    "latest_cover_letter": "",
    "latest_cv_translation": "",
    "latest_cover_letter_translation": "",

    # work-o-bot
    "workobot_mode": "General Help",
    "workobot_messages": [
        {
            "role": "assistant",
            "content": txt("workobot_intro") if "workobot_intro" in UI_TEXT.get(ui_lang(), {}) else "Hi, I’m Work-O-Bot. I can help with interview prep, career communication, and skill gap guidance."
        }
    ],

    # cache
    "dashboard_cache": {},
    "next_best_step_result": "",
    "ai_quality_style": "coach",
    "target_role_context": "",
}

for key, value in defaults.items():
    if key not in st.session_state:
        st.session_state[key] = value

init_beta_analytics()
track_event("app_open", "App")

# =========================================================
# HELPERS
# =========================================================
def google_search_url(query: str) -> str:
    return "https://www.google.com/search?q=" + urllib.parse.quote(query)

def extract_pdf_text(uploaded_file) -> str:
    text = ""
    with pdfplumber.open(uploaded_file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text.strip()

def safe_json_loads(raw: str) -> Dict:
    cleaned = raw.strip()

    # remove ```json fences if present
    cleaned = re.sub(r"^```json\s*", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"^```\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)

    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", cleaned, flags=re.DOTALL)
        if match:
            return json.loads(match.group(0))
        raise


# =========================================================
# AI QUALITY LAYER
# =========================================================
def get_user_status_context() -> str:
    return st.session_state.get("user_status", "Not specified")

def get_target_country_context() -> str:
    return st.session_state.get("country", "Not specified")

def get_target_role_context() -> str:
    return st.session_state.get("target_role_context", "") or st.session_state.get("current_role_detected", "") or "Not specified"

def compact_cv_context(max_chars: int = 3500) -> str:
    cv = st.session_state.get("cv_text", "") or ""
    cv = re.sub(r"\s+", " ", cv).strip()
    return cv[:max_chars] + ("..." if len(cv) > max_chars else "")

def clean_context_value(value, fallback: str = "Not available") -> str:
    value = str(value or "").strip()
    return value if value else fallback

def build_career_intelligence_layer(max_cv_chars: int = 4500) -> str:
    """
    Central context layer used before every AI call.
    This makes WorkZo's answers specific, country-aware, and less generic.
    """
    country = clean_context_value(st.session_state.get("country"), "Not selected")
    target_country = clean_context_value(st.session_state.get("migration_country") or st.session_state.get("country"), country)
    language = clean_context_value(st.session_state.get("preferred_language"), "English")
    user_status = clean_context_value(st.session_state.get("user_status"), "Not selected")
    target_role = clean_context_value(get_target_role_context(), "Not specified")

    return f"""
CAREER INTELLIGENCE LAYER
User selected country: {country}
Target / application country: {target_country}
Preferred output language: {language}
Current career situation: {user_status}
Target role or role direction: {target_role}
Career goal: {clean_context_value(st.session_state.get("career_goal"))}

AI-extracted profile summary:
{clean_context_value(st.session_state.get("profile_summary"))}

AI-extracted current role:
{clean_context_value(st.session_state.get("current_role_detected"))}

AI-extracted key skills:
{clean_context_value(st.session_state.get("key_skills_detected"))}

Detected resume strengths:
{clean_context_value(st.session_state.get("resume_strengths"))}

Detected resume improvements:
{clean_context_value(st.session_state.get("resume_improvements"))}

Current scores:
- Resume score: {clean_context_value(st.session_state.get("cv_score_value"), "Not calculated")}
- ATS score: {clean_context_value(st.session_state.get("ats_score_value"), "Not calculated")}
- Country readiness score: {clean_context_value(st.session_state.get("country_readiness_score_value"), "Not calculated")}

CV / profile text excerpt:
{compact_cv_context(max_cv_chars) or "No CV text available"}
""".strip()

def workzo_expert_context() -> str:
    return build_career_intelligence_layer()

SUPPORTED_UI_LANGUAGES = {"English", "German", "Dutch"}

def normalize_answer_language(language: str) -> str:
    language = (language or "English").strip()
    return language if language else "English"

def language_guard_rules(answer_lang: str) -> str:
    answer_lang = normalize_answer_language(answer_lang)
    if answer_lang == "German":
        return """
STRICT LANGUAGE MODE: German.
- Write the final answer completely in natural German.
- Do not mix English sentence starters, headings, explanations, or filler words.
- Allowed English only: company names, software/tool names, job titles when commonly used, URLs, email addresses, code, and exact keywords from a job ad/CV.
- Use German career terms where natural: Lebenslauf, Anschreiben, Berufserfahrung, Kenntnisse, Fähigkeiten, Ausbildung, Bewerbungsunterlagen.
- If you accidentally produce English headings or mixed English/German, rewrite the whole answer in German before returning it.
""".strip()
    if answer_lang == "Dutch":
        return """
STRICT LANGUAGE MODE: Dutch.
- Write the final answer completely in natural Dutch.
- Do not mix English sentence starters, headings, explanations, or filler words.
- Allowed English only: company names, software/tool names, job titles when commonly used, URLs, email addresses, code, and exact keywords from a job ad/CV.
- Use Dutch career terms where natural: cv, motivatiebrief, werkervaring, vaardigheden, opleiding, sollicitatie.
- If you accidentally produce English headings or mixed English/Dutch, rewrite the whole answer in Dutch before returning it.
""".strip()
    if answer_lang == "English":
        return """
STRICT LANGUAGE MODE: English.
- Write the final answer completely in English.
- Do not switch into German, Dutch, or another language unless the user specifically asks for a translation example.
""".strip()
    return f"""
STRICT LANGUAGE MODE: {answer_lang}.
- Write the final answer completely in natural {answer_lang}.
- Do not mix English sentence starters, headings, explanations, or filler words unless they are exact job titles, software/tool names, company names, URLs, email addresses, code, or keywords copied from the user's CV/job ad.
- Translate all guidance, headings, labels, and conclusions into {answer_lang}.
- If you accidentally produce mixed-language output, rewrite the whole answer in {answer_lang} before returning it.
""".strip()

def quality_system_prompt(answer_lang: str, system_addition: str = "") -> str:
    return f"""
You are WorkZo AI, a senior international career strategist, resume consultant, ATS specialist, interview coach, and job-search advisor.

Your job is NOT to give generic AI advice. Your job is to give precise, personalized, practical guidance based on:
- the user's CV/profile
- the user's status: fresh graduate, migrant, career changer, experienced professional, etc.
- the target country and local hiring expectations
- the user's target role and current experience level

{language_guard_rules(answer_lang)}

Quality rules:
1. Be specific. Avoid generic sentences like "improve your skills" unless you say exactly which skill, why it matters, and what to do next.
2. Use the user's actual CV/profile details whenever available. Reference concrete roles, tools, skills, industries, projects, and gaps from the provided context.
3. Be honest about weak fit, missing experience, language gaps, unrealistic seniority, or template mismatch.
4. Give prioritized actions: what to do first, second, third.
5. For migration/country-specific topics, mention local CV expectations, language expectations, and role-market fit.
6. For fresh graduates, focus on portfolio projects, internships, entry-level roles, keywords, and proof of skill.
7. For career changers, focus on transferable skills, bridge roles, portfolio proof, and realistic role titles.
8. For experienced users, focus on positioning, measurable achievements, leadership/impact, and market fit.
9. Keep output structured with useful headings and concise bullets. Prefer 3-5 strong points over long generic lists.
10. Never invent employers, degrees, certifications, dates, or achievements.
11. If information is missing, clearly say what is missing and give a best-effort recommendation.
12. End with a clear "Next best action" when appropriate.
13. Every answer must feel like it was written for this user, this country, and this career situation.
14. Do not give generic advice such as "network more", "improve your resume", or "learn more skills" without converting it into a concrete action.

Use this context where relevant:
{workzo_expert_context()}

{system_addition}
""".strip()

def build_quality_prompt(task: str, user_input: str, expected_structure: str = "") -> str:
    return f"""
Task: {task}

{build_career_intelligence_layer()}

User input:
{user_input}

Expected output structure:
{expected_structure}

Advanced output rules:
- Start with the most useful answer first, not a long introduction.
- Give practical, tailored, country-aware, and status-aware advice.
- Avoid generic advice. Mention concrete skills, role titles, keywords, CV sections, or actions wherever possible.
- Keep the answer easy to scan with short sections and bullets.
- When useful, include a short "Why this matters" and a "Next best action".
- Follow the selected preferred language strictly. Do not mix languages.
""".strip()


def run_ai_prompt(
    prompt: str,
    system_addition: str = "",
    force_language: str = None,
    force_english: bool = False,
    json_mode: bool = False
) -> str:
    if not can_make_request():
        return "ERROR: Usage limit reached. Please try again later."

    register_request()

    if force_english:
        answer_lang = "English"
    elif force_language:
        answer_lang = normalize_answer_language(force_language)
    else:
        answer_lang = normalize_answer_language(st.session_state.get("preferred_language", "English"))

    strict_system_addition = f"""
{system_addition}

FINAL OUTPUT LANGUAGE: {answer_lang}
You must obey STRICT LANGUAGE MODE. Before you answer, internally check that every heading, bullet, explanation, and closing sentence is in {answer_lang}.
Do not mention this language check to the user.
""".strip()

    system_prompt = quality_system_prompt(answer_lang, strict_system_addition)
    model_name = os.getenv("WORKZO_AI_MODEL") or get_streamlit_secret("WORKZO_AI_MODEL", "gpt-4o-mini")

    try:
        enhanced_user_prompt = f"""
{build_career_intelligence_layer()}

USER REQUEST / FEATURE TASK:
{prompt}

Response requirements:
- Do not answer generically. Use the career intelligence layer above.
- Give the strongest 3-5 insights only unless the task explicitly asks for more.
- Convert broad advice into exact actions the user can take.
- If a CV, job description, target role, or country detail is missing, say that clearly and provide the best next step.
""".strip()

        kwargs = {
            "model": model_name,
            "temperature": 0.2,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": enhanced_user_prompt}
            ]
        }
        if json_mode:
            kwargs["response_format"] = {"type": "json_object"}

        res = client.chat.completions.create(**kwargs)
        return (res.choices[0].message.content or "").strip()
    except Exception as e:
        return f"ERROR: {type(e).__name__}: {str(e)}"

def render_error_or_success(result: str):
    if result.startswith("ERROR:"):
        st.error(txt("ai_unavailable"))
        st.caption(result)
        return False
    return True

def parse_score(text: str) -> Optional[int]:
    if not text:
        return None

    match = re.search(r'(\d{1,3})\s*/\s*100', text)
    if match:
        value = int(match.group(1))
        if 0 <= value <= 100:
            return value

    match = re.search(r'(?<!\d)(\d{1,3})(?!\d)', text)
    if match:
        value = int(match.group(1))
        if 0 <= value <= 100:
            return value

    return None

def numbered_sections_to_markdown(text: str) -> Dict[str, str]:
    """Parse AI sections like `1. Title`, `### 1. Title`, or `**1. Title**`.
    This prevents the same full response appearing in every expander.
    """
    text = text or ""
    sections = {}
    pattern = r'(?m)^\s*(?:#{1,6}\s*)?(?:\*\*)?(\d+)\.\s*(.+?)(?:\*\*)?\s*$'
    matches = list(re.finditer(pattern, text))
    if not matches:
        return {"Result": text.strip()}

    for i, match in enumerate(matches):
        title = match.group(2).strip()
        title = title.replace("**", "").replace("__", "").strip().rstrip(":")
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[start:end].strip()
        body = re.sub(r"^[-–—\s]+", "", body).strip()
        sections[title] = body
    return sections

def render_section_cards(result: str, default_expand: bool = False):
    if not render_error_or_success(result):
        return
    sections = numbered_sections_to_markdown(result)
    for title, body in sections.items():
        with st.expander(title, expanded=default_expand):
            st.write(body if body else "—")


def strip_markdown_for_resume(text_value: str) -> str:
    """Remove AI markdown artifacts so the downloadable CV looks like a real resume, not a chat response."""
    text_value = text_value or ""
    cleaned_lines = []
    for raw in text_value.splitlines():
        line = raw.strip()
        if not line:
            cleaned_lines.append("")
            continue
        line = re.sub(r"^#{1,6}\s*", "", line)
        line = line.replace("**", "").replace("__", "")
        line = re.sub(r"^\s*[-*]\s+", "- ", line)
        line = re.sub(r"\[([^\]]+)\]\(([^\)]+)\)", r"\1: \2", line)
        cleaned_lines.append(line)
    cleaned = "\n".join(cleaned_lines)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
    return cleaned

def country_resume_rules_text(country: str) -> str:
    rules = get_country_cv_rules(country)
    return f"""
Country style rules for {country}:
- Recommended length: about {rules.get('pages', 2)} page(s), depending on experience.
- Recommended sections: {', '.join(rules.get('sections', []))}.
- Photo: {'allowed/optional' if rules.get('photo') else 'do not include'}.
- Date of birth / sensitive personal details: {'allowed only if user already provided and it is appropriate' if rules.get('dob') else 'do not include'}.
- Important exclusions: {rules.get('avoid', 'Avoid sensitive personal details and do not invent information.')}.
""".strip()

def make_docx_from_cv_text(title: str, cv_text: str, template_name: str = "ATS Resume", target_country: str = "") -> bytes:
    """Create a real editable DOCX resume. Falls back to TXT bytes if python-docx is unavailable."""
    cleaned = strip_markdown_for_resume(cv_text)
    if Document is None:
        return cleaned.encode("utf-8")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)

    lines = [line.rstrip() for line in cleaned.splitlines()]
    first_written = False
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        heading_candidate = line.upper() == line and len(line) <= 45 and not line.startswith("-")
        if not first_written:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(18)
            first_written = True
        elif line.startswith("- "):
            p = doc.add_paragraph(style=None)
            p.style = doc.styles["List Bullet"]
            p.add_run(line[2:].strip())
        elif heading_candidate or line.endswith(":"):
            p = doc.add_paragraph()
            run = p.add_run(line.rstrip(":"))
            run.bold = True
            run.font.size = Pt(12)
        else:
            doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def make_styled_pdf_from_cv_text(title: str, cv_text: str, template_name: str = "ATS Resume", target_country: str = "") -> bytes:
    """Create a clean downloadable PDF resume using ReportLab. """
    cleaned = strip_markdown_for_resume(cv_text)
    if SimpleDocTemplate is None:
        return cleaned.encode("utf-8")

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=36, leftMargin=36, topMargin=34, bottomMargin=34)
    styles = getSampleStyleSheet()
    normal = ParagraphStyle("WorkZoNormal", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.5, leading=12, spaceAfter=4)
    bullet = ParagraphStyle("WorkZoBullet", parent=normal, leftIndent=12, bulletIndent=4, spaceAfter=3)
    heading = ParagraphStyle("WorkZoHeading", parent=styles["Heading3"], fontName="Helvetica-Bold", fontSize=11, leading=14, spaceBefore=8, spaceAfter=5)
    name_style = ParagraphStyle("WorkZoName", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=18, leading=22, alignment=1, spaceAfter=6)

    story = []
    lines = [line.strip() for line in cleaned.splitlines()]
    first_written = False
    for line in lines:
        if not line:
            story.append(Spacer(1, 5))
            continue
        safe_line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        is_heading = (line.upper() == line and len(line) <= 45 and not line.startswith("-")) or line.endswith(":")
        if not first_written:
            story.append(Paragraph(safe_line, name_style))
            first_written = True
        elif line.startswith("- "):
            story.append(Paragraph(safe_line[2:].strip(), bullet, bulletText="•"))
        elif is_heading:
            story.append(Paragraph(safe_line.rstrip(":"), heading))
        else:
            story.append(Paragraph(safe_line, normal))
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def make_pdf_from_text(title: str, body: str) -> bytes:
    """
    Create a simple PDF from generated CV text.
    Falls back to encoded text if reportlab is unavailable, but Streamlit label will show dependency need.
    """
    if SimpleDocTemplate is None:
        return body.encode("utf-8")

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=42, leftMargin=42, topMargin=42, bottomMargin=42)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph(title, styles["Title"]))
    story.append(Spacer(1, 12))

    for raw_line in body.splitlines():
        line = raw_line.strip()
        if not line:
            story.append(Spacer(1, 8))
            continue
        safe_line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if line.endswith(":") or line.isupper():
            story.append(Paragraph(f"<b>{safe_line}</b>", styles["Heading3"]))
        else:
            story.append(Paragraph(safe_line, styles["BodyText"]))
            story.append(Spacer(1, 4))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def get_section_text(result: str, names: List[str], fallback_to_full: bool = False) -> str:
    sections = numbered_sections_to_markdown(result)
    lower_map = {k.lower().strip(): v for k, v in sections.items()}
    for name in names:
        value = sections.get(name, "").strip()
        if value:
            return value
        value = lower_map.get(name.lower().strip(), "").strip()
        if value:
            return value
    return result.strip() if fallback_to_full else ""


def show_gauge(score: Optional[int], title: str):
    if score is None:
        st.info(txt("not_analyzed"))
        return

    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=score,
        title={"text": title},
        gauge={
            "axis": {"range": [0, 100]},
            "bar": {"thickness": 0.35},
            "steps": [
                {"range": [0, 40], "color": "#f8d7da"},
                {"range": [40, 70], "color": "#fff3cd"},
                {"range": [70, 100], "color": "#d1e7dd"},
            ],
        }
    ))
    fig.update_layout(height=260, margin=dict(l=20, r=20, t=50, b=20))
    st.plotly_chart(fig, use_container_width=True)

def make_cv_hash(cv_text: str, country: str) -> str:
    raw = f"{SCORING_VERSION}||{country.strip()}||{st.session_state.get('migration_country', '').strip()}||{st.session_state.get('user_status', '').strip()}||{cv_text.strip()}"
    return hashlib.md5(raw.encode("utf-8")).hexdigest()

def bool_score(flag: bool, points: int) -> int:
    return points if flag else 0

def calculate_resume_score(features: Dict) -> int:
    """
    Practical resume quality score. This is not random AI scoring.
    It rewards complete, readable, achievement-based CVs and penalizes missing basics.
    """
    score = 10
    score += round(features.get("contact_score", 0) * 10)
    score += round(features.get("summary_score", 0) * 10)
    score += round(features.get("experience_score", 0) * 18)
    score += round(features.get("skills_score", 0) * 12)
    score += round(features.get("education_score", 0) * 8)
    score += round(features.get("quantified_score", 0) * 16)
    score += round(features.get("bullet_score", 0) * 10)
    score += round(features.get("date_score", 0) * 8)
    score += round(features.get("formatting_score", 0) * 8)

    action_verb_hits = int(features.get("action_verb_hits", 0))
    if action_verb_hits >= 10:
        score += 6
    elif action_verb_hits >= 6:
        score += 4
    elif action_verb_hits >= 3:
        score += 2

    if features.get("linkedin_present", False):
        score += 2

    word_count = int(features.get("word_count", 0))
    if word_count < 120:
        score -= 18
    elif word_count < 180:
        score -= 10
    elif word_count > 1000:
        score -= 10
    elif word_count > 800:
        score -= 5

    long_lines = int(features.get("long_lines", 0))
    if long_lines > 6:
        score -= 8
    elif long_lines > 3:
        score -= 4

    if not features.get("contact_info_present", False):
        score -= 12
    if not features.get("professional_summary_present", False):
        score -= 8
    if not features.get("work_experience_present", False):
        score -= 15
    if not features.get("skills_present", False):
        score -= 10
    if not features.get("bullet_points_present", False):
        score -= 7
    if not features.get("quantified_achievements_present", False):
        score -= 8

    return max(20, min(score, 96))


def calculate_ats_score(features: Dict) -> int:
    """
    ATS score focuses on machine readability, standard sections, keywords, dates,
    and simple formatting. It is stricter than the general resume score.
    """
    score = 8
    score += 18 if features.get("standard_headings_present", False) else 4
    score += round(features.get("contact_score", 0) * 10)
    score += 14 if features.get("plain_text_readable", False) else 4
    score += round(features.get("skills_score", 0) * 12)
    score += round(features.get("experience_score", 0) * 14)
    score += round(features.get("education_score", 0) * 8)
    score += round(features.get("bullet_score", 0) * 10)
    score += round(features.get("date_score", 0) * 10)
    score += min(8, int(features.get("action_verb_hits", 0)))

    word_count = int(features.get("word_count", 0))
    if word_count < 120:
        score -= 15
    elif word_count < 180:
        score -= 8
    elif word_count > 1000:
        score -= 10

    long_lines = int(features.get("long_lines", 0))
    if long_lines > 6:
        score -= 10
    elif long_lines > 3:
        score -= 5

    if not features.get("contact_info_present", False):
        score -= 12
    if not features.get("standard_headings_present", False):
        score -= 12
    if not features.get("date_ranges_present", False):
        score -= 8
    if not features.get("skills_present", False):
        score -= 10
    if not features.get("plain_text_readable", False):
        score -= 8

    return max(20, min(score, 95))


def calculate_country_readiness_score(features: Dict, target_country: str, user_status: str) -> int:
    """
    Resume readiness for selected migration/target country.
    This is intentionally stricter than the general resume score.
    """
    base = calculate_resume_score(features)

    score = int(base * 0.72)

    # Country-template/readability signals
    if features.get("contact_info_present", False):
        score += 6
    if features.get("standard_headings_present", False):
        score += 8
    if features.get("plain_text_readable", False):
        score += 8
    if features.get("date_ranges_present", False):
        score += 5
    if features.get("skills_present", False):
        score += 5
    if features.get("quantified_achievements_present", False):
        score += 6

    country = (target_country or "").lower()
    status = (user_status or "").lower()

    # Migrating users need stronger country adaptation.
    if "migrate" in status or "abroad" in status:
        score -= 5

    # Local language/country expectations: not exact, but useful early signal.
    cv_raw = (st.session_state.get("cv_text", "") or "").lower()
    if country in {"germany", "austria", "switzerland"}:
        if "german" in cv_raw or "deutsch" in cv_raw:
            score += 5
        else:
            score -= 6
    if country in {"united kingdom", "united states", "canada", "australia", "ireland"}:
        if "english" in cv_raw:
            score += 4

    if not features.get("work_experience_present", False):
        score -= 8
    if int(features.get("word_count", 0)) < 180:
        score -= 6

    return max(25, min(score, 92))

def analyze_resume_dashboard_stable(cv_text: str, force_refresh: bool = False):
    cv_key = make_cv_hash(cv_text, st.session_state.country)

    if not force_refresh and cv_key in st.session_state.dashboard_cache:
        cached = st.session_state.dashboard_cache[cv_key]
        apply_dashboard_cache(cached)
        return

    prompt = f"""
Analyze this CV like a senior international career advisor and return ONLY valid JSON.
Do not add markdown, explanations, or code fences.

The analysis must be personalized to the user status and target country.
Do not give generic strengths like "good communication" unless the CV actually shows it.

JSON format:
{{
  "professional_summary": "short paragraph",
  "current_or_likely_role": "role name",
  "key_skills": ["skill1", "skill2", "skill3", "skill4", "skill5"],
  "suggested_roles": ["role1", "role2", "role3"],
  "best_fit_countries": ["Country - specific reason based on CV evidence", "Country - specific reason based on CV evidence", "Country - specific reason based on CV evidence"],
  "country_cv_readiness": "short country-specific verdict on whether this CV fits the target country",
  "status_guidance": "short advice based on the user's career status",
  "next_actions": ["action1", "action2", "action3", "action4"],
  "strengths": ["specific strength from CV", "specific strength from CV", "specific strength from CV", "specific strength from CV"],
  "improvements": ["specific improvement", "specific improvement", "specific improvement", "specific improvement"],
  "next_best_actions": ["action1", "action2", "action3"],
  "country_readiness_notes": ["note1", "note2", "note3"]
}}

Rules:
- Be concise, practical, and specific.
- Consider the user's status from onboarding.
- Include next_best_actions that tell the user exactly what to do next.
- Include country_readiness_notes for the chosen country.
- For best_fit_countries, suggest only the top 3 countries where this CV/profile is likely to be competitive.
- Do not give generic country names. Each reason must mention a specific CV signal such as role background, language, tools, domain, education, or transferable experience.
- Include the selected target country if it is realistic; otherwise explain why another market fits better.
- Keep each country reason under 18 words.
- Do not generate any numeric score.
- Do not add keys outside this JSON structure.
- Return valid JSON only.
- Keep the JSON keys exactly in English, but write all JSON values in the selected AI response language.
- Do not mix languages inside the JSON values.

Country context: {st.session_state.country}
User status: {st.session_state.get('user_status', 'Not specified')}
Career goal: {st.session_state.get('career_goal', 'Not specified')}

CV:
{cv_text}
"""
    result = run_ai_prompt(prompt, force_language=st.session_state.get("preferred_language", "English"), json_mode=True)

    if result.startswith("ERROR:"):
        return

    try:
        parsed = safe_json_loads(result)
    except Exception:
        st.warning("The resume analysis response could not be parsed. Please try refresh once.")
        return

    features = analyze_cv_text_features(cv_text)

    cached = {
        "profile_summary": parsed.get("professional_summary", "").strip(),
        "current_role_detected": parsed.get("current_or_likely_role", "").strip(),
        "key_skills_detected": parsed.get("key_skills", []),
        "suggested_roles_detected": parsed.get("suggested_roles", []),
        "best_fit_countries_detected": parsed.get("best_fit_countries", []),
        "country_cv_readiness": parsed.get("country_cv_readiness", ""),
        "status_guidance": parsed.get("status_guidance", ""),
        "next_actions_detected": parsed.get("next_actions", []),
        "resume_strengths": parsed.get("strengths", []),
        "resume_improvements": parsed.get("improvements", []),
        "next_best_actions_detected": parsed.get("next_best_actions", []),
        "country_readiness_notes_detected": parsed.get("country_readiness_notes", []),
        "country_readiness_score_value": calculate_country_readiness_score(features, st.session_state.get("migration_country") or st.session_state.country, st.session_state.get("user_status", "")),
        "cv_profile_raw": json.dumps({"ai_extract": parsed, "text_features": features}, indent=2),
        "cv_score_value": calculate_resume_score(features),
        "ats_score_value": calculate_ats_score(features),
    }

    st.session_state.dashboard_cache[cv_key] = cached
    apply_dashboard_cache(cached)
    st.session_state.dashboard_cv_hash = cv_key

def apply_dashboard_cache(cached: Dict):
    st.session_state.profile_summary = cached.get("profile_summary", "")
    st.session_state.current_role_detected = cached.get("current_role_detected", "")

    skills = cached.get("key_skills_detected", [])
    roles = cached.get("suggested_roles_detected", [])
    strengths = cached.get("resume_strengths", [])
    improvements = cached.get("resume_improvements", [])
    best_fit_countries = cached.get("best_fit_countries_detected", [])
    next_best_actions = cached.get("next_best_actions_detected", [])
    country_readiness_notes = cached.get("country_readiness_notes_detected", [])
    next_actions = cached.get("next_actions_detected", [])

    st.session_state.key_skills_detected = "\n".join([f"- {x}" for x in skills]) if isinstance(skills, list) else str(skills)
    st.session_state.suggested_roles_detected = "\n".join([f"- {x}" for x in roles]) if isinstance(roles, list) else str(roles)
    st.session_state.resume_strengths = "\n".join([f"- {x}" for x in strengths]) if isinstance(strengths, list) else str(strengths)
    st.session_state.resume_improvements = "\n".join([f"- {x}" for x in improvements]) if isinstance(improvements, list) else str(improvements)
    st.session_state.best_fit_countries_detected = "\n".join([f"- {x}" for x in best_fit_countries]) if isinstance(best_fit_countries, list) else str(best_fit_countries)
    st.session_state.next_best_step_result = "\n".join([f"- {x}" for x in next_best_actions]) if isinstance(next_best_actions, list) else str(next_best_actions)
    st.session_state.country_readiness_notes = "\n".join([f"- {x}" for x in country_readiness_notes]) if isinstance(country_readiness_notes, list) else str(country_readiness_notes)

    st.session_state.cv_profile_raw = cached.get("cv_profile_raw", "")
    st.session_state.cv_score_value = cached.get("cv_score_value")
    st.session_state.ats_score_value = cached.get("ats_score_value")
    st.session_state.country_readiness_score_value = cached.get("country_readiness_score_value")

def reset_onboarding():
    preserve = {
        "request_count": st.session_state.request_count,
        "first_request_time": st.session_state.first_request_time,
        "ui_language": st.session_state.ui_language,
        "response_language": st.session_state.response_language,
        "dashboard_cache": st.session_state.dashboard_cache,
        "geo_country_suggestion": st.session_state.get("geo_country_suggestion", "Germany"),
        "geo_language_suggestion": st.session_state.get("geo_language_suggestion", "English"),
    }
    for key in list(st.session_state.keys()):
        del st.session_state[key]

    for key, value in defaults.items():
        st.session_state[key] = value

    for key, value in preserve.items():
        st.session_state[key] = value

    st.session_state.page = "onboarding"
    st.session_state.onboarding_complete = False

def set_new_resume_and_refresh(new_resume_text: str):
    st.session_state.cv_text = new_resume_text.strip()

    st.session_state.profile_summary = ""
    st.session_state.current_role_detected = ""
    st.session_state.key_skills_detected = ""
    st.session_state.suggested_roles_detected = ""
    st.session_state.best_fit_countries_detected = ""
    st.session_state.next_actions_detected = ""
    st.session_state.country_cv_readiness = ""
    st.session_state.status_guidance = ""
    st.session_state.cv_profile_raw = ""
    st.session_state.cv_score_value = None
    st.session_state.ats_score_value = None
    st.session_state.resume_strengths = ""
    st.session_state.resume_improvements = ""

    analyze_resume_dashboard_stable(st.session_state.cv_text, force_refresh=True)



def render_metric_card(label: str, value: str, foot: str = ""):
    st.markdown(f"""
    <div class="metric-card">
        <div class="metric-label">{label}</div>
        <div class="metric-value">{value}</div>
        <div class="metric-foot">{foot}</div>
    </div>
    """, unsafe_allow_html=True)


def render_feature_tile(icon: str, title: str, copy: str):
    st.markdown(f"""
    <div class="feature-tile">
        <div style="font-size:1.45rem; margin-bottom:8px;">{icon}</div>
        <div class="feature-title">{title}</div>
        <div class="feature-copy">{copy}</div>
    </div>
    """, unsafe_allow_html=True)


def generate_next_best_steps(cv_text: str, country_name: str, current_role: str, suggested_roles: str, user_status: str = "", career_goal: str = "") -> str:
    prompt = f"""
Country: {country_name}
User status: {user_status or 'Not specified'}
Career goal: {career_goal or 'Not specified'}
Current role: {current_role or 'Not clear'}
Suggested roles:
{suggested_roles}
Candidate CV:
{cv_text}

Create a personalized action center. Adapt the advice based on the user status:
st.info("⚠️ WorkZo AI is currently in Beta. Some results may not be perfect yet. Your feedback helps improve the tool.")

- Fresh graduate: portfolio projects, internships, entry-level search, skills proof.
- Career changer: transferable skills, bridge role, portfolio, realistic job titles.
- Migrating to another country: country CV format, language expectations, local job titles, application documents.
- Returning after career break: confidence positioning, gap explanation, restart strategy.
- Experienced professional: positioning, seniority, specialization, leadership proof.

Return in this exact structure:
1. Resume Readiness Verdict
2. Best Immediate Goal
3. This Week's 3 Priority Actions
4. Best Role Cluster to Target
5. One Resume Fix That Will Help Most
6. One Skill or Proof to Build Next
7. Country-Specific Advice
8. One Message to Send Today

Keep it practical, specific, and motivating.
"""
    result = run_ai_prompt(prompt)
    return result

# =========================================================
# HEADER
# =========================================================
# HEADER
# =========================================================
maybe_scroll_to_top()
header_col1, header_col2 = st.columns([1, 7])

with header_col1:
    if os.path.exists(LOGO_PATH):
        st.image(LOGO_PATH, width=72)
    else:
        st.markdown("<div style='font-size:46px; line-height:1.2;'>🚀</div>", unsafe_allow_html=True)

with header_col2:
    title_col, badge_col = st.columns([3, 1])
    with title_col:
        st.markdown("<h1 style='margin:0; letter-spacing:0.02em;'>WORKZO AI</h1>", unsafe_allow_html=True)
    with badge_col:
        st.markdown('<span class="beta-badge">BETA V8.7</span>', unsafe_allow_html=True)
    st.markdown(f"""
    <div style="margin-top:6px;">
        <h4 style="margin:0;">{txt("title")}</h4>
        <p style="color: gray; margin:6px 0 0 0;">{txt("subtitle")}</p>
    </div>
    """, unsafe_allow_html=True)


# =========================================================
# ONBOARDING
# =========================================================
def show_onboarding():
    maybe_scroll_to_top()
    st.subheader(txt("onboarding_title"))
    st.caption(txt("onboarding_subtitle"))

    st.info(txt("beta_privacy_note"))

    language_list = language_options if language_options else ["English", "German", "Dutch"]
    lang_default = st.session_state.get("preferred_language", "English")
    if lang_default not in language_list:
        lang_default = "English" if "English" in language_list else language_list[0]

    preferred_language = st.selectbox(
        txt("preferred_language"),
        language_list,
        index=language_list.index(lang_default),
        key="onboarding_preferred_language",
        help=txt("language_help")
    )
    set_single_preferred_language(preferred_language)

    country_index = country_options.index(st.session_state.country) if st.session_state.country in country_options else 0
    country = st.selectbox(txt("country"), country_options, index=country_index, key="onboarding_country")

    status_options = [
        txt("select_optional"),
        txt("student_thesis_internship"),
        txt("local_jobseeker"),
        txt("fresh_graduate"),
        txt("career_changer"),
        txt("migrant"),
        txt("experienced"),
        txt("returning"),
    ]
    status_display = st.selectbox(
        txt("user_status"),
        status_options,
        index=0,
        key="onboarding_user_status",
        format_func=lambda value: value,
        help=txt("status_help")
    )
    status_map = {
        txt("student_thesis_internship"): STUDENT_STATUS_INTERNAL,
        txt("local_jobseeker"): "Looking for jobs locally",
        txt("fresh_graduate"): "Fresh graduate / entry level",
        txt("career_changer"): "Career changer",
        txt("migrant"): "Planning to migrate / applying abroad",
        txt("experienced"): "Experienced professional",
        txt("returning"): "Returning after a career break",
    }
    user_status = "" if status_display == txt("select_optional") else status_map.get(status_display, status_display).strip()

    migration_country = ""
    if "migrate" in user_status.lower() or "abroad" in user_status.lower():
        migration_default = st.session_state.get("migration_country") or st.session_state.get("country") or "Germany"
        migration_index = country_options.index(migration_default) if migration_default in country_options else 0
        migration_country = st.selectbox(
            txt("migration_country"),
            country_options,
            index=migration_index
        )

    st.markdown(f"### {txt('resume_input')}")
    st.caption(txt("resume_choice_caption"))

    if "cv_mode" not in st.session_state:
        st.session_state.cv_mode = ""

    c_upload, c_create = st.columns(2)
    with c_upload:
        st.markdown(f"""
        <div class="choice-card">
            <div class="choice-card-title">📄 {html.escape(txt('upload_cv_title'))}</div>
            <div class="choice-card-copy">{html.escape(txt('upload_cv_desc'))}</div>
        </div>
        """, unsafe_allow_html=True)

        uploaded_direct = st.file_uploader(
            "Upload",
            type=["pdf", "txt"],
            key="upload_cv_direct",
            label_visibility="collapsed"
        )

        if uploaded_direct is not None:
            st.session_state.cv_mode = "Upload CV"
            st.session_state.uploaded_file_direct = uploaded_direct

    with c_create:
        st.markdown(f"""
        <div class="choice-card">
            <div class="choice-card-title">✍️ {html.escape(txt('create_cv_title'))}</div>
            <div class="choice-card-copy">{html.escape(txt('create_cv_desc'))}</div>
        </div>
        """, unsafe_allow_html=True)
        if st.button(txt("create_cv"), key="choose_create_cv", use_container_width=True):
            st.session_state.cv_mode = "Create CV"

    cv_mode = st.session_state.get("cv_mode", "")

    if not cv_mode:
        st.info(txt("choose_resume_continue"))
        return

    with st.form("onboarding_form_v48"):
        cv_text_input = ""
        full_name = ""
        email = ""
        phone = ""
        location = ""
        summary = ""
        skills = ""
        experience = ""
        projects = ""
        certifications = ""
        education = ""
        languages = ""
        extra_info = ""

        if cv_mode == "Upload CV":
            uploaded_file = st.session_state.get("uploaded_file_direct")

            if uploaded_file is not None:
                if uploaded_file.size > 5 * 1024 * 1024:
                    st.error(txt("file_too_large"))
                elif uploaded_file.type == "text/plain":
                    cv_text_input = organize_cv_for_display(uploaded_file.read().decode("utf-8", errors="ignore"))
                elif uploaded_file.type == "application/pdf":
                    try:
                        cv_text_input = organize_cv_for_display(extract_pdf_text(uploaded_file))
                    except Exception as e:
                        st.error(f"{txt('pdf_read_error')}: {e}")
                else:
                    st.error(txt("unsupported_file"))
            else:
                st.info(txt("upload_cv_instruction"))

        else:
            st.markdown(f"#### {txt('guided_cv_builder')}")
            st.caption(txt("guided_cv_caption"))

            c1, c2 = st.columns(2)
            with c1:
                full_name = st.text_input(txt("full_name"), placeholder="Example: Alex Morgan")
                email = st.text_input(txt("email"), placeholder="Example: alex.morgan@email.com")
                phone = st.text_input(txt("phone"), placeholder="Example: +49 123 456789")
                location = st.text_input(txt("location"), placeholder="Example: Berlin, Germany")
            with c2:
                languages = st.text_area(txt("languages_label"), placeholder="Example: English B2, German A2, Spanish native")
                certifications = st.text_area(txt("cert_courses"), placeholder="Example: Google Data Analytics Certificate, Excel Advanced, AWS Cloud Practitioner basics")

            summary = st.text_area(txt("summary"), placeholder="Example: Customer support professional with experience in SaaS tools, now transitioning into data analysis.")
            skills = st.text_area(txt("skills"), placeholder="Example: SQL, Excel, Python, Tableau, communication, problem solving, CRM tools")
            experience = st.text_area(txt("experience"), placeholder="Example: Technical Support Associate, ABC Software, 2020–2024 — handled customer tickets, troubleshooting, and product training.")
            projects = st.text_area(txt("projects_label"), placeholder="Example: Built a sales dashboard in Tableau using sample sales data to track monthly revenue and customer trends.")
            education = st.text_area(txt("education"), placeholder="Example: Data Science Bootcamp, Online Academy, 2024")
            extra_info = st.text_area(txt("extra_cv_info"), placeholder="Example: Target role: Junior Data Analyst. Open to remote/hybrid roles in Germany.")

            if any([full_name.strip(), email.strip(), phone.strip(), location.strip(), summary.strip(), skills.strip(), experience.strip(), projects.strip(), certifications.strip(), education.strip(), languages.strip(), extra_info.strip()]):
                cv_text_input = build_created_cv_text(
                    full_name, email, phone, location, summary, skills, experience, education, projects, certifications, languages, extra_info
                )

        submitted = st.form_submit_button(txt("continue"))

        if submitted:
            track_button_click("Continue to Dashboard", "Onboarding", {"cv_mode": cv_mode, "selected_status": user_status})
            if not country:
                st.warning(txt("warn_country"))
                return

            if cv_mode == "Upload CV":
                if not cv_text_input.strip():
                    st.warning(txt("warn_upload_cv"))
                    return
            else:
                if not full_name.strip():
                    st.warning(txt("warn_full_name"))
                    return
                if not summary.strip() and not experience.strip() and not skills.strip():
                    st.warning("Please enter at least a summary, experience, or skills so WorkZo can build your CV.")
                    return

                with st.spinner("Creating your CV draft with AI..."):
                    ai_cv = generate_cv_from_user_details(cv_text_input, migration_country if migration_country else country, user_status)
                    if ai_cv and not ai_cv.startswith("ERROR:"):
                        st.session_state.created_cv_ai_output = ai_cv
                        cv_text_input = ai_cv

            st.session_state.country = country
            st.session_state.user_status = user_status
            st.session_state.migration_country = migration_country if migration_country else country
            st.session_state.cv_mode = cv_mode
            st.session_state.cv_text = clean_cv_text(cv_text_input)
            track_event("cv_ready", "Onboarding", {"cv_mode": cv_mode, "user_status": user_status, "target_country": st.session_state.get("migration_country", country)})
            st.session_state.onboarding_complete = True
            st.session_state.page = "dashboard"

            with st.spinner("Reading your resume and preparing dashboard..."):
                analyze_resume_dashboard_stable(st.session_state.cv_text, force_refresh=True)

            st.rerun()

# =========================================================
# WORK-O-BOT
# =========================================================
# WORK-O-BOT
# =========================================================
def run_workobot(user_message: str, mode: str):
    task_map = {
        "General Help": "Act as a career strategist and answer the user's question with personalized, practical guidance.",
        "German Practice": "Act as a workplace German coach. Correct the user's German, explain briefly, and give practice sentences for career situations.",
        "Mock Test Prep": "Act as a language/mock-test coach. Create exam-style tasks, model answers, corrections, and scoring guidance.",
        "Interview Prep": "Act as a realistic interviewer and interview coach. Give questions, evaluate answers, and improve responses using STAR/CAR structure.",
        "Career Communication": "Act as a professional communication coach. Improve emails, LinkedIn messages, HR replies, and interview answers.",
        "Skill Gap Help": "Act as a skill-gap strategist. Identify realistic gaps, priority learning path, portfolio proof, and role-targeting advice.",
    }

    expected = """
1. Direct Answer
2. Personalized Analysis
3. Suggested Response / Example
4. Mistakes or Risks to Avoid
5. Practice / Next Best Action
"""

    prompt = build_quality_prompt(
        task=task_map.get(mode, task_map["General Help"]),
        user_input=user_message,
        expected_structure=expected
    )

    result = run_ai_prompt(
        prompt,
        system_addition="""
For Work-O-Bot:
- Sound like a real coach, not a generic chatbot.
- Ask at most one follow-up question only if it is truly needed.
- If the user is practicing language/interview, include a corrected version, a natural version, and one practice task.
- If the user asks about jobs/career, include realistic role levels and target-country fit.
"""
    )
    if result.startswith("ERROR:"):
        return "I couldn’t answer right now because the AI service is unavailable."
    return result

def show_workobot():
    st.subheader(txt("workobot"))
    st.caption(txt("workobot_sub"))

    mode = st.selectbox(
        txt("chat_mode"),
        [
            "General Help",
            "German Practice",
            "Mock Test Prep",
            "Interview Prep",
            "Career Communication",
            "Skill Gap Help",
        ],
        index=[
            "General Help",
            "German Practice",
            "Mock Test Prep",
            "Interview Prep",
            "Career Communication",
            "Skill Gap Help",
        ].index(st.session_state.get("workobot_mode", "General Help"))
    )
    st.session_state.workobot_mode = mode

    st.markdown(f"### {txt('quick_prompts')}")
    qp1, qp2, qp3, qp4 = st.columns(4)
    quick_prompt = None
    country_quick_prompts = get_country_specific_quick_prompts(st.session_state.country)

    with qp1:
        if st.button(country_quick_prompts[0][0], use_container_width=True, key=f'workobot_quick_prompt_0_{st.session_state.get("country", "global")}_{mode}'):
            quick_prompt = country_quick_prompts[0][1]
    with qp2:
        if st.button(country_quick_prompts[1][0], use_container_width=True, key=f'workobot_quick_prompt_1_{st.session_state.get("country", "global")}_{mode}'):
            quick_prompt = country_quick_prompts[1][1]
    with qp3:
        if st.button(country_quick_prompts[2][0], use_container_width=True, key=f'workobot_quick_prompt_2_{st.session_state.get("country", "global")}_{mode}'):
            quick_prompt = country_quick_prompts[2][1]
    with qp4:
        if st.button(country_quick_prompts[3][0], use_container_width=True, key=f'workobot_quick_prompt_3_{st.session_state.get("country", "global")}_{mode}'):
            quick_prompt = country_quick_prompts[3][1]

    right1, right2 = st.columns([1, 5])
    with right1:
        if st.button(txt("clear_chat"), key="workobot_clear_chat"):
            st.session_state.workobot_messages = [
                {
                    "role": "assistant",
                    "content": txt("workobot_intro") if "workobot_intro" in UI_TEXT.get(ui_lang(), {}) else "Hi, I’m Work-O-Bot. I can help with interview prep, career communication, and skill gap guidance."
                }
            ]
            st.rerun()

    for msg in st.session_state.workobot_messages:
        with st.chat_message(msg["role"]):
            st.write(msg["content"])

    user_input = st.chat_input(txt("chat_placeholder"))
    final_input = quick_prompt or user_input

    if final_input:
        st.session_state.workobot_messages.append({"role": "user", "content": final_input})
        with st.chat_message("user"):
            st.write(final_input)

        with st.chat_message("assistant"):
            with st.spinner(txt("thinking") if "thinking" in UI_TEXT.get(ui_lang(), {}) else "Work-O-Bot is thinking..."):
                reply = run_workobot(final_input, mode)
                st.write(reply)

        st.session_state.workobot_messages.append({"role": "assistant", "content": reply})

def get_country_specific_quick_prompts(country_name: str):
    country = (country_name or "").strip().lower()

    if country == "germany":
        return [
            ("B1/B2 mock questions", "Give me a German B1/B2 mock test practice set with example questions and answers."),
            ("Career communication", "Help me improve my career communication with 3 examples relevant to Germany."),
            ("Skill gap help", "Based on my CV and the German market, what skill gap should I fix first?"),
            ("German for interviews", "Give me German interview questions and sample answers for my profile."),
        ]

    if country == "india":
        return [
            ("Interview questions", "Give me interview questions and strong sample answers for roles relevant to my profile in India."),
            ("HR communication", "Help me improve HR and recruiter communication with 3 examples relevant to India."),
            ("Skill gap help", "Based on my CV and the Indian market, what skill gap should I fix first?"),
            ("Job search plan", "Give me a practical interview and job search plan for my profile in India."),
        ]

    return [
        ("Interview questions", f"Give me interview questions and strong sample answers for roles relevant to my profile in {country_name}."),
        ("Career communication", f"Help me improve my career communication with 3 examples relevant to {country_name}."),
        ("Skill gap help", f"Based on my CV and the {country_name} market, what skill gap should I fix first?"),
        ("Job search plan", f"Give me a practical interview and job search plan for my profile in {country_name}."),
    ]

def extract_best_effort_cover_letter_sections(result: str):
    sections = numbered_sections_to_markdown(result)
    normalized = {k.strip().lower(): v.strip() for k, v in sections.items()}

    def first_match(candidates):
        for key, value in normalized.items():
            for candidate in candidates:
                if candidate in key and value:
                    return value
        return ""

    full_letter = first_match(["full cover letter", "cover letter", "letter"])
    short_email = first_match(["short email", "email version", "email"])
    tips = first_match(["customization tips", "tips", "customisation tips"])

    if not full_letter and sections:
        full_letter = next((v.strip() for v in sections.values() if v.strip()), "")

    return sections, full_letter, short_email, tips




# =========================================================
# VISUAL CV TEMPLATE PREVIEW
# =========================================================
def parse_cv_sections_for_template(cv_text: str) -> Dict[str, str]:
    """
    Lightweight parser to organize generated CV text into visual preview sections.
    It works best with the AI-generated CV draft but also handles plain text.
    """
    sections = {
        "header": "",
        "summary": "",
        "skills": "",
        "experience": "",
        "projects": "",
        "education": "",
        "certifications": "",
        "languages": "",
        "other": "",
    }

    text = clean_cv_text(cv_text)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return sections

    sections["header"] = "\n".join(lines[:4])
    current = "summary"

    heading_map = {
        "summary": ["summary", "profile", "professional profile", "career objective", "objective"],
        "skills": ["skills", "core skills", "technical skills", "kenntnisse"],
        "experience": ["experience", "work experience", "professional experience", "berufserfahrung"],
        "projects": ["projects", "project", "projekte"],
        "education": ["education", "ausbildung", "academic"],
        "certifications": ["certification", "certifications", "courses", "training", "weiterbildung"],
        "languages": ["languages", "sprachen"],
    }

    bucket_lines = {k: [] for k in sections if k != "header"}

    for line in lines[4:]:
        normalized = re.sub(r"[^a-zA-Z ]", "", line).strip().lower()
        matched = False
        for key, labels in heading_map.items():
            if normalized in labels or any(normalized.startswith(label) for label in labels):
                current = key
                matched = True
                break
        if not matched:
            bucket_lines.setdefault(current, []).append(line)

    for key, vals in bucket_lines.items():
        sections[key] = "\n".join(vals).strip()

    return sections

def html_escape(text_value: str) -> str:
    return (text_value or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

def lines_to_html(text_value: str) -> str:
    lines = [html_escape(x.strip()) for x in (text_value or "").splitlines() if x.strip()]
    if not lines:
        return "<p style='color:#64748b;'>—</p>"
    html = ""
    for line in lines:
        if line.startswith(("-", "•", "*")):
            html += f"<li>{line.lstrip('-•* ').strip()}</li>"
        else:
            html += f"<p>{line}</p>"
    if "<li>" in html:
        html = html.replace("<li>", "<ul><li>", 1)
        html = html[::-1].replace(">il/<", ">lu/<>il/<", 1)[::-1]
    return html

def resolve_template_style(template_name: str) -> str:
    """Map template names to visibly different preview/download layouts."""
    name = (template_name or "").lower()
    if any(x in name for x in ["career change", "pivot", "returning", "newcomer"]):
        return "Career Pivot"
    if any(x in name for x in ["graduate", "student", "intern", "thesis", "fresher", "stage", "werkstudent", "entry-level"]):
        return "Graduate Portfolio"
    if any(x in name for x in ["ats", "one-page", "one page", "minimal", "no photo"]):
        return "Minimal ATS"
    if any(x in name for x in ["modern", "portfolio", "netherlands", "dutch", "singapore"]):
        return "Creative Modern"
    if any(x in name for x in ["german", "lebenslauf", "austria", "swiss", "dach", "france", "french", "classique"]):
        return "German-Style Lebenslauf"
    if any(x in name for x in ["professional", "resume", "canadian", "canada", "australia", "uk", "us", "usa"]):
        return "Executive Slate"
    return "Minimal ATS"

def build_visual_cv_html(cv_text: str, template_name: str, target_country: str) -> str:
    style_key = resolve_template_style(template_name)
    data = parse_cv_sections_for_template(cv_text)
    header_lines = [x.strip() for x in data.get("header", "").splitlines() if x.strip()]
    name = header_lines[0] if header_lines else "Your Name"
    title = header_lines[1] if len(header_lines) > 1 else f"CV for {target_country}"
    contact = " · ".join(header_lines[2:]) if len(header_lines) > 2 else ""

    summary = lines_to_html(data.get("summary"))
    skills = lines_to_html(data.get("skills"))
    experience = lines_to_html(data.get("experience"))
    projects = lines_to_html(data.get("projects"))
    education = lines_to_html(data.get("education"))
    certifications = lines_to_html(data.get("certifications"))
    languages = lines_to_html(data.get("languages"))
    other = lines_to_html(data.get("other"))

    base_css = """
    <style>
      .cv-page {
        width: 794px;
        min-height: 1123px;
        margin: 0 auto 24px auto;
        background: white;
        color: #111827;
        box-shadow: 0 18px 50px rgba(0,0,0,0.35);
        border-radius: 8px;
        overflow: hidden;
        font-family: Arial, Helvetica, sans-serif;
      }
      .cv-page p { margin: 0 0 7px 0; line-height: 1.38; font-size: 13px; }
      .cv-page ul { margin: 0 0 8px 18px; padding: 0; }
      .cv-page li { font-size: 13px; margin-bottom: 5px; line-height: 1.35; }
      .cv-section-title { font-size: 12px; letter-spacing: 1.4px; text-transform: uppercase; font-weight: 800; margin: 16px 0 8px; }
      .cv-small { color:#64748b; font-size:12px; }
      @media (max-width: 850px) {
        .cv-page { width: 100%; min-height: auto; border-radius: 0; }
      }
    </style>
    """

    if style_key in ["Executive Slate", "German-Style Lebenslauf"]:
        return base_css + f"""
        <div class="cv-page">
          <div style="padding:42px 46px 20px; border-bottom:4px solid #1f2937;">
            <div style="font-size:34px; font-weight:800; letter-spacing:1px;">{html_escape(name)}</div>
            <div style="font-size:15px; color:#334155; margin-top:6px;">{html_escape(title)}</div>
            <div class="cv-small" style="margin-top:10px;">{html_escape(contact)}</div>
          </div>
          <div style="display:grid; grid-template-columns: 32% 68%;">
            <aside style="background:#f1f5f9; padding:26px 24px; min-height:900px;">
              <div class="cv-section-title">Skills</div>{skills}
              <div class="cv-section-title">Languages</div>{languages}
              <div class="cv-section-title">Education</div>{education}
              <div class="cv-section-title">Certifications</div>{certifications}
            </aside>
            <main style="padding:26px 32px;">
              <div class="cv-section-title">Profile</div>{summary}
              <div class="cv-section-title">Experience</div>{experience}
              <div class="cv-section-title">Projects</div>{projects}
              <div class="cv-section-title">Additional</div>{other}
            </main>
          </div>
        </div>
        """

    if style_key == "Minimal ATS":
        return base_css + f"""
        <div class="cv-page" style="padding:46px 56px;">
          <div style="border-bottom:2px solid #e5e7eb; padding-bottom:14px;">
            <div style="font-size:31px; font-weight:800;">{html_escape(name)}</div>
            <div style="font-size:15px; margin-top:5px;">{html_escape(title)}</div>
            <div class="cv-small" style="margin-top:8px;">{html_escape(contact)}</div>
          </div>
          <div class="cv-section-title">Professional Summary</div>{summary}
          <div class="cv-section-title">Core Skills</div>{skills}
          <div class="cv-section-title">Professional Experience</div>{experience}
          <div class="cv-section-title">Projects</div>{projects}
          <div class="cv-section-title">Education</div>{education}
          <div class="cv-section-title">Certifications</div>{certifications}
          <div class="cv-section-title">Languages</div>{languages}
        </div>
        """

    if style_key == "Creative Modern":
        return base_css + f"""
        <div class="cv-page">
          <div style="background:linear-gradient(135deg,#2563eb,#14b8a6); color:white; padding:42px 46px;">
            <div style="font-size:35px; font-weight:900;">{html_escape(name)}</div>
            <div style="font-size:16px; margin-top:6px; opacity:.95;">{html_escape(title)}</div>
            <div style="font-size:12px; margin-top:10px; opacity:.9;">{html_escape(contact)}</div>
          </div>
          <div style="padding:30px 42px; display:grid; grid-template-columns: 1.2fr .8fr; gap:30px;">
            <main>
              <div class="cv-section-title" style="color:#2563eb;">Profile</div>{summary}
              <div class="cv-section-title" style="color:#2563eb;">Experience</div>{experience}
              <div class="cv-section-title" style="color:#2563eb;">Projects</div>{projects}
            </main>
            <aside style="border-left:1px solid #e5e7eb; padding-left:24px;">
              <div class="cv-section-title" style="color:#0f766e;">Skills</div>{skills}
              <div class="cv-section-title" style="color:#0f766e;">Education</div>{education}
              <div class="cv-section-title" style="color:#0f766e;">Certifications</div>{certifications}
              <div class="cv-section-title" style="color:#0f766e;">Languages</div>{languages}
            </aside>
          </div>
        </div>
        """

    if style_key == "Career Pivot":
        return base_css + f"""
        <div class="cv-page" style="padding:42px 50px;">
          <div style="background:#f8fafc; border-left:6px solid #7c3aed; padding:20px 24px; margin-bottom:22px;">
            <div style="font-size:32px; font-weight:900;">{html_escape(name)}</div>
            <div style="font-size:15px; color:#4c1d95; margin-top:6px;">{html_escape(title)}</div>
            <div class="cv-small" style="margin-top:8px;">{html_escape(contact)}</div>
          </div>
          <div class="cv-section-title" style="color:#7c3aed;">Career Change Profile</div>{summary}
          <div class="cv-section-title" style="color:#7c3aed;">Transferable + Target Skills</div>{skills}
          <div class="cv-section-title" style="color:#7c3aed;">Relevant Projects / Training</div>{projects}
          <div class="cv-section-title" style="color:#7c3aed;">Previous Experience Reframed</div>{experience}
          <div style="display:grid; grid-template-columns:1fr 1fr; gap:24px;">
            <div><div class="cv-section-title" style="color:#7c3aed;">Education</div>{education}</div>
            <div><div class="cv-section-title" style="color:#7c3aed;">Languages</div>{languages}</div>
          </div>
        </div>
        """

    # Graduate Portfolio
    return base_css + f"""
    <div class="cv-page">
      <div style="padding:38px 46px; border-bottom:1px solid #e5e7eb;">
        <div style="font-size:33px; font-weight:900;">{html_escape(name)}</div>
        <div style="font-size:15px; color:#0369a1; margin-top:6px;">{html_escape(title)}</div>
        <div class="cv-small" style="margin-top:8px;">{html_escape(contact)}</div>
      </div>
      <div style="padding:28px 42px;">
        <div class="cv-section-title" style="color:#0369a1;">Career Objective</div>{summary}
        <div class="cv-section-title" style="color:#0369a1;">Education</div>{education}
        <div class="cv-section-title" style="color:#0369a1;">Projects</div>{projects}
        <div class="cv-section-title" style="color:#0369a1;">Technical Skills</div>{skills}
        <div class="cv-section-title" style="color:#0369a1;">Experience</div>{experience}
        <div class="cv-section-title" style="color:#0369a1;">Certifications & Languages</div>{certifications}{languages}
      </div>
    </div>
    """


def get_cv_template_options(target_country: str, user_status: str) -> Dict[str, str]:
    """Return templates that are genuinely different by country expectations."""
    country = (target_country or "").strip().lower()
    status = (user_status or "").strip().lower()

    country_templates = {
        "germany": {
            "German Lebenslauf - Structured": "German-style CV: clear sections, reverse chronology, languages, education, skills, and optional professional photo/personal details only if the user chooses.",
            "German ATS - No Photo": "Modern German application CV without photo/date of birth; safer for international companies and online portals.",
            "German Student / Werkstudent CV": "Best for students, thesis roles, internships, Praktikum, Werkstudent, Bachelorarbeit or Masterarbeit applications.",
        },
        "austria": {
            "Austria Lebenslauf - Structured": "Austrian-style CV with clear chronology, education, experience, skills, and language levels.",
            "DACH ATS - No Photo": "Clean DACH-market CV for online applications and international employers.",
            "Austria Student / Internship CV": "Education-first CV for Praktikum, trainee, thesis and junior opportunities.",
        },
        "switzerland": {
            "Swiss Professional CV": "Structured Swiss CV with strong profile, reverse chronology, languages and skills; concise and formal.",
            "Swiss International ATS": "Clean no-photo ATS version for multinational Swiss companies.",
            "Swiss Student / Internship CV": "Student-friendly CV for internships, thesis roles and graduate programmes.",
        },
        "united states": {
            "US One-Page Resume": "US resume style: one page, no photo, no date of birth, impact bullets, keywords and measurable results.",
            "US ATS Resume": "ATS-first American resume for online applications with simple formatting and strong keyword alignment.",
            "US Entry-Level Resume": "Education/projects-first resume for students, new grads and early-career applicants.",
        },
        "united kingdom": {
            "UK Two-Page CV": "UK CV style: professional profile, key skills, reverse chronology, no photo and no personal data.",
            "UK ATS CV": "Clean UK CV for job boards and recruiters, focused on evidence, skills and achievements.",
            "UK Graduate CV": "Graduate-friendly UK CV focused on education, projects, placements and transferable skills.",
        },
        "canada": {
            "Canadian Professional Resume": "Balanced Canadian resume: no photo, no date of birth, achievement-focused, recruiter-friendly, usually 1–2 pages.",
            "Canadian ATS Keyword Resume": "Strict ATS Canadian resume with simple formatting, strong keywords, skills and measurable impact bullets.",
            "Canadian Newcomer / Career Pivot Resume": "Useful for immigrants, newcomers, and career changers: highlights transferable experience and Canadian-style wording.",
        },
        "india": {
            "India Professional Resume": "Indian market resume with profile, skills, projects, experience and education; suitable for job portals.",
            "India ATS Resume": "Clean Naukri/LinkedIn-style resume with keywords, tools, achievements and project proof.",
            "India Fresher Resume": "Fresher-focused resume emphasizing education, projects, internships, tools and certifications.",
        },
        "france": {
            "France CV - Classique": "French-style CV with Profil, Expérience, Formation, Compétences and Langues; concise and structured.",
            "France ATS - International": "No-photo international French-market CV for online applications and multinational companies.",
            "France Stage / Alternance CV": "Student CV for stage, alternance, apprentissage and junior roles.",
        },
        "netherlands": {
            "Netherlands CV - Direct": "Dutch-style CV: direct, skills-focused, concise, no unnecessary personal details.",
            "Netherlands ATS CV": "Clean ATS CV for Dutch job portals and international companies.",
            "Netherlands Internship / Stage CV": "Stage/afstudeerstage focused CV for students and graduates.",
        },
        "australia": {
            "Australia Resume": "Australian resume style: 2–3 pages if needed, no photo, achievement-focused and recruiter-friendly.",
            "Australia ATS Resume": "Clean ATS version for Seek/LinkedIn-style applications.",
            "Australia Graduate Resume": "Graduate-program format with education, projects, placements and skills.",
        },
        "singapore": {
            "Singapore Professional Resume": "Concise Singapore-market resume with profile, skills, experience and education.",
            "Singapore ATS Resume": "ATS-friendly format for job portals and multinational employers.",
            "Singapore Graduate Resume": "Graduate/internship-oriented format with education, projects, internships and tools.",
        },
    }

    templates = country_templates.get(country, {
        "International ATS CV": "Safe global CV format: no photo, no sensitive personal details, clean sections and ATS-friendly layout.",
        "International Professional CV": "General professional CV for global applications with summary, skills, experience, education and projects.",
        "International Graduate CV": "Student/new graduate version focused on education, projects, internships and certifications.",
    })

    def priority(item):
        name = item[0].lower()
        if any(x in status for x in ["student", "thesis", "internship", "graduate", "entry"]):
            return 0 if any(x in name for x in ["student", "intern", "graduate", "fresher", "stage", "werkstudent"]) else 1
        if any(x in status for x in ["career changer", "returning", "break"]):
            return 0 if any(x in name for x in ["ats", "newcomer", "international"]) else 1
        return 0

    return dict(sorted(templates.items(), key=priority))

def get_template_instructions(template_name: str, target_country: str) -> str:
    country = (target_country or "the selected country").strip()
    name = (template_name or "").lower()
    country_rules = get_country_cv_rules(country)
    base = f"""
Target country: {country}
STRICT COUNTRY RULE: This CV must be written for {country} only. Do not mention Germany, DACH, the USA, Canada, or any other country unless it is the selected target country.
Recommended length: about {country_rules.get('pages', 2)} page(s), unless the user's experience requires otherwise.
Recommended sections: {', '.join(country_rules.get('sections', []))}.
Photo commonly expected: {'Yes' if country_rules.get('photo') else 'No'}.
Date of birth/personal details commonly expected: {'Yes' if country_rules.get('dob') else 'No'}.
Important exclusions: {country_rules.get('avoid', 'Avoid sensitive personal details and do not invent information.')}.
"""

    if any(x in name for x in ["german", "lebenslauf", "austria", "swiss", "dach"]):
        return base + """
Use a DACH-style structured CV with clear chronology, languages, education and professional experience.
Keep it formal, concise and readable. Do not invent photo, date of birth, nationality, marital status or visa details.
"""
    if any(x in name for x in ["us", "canadian", "canada", "australia", "uk", "ats", "one-page", "one page"]):
        return base + """
Use a clean ATS-friendly resume/CV. Avoid photo, date of birth, marital status, nationality and other sensitive personal details.
Prioritize achievements, action verbs, measurable impact and job-specific keywords.
"""
    if any(x in name for x in ["france", "french", "classique"]):
        return base + """
Use a concise French-market structure: Profil, Expérience professionnelle, Formation, Compétences, Langues.
Keep wording direct and professional. Photo is optional; do not invent personal details.
"""
    if any(x in name for x in ["netherlands", "dutch"]):
        return base + """
Use a direct Dutch-market CV: concise profile, skills, work experience, education and languages.
Avoid unnecessary personal details and keep achievements easy to scan.
"""
    if any(x in name for x in ["india", "fresher"]):
        return base + """
Use an Indian job-portal friendly resume with strong skills, projects, education, experience and certifications.
For freshers, place education, projects and tools high on the page.
"""
    if any(x in name for x in ["student", "intern", "graduate", "stage", "werkstudent", "thesis"]):
        return base + """
Use a student/graduate CV format. Put education, thesis/internship/project work, tools, coursework and learning proof near the top.
Avoid senior-sounding claims unless the CV proves them.
"""

    return base + """
Use a safe international ATS CV structure with summary, skills, experience, education, projects, certifications and languages.
Do not invent any personal details, employers, dates or achievements.
"""

def build_template_sample_cv(template_name: str, target_country: str) -> str:
    style = resolve_template_style(template_name)
    if style == "Graduate Portfolio":
        return f"""
Alex Morgan
Graduate Data Analyst Candidate
alex.morgan@email.com · Toronto, ON · linkedin.com/in/alexmorgan

Career Objective
Entry-level candidate with hands-on coursework, portfolio projects and strong motivation to build practical experience in {target_country}.

Education
Bachelor of Science in Computer Science | Example University | 2024
- Relevant coursework: databases, statistics, programming and data visualization.

Projects
- Sales Dashboard Project: built a dashboard using SQL and Tableau to identify monthly trends.
- Customer Sentiment Project: analyzed comments with Python and presented findings clearly.

Technical Skills
- Python, SQL, Excel, Tableau, data cleaning, reporting

Experience
Student Assistant | Example Organization | 2023-2024
- Supported data entry, documentation and team coordination.

Languages
English - Professional
""".strip()
    if style == "Career Pivot":
        return f"""
Alex Morgan
Customer Support Specialist transitioning to Data Analyst
alex.morgan@email.com · Toronto, ON · linkedin.com/in/alexmorgan

Career Change Profile
Customer-facing technical professional moving into data analysis, combining problem-solving, stakeholder communication and practical analytics training.

Transferable + Target Skills
- Customer troubleshooting, stakeholder communication, SQL, Python, Excel, dashboards

Relevant Projects / Training
- Data Analysis Bootcamp: completed projects using Python, SQL and visualization tools.
- Support Metrics Analysis: analyzed ticket trends and created improvement recommendations.

Previous Experience Reframed
Technical Support Specialist | Example Company | 2020-2024
- Resolved technical issues, documented patterns and collaborated with product teams to improve user experience.

Education
Professional Certificate in Data Analytics | 2024

Languages
English - Professional
""".strip()
    if style == "Minimal ATS":
        return f"""
Alex Morgan
Data Analyst
alex.morgan@email.com · Toronto, ON · linkedin.com/in/alexmorgan

Professional Summary
Data-focused professional with experience in reporting, customer insights and process improvement. Skilled in turning business questions into clear analysis and practical recommendations.

Core Skills
- SQL, Python, Excel, Tableau, data cleaning, reporting, stakeholder communication

Professional Experience
Data Analyst | Example Company | 2022-2024
- Built weekly reports that helped the team monitor service performance and identify recurring issues.
- Cleaned and analyzed customer data to support process improvement decisions.

Projects
- Operations Dashboard: created KPI dashboard for ticket volume, resolution time and customer trends.

Education
Bachelor of Science | Example University | 2020

Certifications
Data Analytics Certificate | 2024

Languages
English - Professional
""".strip()
    return f"""
Alex Morgan
Data & Operations Professional
alex.morgan@email.com · Toronto, ON · linkedin.com/in/alexmorgan

Profile
Results-oriented professional with experience improving service processes, analyzing operational information and communicating insights to non-technical teams.

Skills
- Data analysis, SQL, Python, Excel, Tableau, documentation, customer communication

Experience
Operations Analyst | Example Company | 2021-2024
- Improved reporting clarity by organizing recurring metrics into a reusable dashboard.
- Partnered with internal teams to identify process gaps and document practical solutions.

Projects
- Customer Trend Report: analyzed support data and summarized top improvement areas.

Education
Bachelor's Degree | Example University | 2020

Languages
English - Professional
""".strip()

def render_cv_template_preview(template_name: str, target_country: str, user_status: str):
    st.markdown("### Template Preview")
    template_description = get_cv_template_options(target_country, user_status).get(template_name, "")
    visual_style = resolve_template_style(template_name)
    st.markdown(f"""
<div class="glass-card">
  <div class="section-title">{html_escape(template_name)}</div>
  <div class="small-muted">Target country: {html_escape(target_country)} • Layout style: {html_escape(visual_style)}</div>
  <br>
  <div>{html_escape(template_description)}</div>
</div>
""", unsafe_allow_html=True)

    sample_cv = build_template_sample_cv(template_name, target_country)
    st.components.v1.html(build_visual_cv_html(sample_cv, template_name, target_country), height=620, scrolling=True)

    with st.expander("What this template includes", expanded=False):
        st.markdown(get_template_instructions(template_name, target_country))

def generate_country_cv_template(cv_text: str, target_country: str, user_status: str, template_name: str = "ATS Classic") -> str:
    selected_country = (target_country or "International").strip()
    template_instructions = get_template_instructions(template_name, selected_country)
    rules_text = country_resume_rules_text(selected_country)

    prompt = build_quality_prompt(
        task=f"Rebuild the user's CV into a {selected_country}-specific resume/CV using the selected template.",
        user_input=cv_text,
        expected_structure="""
1. Country Fit Notes
2. Full Resume Draft
3. WorkZo Changes
4. Missing Details to Confirm
"""
    )
    prompt += f"""

ABSOLUTE SELECTED COUNTRY FOR THIS CV: {selected_country}
IMPORTANT: Ignore any other country found in the career intelligence layer, profile, phone number, address, previous residence, or user history.
If the user lives in Germany but selected Canada, create a Canadian resume. Never say it is for Germany.

User status: {user_status}
Selected template: {template_name}

{rules_text}

Template instructions:
{template_instructions}

Output rules:
- Use these exact section headings: 1. Country Fit Notes, 2. Full Resume Draft, 3. WorkZo Changes, 4. Missing Details to Confirm.
- Do not use markdown symbols like ###, **, or bullet characters outside the resume bullets.
- Section 1 must be 3 short bullets only: country fit, removed/kept details, ideal length.
- Section 2 must be the full copy-ready resume/CV only.
- Use only the user's real information from the source CV.
- Do not invent employers, dates, degrees, certifications, tools, achievements, nationality, visa status, or language level.
- Remove or avoid personal details that are not appropriate for {selected_country}.
- For Canada/USA/UK/Australia/Singapore: do not include photo, date of birth, marital status, gender, religion, nationality, or full street address.
- For Germany/Austria/Switzerland: personal details are optional; do not invent them.
- The source CV may be extracted from a two-column PDF and may look mixed up. Reconstruct it logically into clean sections before rewriting.
- Keep contact details together at the top. Do not let phone/email/address get mixed into the profile summary.
- Keep skills grouped under skills, not inside experience. Keep education under education.
- Make the CV ATS-friendly, achievement-focused, and easy to scan.
- If something is missing, put it only in section 4.
"""
    return run_ai_prompt(
        prompt,
        system_addition=f"The selected CV country is {selected_country}. For this CV generation task, {selected_country} overrides every other country signal. Never mention Germany unless selected_country is Germany."
    )

def update_cv_with_ai(current_cv: str, update_notes: str, job_description: str = "") -> str:
    prompt = build_quality_prompt(
        task="Update and improve the user's CV using their instructions and optional job description.",
        user_input=f"Current CV:\n{current_cv}\n\nUpdate notes:\n{update_notes}\n\nJob description:\n{job_description}",
        expected_structure="""
1. Updated CV
2. Key Changes Made
3. Missing Details to Confirm
4. ATS / Country Suggestions
"""
    )
    prompt += """
Rules:
- Do not invent experience, employers, dates, or certifications.
- If a job description is provided, tailor keywords and positioning.
- If the user asks to add language, experience, or a new skill, include it only as stated.
- Keep the CV clean and ATS-friendly.
"""
    return run_ai_prompt(prompt)




def extract_job_post_from_url(url: str) -> str:
    """Best-effort job-post text extractor using only standard library.
    Some sites block scraping; users can still paste the job text manually.
    """
    url = (url or "").strip()
    if not url:
        return ""
    if not url.startswith(("http://", "https://")):
        url = "https://" + url
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=12) as response:
            raw = response.read(700000).decode("utf-8", errors="ignore")
        raw = re.sub(r"(?is)<(script|style|noscript).*?>.*?</\1>", " ", raw)
        raw = re.sub(r"(?is)<br\s*/?>", "\n", raw)
        raw = re.sub(r"(?is)</p>|</div>|</li>|</h[1-6]>", "\n", raw)
        text_only = re.sub(r"(?is)<[^>]+>", " ", raw)
        text_only = html.unescape(text_only)
        text_only = re.sub(r"[ \t]+", " ", text_only)
        text_only = re.sub(r"\n\s*\n+", "\n", text_only).strip()
        return text_only[:12000]
    except Exception:
        return ""

def combine_job_link_and_text(job_url: str, pasted_text: str) -> str:
    pasted_text = (pasted_text or "").strip()
    job_url = (job_url or "").strip()
    extracted = extract_job_post_from_url(job_url) if job_url else ""
    if extracted and pasted_text:
        return f"Job post link: {job_url}\n\nExtracted from link:\n{extracted}\n\nExtra details pasted by user:\n{pasted_text}"
    if extracted:
        return f"Job post link: {job_url}\n\nExtracted from link:\n{extracted}"
    return pasted_text

# =========================================================
# DOCUMENT TOOLS
# =========================================================
def show_document_tools():
    st.subheader(txt("document_tools"))
    st.caption(txt("document_hub_caption"))

    tabs = st.tabs([
        "Improve CV for a Job",
        "Update Resume Details",
        txt("cv_template_builder"),
        "Cover Letter Generator + Language",
        txt("cv_translator"),
    ])

    # -----------------------------------------------------
    # 1. Improve CV for a Job
    # -----------------------------------------------------
    with tabs[0]:
        st.markdown("### Improve CV for a Job")
        st.caption("Paste the job description. WorkZo will tailor your CV for that role.")
        current_cv = st.text_area(txt("your_cv"), value=organize_cv_for_display(st.session_state.cv_text), height=260, key="improve_cv_for_job_cv")
        job_desc_cv = st.text_area("Paste the job description", height=220, key="improve_cv_for_job_desc")

        if st.button(txt("improve"), key="btn_improve_cv_for_job_v49"):
            job_desc_combined = (job_desc_cv or "").strip()
            if not current_cv.strip() or not job_desc_combined.strip():
                st.warning("Please provide your CV and paste the job description.")
            else:
                with st.spinner("Improving CV for this job..."):
                    prompt = build_quality_prompt(
                        task="Tailor the user's CV to a specific job description.",
                        user_input=f"CV:\n{current_cv}\n\nJob description:\n{job_desc_combined}",
                        expected_structure="""
1. Fit Verdict
2. Better Professional Summary
3. Rewritten Experience Bullet Points
4. Missing Keywords from the Job
5. Skills to Move Higher
6. ATS-Friendly Suggestions
7. What Not to Change
8. Copy-Ready Tailored CV Sections
"""
                    )
                    prompt += """
Rules:
- Do not invent experience.
- Make it job-description specific.
- Focus on keywords, proof, and positioning.
"""
                    result = run_ai_prompt(prompt)
                    if render_error_or_success(result):
                        st.session_state.latest_cv_analysis = result
                        render_section_cards(result, default_expand=True)

    # -----------------------------------------------------
    # 2. Update Resume Details
    # -----------------------------------------------------
    with tabs[1]:
        st.markdown("### Update Resume Details")
        st.caption("Use this when your CV information changed: new experience, language level, tools, projects, certificates, or career break.")
        current_resume = st.text_area("Current Resume", value=organize_cv_for_display(st.session_state.cv_text), height=260, key="update_resume_current_v49")
        update_notes = st.text_area(
            "What changed?",
            height=160,
            key="update_resume_notes_v49",
            placeholder="Example: Add language level, add a recent project, update years of experience, add a new certificate..."
        )
        optional_job_desc = st.text_area(
            "Optional job description",
            height=160,
            key="update_resume_optional_job_v49",
            placeholder="Paste a job description only if you also want this update tailored to a role."
        )

        if st.button("Update My Resume", key="btn_update_resume_v49"):
            if not current_resume.strip() or not update_notes.strip():
                st.warning("Please provide your current resume and what changed.")
            else:
                with st.spinner("Updating resume..."):
                    result = update_cv_with_ai(current_resume, update_notes, optional_job_desc)
                    if render_error_or_success(result):
                        st.session_state.latest_cv_analysis = result
                        render_section_cards(result, default_expand=True)

                        updated_cv = get_section_text(result, ["Updated CV", "Full Updated Resume", "Full CV"])
                        if updated_cv:
                            st.markdown("### Copy-ready updated CV")
                            st.text_area("Updated CV", value=updated_cv, height=340, key="updated_cv_text_area_v49")
                            if st.button("Save as Dashboard Resume", key="save_updated_cv_dashboard_v49"):
                                set_new_resume_and_refresh(updated_cv)
                                st.success("Saved as your dashboard resume.")
                                st.rerun()

    # -----------------------------------------------------
    # 3. Country CV Template Builder
    # -----------------------------------------------------
    with tabs[2]:
        st.markdown(f"### {txt('cv_template_builder')}")
        st.caption("Choose a country and template. WorkZo will rebuild your resume in the style expected for that job market.")

        target_country = st.selectbox(
            "Target CV country",
            country_options,
            index=country_options.index(st.session_state.get("migration_country") or st.session_state.country)
            if (st.session_state.get("migration_country") or st.session_state.country) in country_options else 0,
            key="country_template_target_v51"
        )
        # Keep the selected CV country available to AI and downloads.
        st.session_state.migration_country = target_country

        user_status_for_template = st.session_state.get("user_status", "Not specified")
        template_options = get_cv_template_options(target_country, user_status_for_template)
        template_names = list(template_options.keys())

        selected_template = st.selectbox(
            "Choose resume template style",
            template_names,
            index=0,
            key="country_cv_template_choice_v51",
            help="Templates are suggested based on the target country and your career situation."
        )

        render_cv_template_preview(selected_template, target_country, user_status_for_template)

        cv_template_input = st.text_area(
            "Source CV / profile text (auto-cleaned from uploaded CV)",
            value=organize_cv_for_display(st.session_state.cv_text),
            height=260,
            key="country_template_cv_v51"
        )

        if st.button("Generate Resume in Selected Template", key="btn_country_cv_template_preview_v51"):
            track_button_click("Generate Country CV Template", "Document Tools", {"template": selected_template, "target_country": target_country})
            if not cv_template_input.strip():
                st.warning("Please provide your CV.")
            else:
                with st.spinner("Building country-specific CV in selected template..."):
                    result = generate_country_cv_template(
                        cv_template_input,
                        target_country,
                        user_status_for_template,
                        selected_template
                    )
                    if render_error_or_success(result):
                        full_cv = get_section_text(result, ["Full Resume Draft", "Full Country-Specific CV Draft", "Rebuilt CV in Target Country Style", "Full CV"], fallback_to_full=True)
                        st.session_state.generated_country_cv_result = result
                        st.session_state.generated_country_cv_text = strip_markdown_for_resume(full_cv)
                        st.session_state.generated_country_cv_template = selected_template
                        st.session_state.generated_country_cv_country = target_country

        if st.session_state.get("generated_country_cv_text"):
            st.markdown("### Edit Resume Text")
            edited_cv = st.text_area(
                "Make changes here, then click Update Preview below",
                value=st.session_state.get("country_cv_edit_buffer", st.session_state.generated_country_cv_text),
                height=360,
                key="country_cv_editable_text_v51"
            )
            if st.button("Update Preview with Edited Resume", key="btn_update_country_cv_preview_v77"):
                st.session_state.generated_country_cv_text = strip_markdown_for_resume(edited_cv)
                st.session_state.country_cv_edit_buffer = st.session_state.generated_country_cv_text
                st.success("Preview updated with your edits.")
                st.rerun()

            preview_cv = st.session_state.generated_country_cv_text
            st.markdown("### Visual Resume Preview")
            st.components.v1.html(
                build_visual_cv_html(
                    preview_cv,
                    st.session_state.get("generated_country_cv_template", selected_template),
                    st.session_state.get("generated_country_cv_country", target_country)
                ),
                height=850,
                scrolling=True
            )

            result = st.session_state.get("generated_country_cv_result", "")
            with st.expander("Template / Country Fit Notes", expanded=False):
                notes = get_section_text(result, ["Country Fit Notes", "Template and Country Fit Notes", "Country CV Format Notes"])
                st.write(notes)

            with st.expander("What WorkZo changed", expanded=False):
                changes = get_section_text(result, ["WorkZo Changes", "What Was Changed"])
                st.write(changes)

            with st.expander("Missing details to confirm", expanded=False):
                missing = get_section_text(result, ["Missing Details to Confirm"])
                st.write(missing)

            download_country = st.session_state.get('generated_country_cv_country', target_country)
            download_template = st.session_state.get('generated_country_cv_template', selected_template)
            clean_download_cv = strip_markdown_for_resume(st.session_state.generated_country_cv_text)
            safe_file_base = f"workzo_{download_country.lower().replace(' ', '_')}_{download_template.lower().replace(' ', '_')}_cv"

            pdf_data = make_styled_pdf_from_cv_text(
                f"WorkZo CV - {download_country} - {download_template}",
                clean_download_cv,
                download_template,
                download_country
            )
            docx_data = make_docx_from_cv_text(
                f"WorkZo CV - {download_country} - {download_template}",
                clean_download_cv,
                download_template,
                download_country
            )

            dl_col1, dl_col2 = st.columns(2)
            with dl_col1:
                if SimpleDocTemplate is not None:
                    st.download_button(
                        label="Download Resume as PDF",
                        data=pdf_data,
                        file_name=f"{safe_file_base}.pdf",
                        mime="application/pdf",
                        key="download_country_cv_pdf_v77"
                    )
            with dl_col2:
                if Document is not None:
                    st.download_button(
                        label="Download Resume as DOCX",
                        data=docx_data,
                        file_name=f"{safe_file_base}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_country_cv_docx_v77"
                    )

    # -----------------------------------------------------
    # 4. Cover Letter Generator
    # -----------------------------------------------------
    with tabs[3]:
        company_name = st.text_input("Company Name", key="doc_tools_company_name")
        role_name = st.text_input(txt("target_role"), key="doc_tools_role_name")
        cover_letter_language = st.selectbox(
            "Cover letter language",
            language_options,
            index=language_options.index(st.session_state.get("preferred_language", "English")) if st.session_state.get("preferred_language", "English") in language_options else 0,
            key="cover_letter_language_v60"
        )
        job_desc_letter = st.text_area("Paste the job description", height=220, key="doc_tools_job_desc")

        if st.button(txt("generate"), key="btn_cover_letter_v49"):
            track_button_click("Generate Cover Letter", "Document Tools")
            job_desc_letter_combined = (job_desc_letter or "").strip()
            if not role_name.strip() or not job_desc_letter_combined.strip():
                st.warning("Please enter a target role and paste the job description.")
            elif not st.session_state.cv_text.strip():
                st.warning("Please upload or create a CV first.")
            else:
                with st.spinner("Generating cover letter..."):
                    prompt = f"""
{txt('country_label')}: {st.session_state.country}
Target role: {role_name}
Company: {company_name if company_name.strip() else "Not specified"}

Candidate CV:
{st.session_state.cv_text}

Job description:
{job_desc_letter_combined}

Write a strong, personalized professional cover letter in {cover_letter_language}.

Quality requirements:
- Use the candidate's actual CV details.
- Connect 3 to 5 specific candidate strengths to the job description.
- Avoid generic sentences.
- Sound natural, confident, and human.
- Keep it suitable for the selected country and role level.
- Write the final cover letter and email version fully in {cover_letter_language}.

Return in this exact structure:

1. Full Cover Letter
2. Short Email Version
3. 3 Customization Tips

Important:
- Do not leave any section empty.
- Write complete content for each section.
"""
                    result = run_ai_prompt(prompt, force_language=cover_letter_language)
                    if render_error_or_success(result):
                        if not result.strip():
                            st.error("Cover letter generation returned an empty response. Please try again.")
                        else:
                            st.session_state.latest_cover_letter = result
                            render_section_cards(result, default_expand=True)

                            sections = numbered_sections_to_markdown(result)
                            full_letter = sections.get("Full Cover Letter", "").strip()
                            short_email = sections.get("Short Email Version", "").strip()

                            if full_letter:
                                st.markdown("### Full Cover Letter")
                                st.text_area("Generated Cover Letter", value=full_letter, height=320)
                                st.download_button(
                                    "Download Cover Letter as TXT",
                                    data=full_letter,
                                    file_name="cover_letter.txt",
                                    mime="text/plain",
                                    key="download_cover_letter_txt_v49"
                                )
                            if short_email:
                                st.markdown("### Short Email Version")
                                st.text_area("Generated Short Email", value=short_email, height=180)

    # -----------------------------------------------------
    # 5. CV Translator
    # -----------------------------------------------------
    with tabs[4]:
        col1, col2 = st.columns(2)
        with col1:
            default_src = language_options.index("English") if "English" in language_options else 0
            source_lang = st.selectbox(txt("translator_source"), language_options, index=default_src, key="cv_src_lang")
        with col2:
            target_default = language_options.index("German") if "German" in language_options else 0
            target_lang = st.selectbox(txt("translator_target"), language_options, index=target_default, key="cv_tgt_lang")

        cv_input = st.text_area(txt("cv_input"), value=organize_cv_for_display(st.session_state.cv_text), height=280)

        if st.button(txt("translate"), key="btn_cv_translate_v49"):
            track_button_click("Translate CV", "Document Tools")
            if not cv_input.strip():
                st.warning("Please provide your CV.")
            else:
                with st.spinner("Translating CV..."):
                    prompt = f"""
Translate and localize this CV from {source_lang} to {target_lang} for professional job applications.

Quality requirements:
- Translate naturally, not word-for-word.
- Keep the CV accurate. Do not invent experience, tools, dates, employers, or achievements.
- Adapt section headings to the target language and target-country CV style.
- Preserve achievements, metrics, tools, and responsibilities.
- Improve awkward phrasing while keeping the original meaning.
- Keep bullet points readable and ATS-friendly.
- Do not merge the CV into one paragraph.

CV:
{cv_input}
"""
                    result = run_ai_prompt(
                        prompt,
                        system_addition=f"Translate the content fully into {target_lang}. Return only the translated CV in {target_lang}.",
                        force_language=target_lang
                    )

                    if render_error_or_success(result):
                        st.session_state.latest_cv_translation = result
                        st.markdown(f"### {txt('copy_ready')}")
                        st.text_area("Translated CV", value=result, height=350)

def render_country_fit_cards(country_text: str):
    rows = []
    for raw in str(country_text or "").splitlines():
        item = raw.strip().lstrip("-• ").strip()
        if item:
            rows.append(item)
    if not rows:
        st.info(txt("not_analyzed"))
        return

    st.caption(txt("country_fit_note"))
    for item in rows[:3]:
        if " - " in item:
            country, reason = item.split(" - ", 1)
        elif ":" in item:
            country, reason = item.split(":", 1)
        else:
            country, reason = item, ""
        reason = reason.strip()
        if len(reason) > 120:
            reason = reason[:120].rsplit(" ", 1)[0] + "..."
        card_html = f"""
<div class="card">
  <div class="section-title">{html.escape(country.strip())}</div>
  <div class="small-muted">{html.escape(reason or txt('view_details'))}</div>
</div>
"""
        st.markdown(card_html, unsafe_allow_html=True)

    if len(rows) > 3:
        with st.expander(txt("country_fit_details"), expanded=False):
            for item in rows[3:]:
                st.markdown(f"- {item}")

# =========================================================
# DASHBOARD
# =========================================================
# DASHBOARD
# =========================================================
def get_workflow_progress() -> int:
    """Simple visible progress for the user journey."""
    cv_ready = bool(str(st.session_state.get("cv_text", "")).strip())
    scores_ready = bool(st.session_state.get("cv_score_value") or st.session_state.get("ats_score_value"))
    used_tool = any(bool(str(st.session_state.get(k, "")).strip()) for k in [
        "latest_job_analysis", "latest_cv_analysis", "latest_cover_letter", "latest_cv_translation", "next_best_step_result"
    ])
    completed_action = any(bool(str(st.session_state.get(k, "")).strip()) for k in [
        "latest_generated_cv", "latest_cover_letter", "latest_job_analysis"
    ])
    if not cv_ready:
        return 0
    if completed_action:
        return 100
    if used_tool:
        return 75
    if scores_ready:
        return 55
    return 30



def get_nav_items() -> List[Tuple[str, str, str]]:
    return [
        ("dashboard", "🏠 " + txt("dashboard"), txt("dashboard_desc_short")),
        ("job_assist", "🔎 " + txt("job_assist"), txt("job_assist_desc_short")),
        ("cv_documents", "📄 " + txt("cv_documents"), txt("cv_documents_desc_short")),
        ("workobot", "🤖 " + txt("workobot"), txt("workobot_desc_short")),
    ]

def set_single_preferred_language(language: str):
    """Preferred Language is the only visible language setting.
    It controls app labels where supported, AI replies, and generated documents.
    """
    language = language or "English"
    st.session_state.preferred_language = language
    st.session_state.ui_language = language if language in UI_TEXT else "English"
    st.session_state.response_language = language

def go_to_nav(page_key: str):
    queue_navigation(page_key)
    st.rerun()

def get_recommended_next_action() -> Dict[str, str]:
    """Return one clear SaaS-style next action for the dashboard."""
    cv_uploaded = bool(str(st.session_state.get("cv_text", "")).strip())
    ats_score = st.session_state.get("ats_score_value")
    latest_job = bool(str(st.session_state.get("latest_job_analysis", "")).strip())

    try:
        ats_score_number = int(ats_score or 0)
    except Exception:
        ats_score_number = 0

    if not cv_uploaded:
        return {"title": txt("next_action_upload_title"), "desc": txt("next_action_upload_desc"), "button": txt("next_action_upload_button"), "target": "onboarding"}
    if ats_score_number and ats_score_number < 75:
        return {"title": txt("next_action_improve_title"), "desc": txt("next_action_improve_desc"), "button": txt("next_action_improve_button"), "target": "cv_documents"}
    if not latest_job:
        return {"title": txt("next_action_job_title"), "desc": txt("next_action_job_desc"), "button": txt("next_action_job_button"), "target": "job_assist"}
    return {"title": txt("next_action_interview_title"), "desc": txt("next_action_interview_desc"), "button": txt("next_action_interview_button"), "target": "workobot"}

def render_recommended_next_action():
    action = get_recommended_next_action()
    st.markdown(f"""
    <div class='next-action-card'>
        <div class='next-action-label'>{txt('recommended_next_step')}</div>
        <div class='next-action-title'>{html.escape(action['title'])}</div>
        <div class='next-action-copy'>{html.escape(action['desc'])}</div>
    </div>
    """, unsafe_allow_html=True)
    if action["target"] == "onboarding":
        if st.button(action["button"], use_container_width=True, key="dashboard_recommended_next_action"):
            track_button_click(action["button"], "Dashboard")
            reset_onboarding()
            st.rerun()
    else:
        st.button(
            action["button"],
            use_container_width=True,
            key="dashboard_recommended_next_action",
            on_click=queue_navigation,
            args=(action["target"],),
        )

def show_dashboard():
    consume_pending_navigation()
    maybe_scroll_to_top()
    with st.sidebar:
        st.button("🏠 WORKZO AI", key="sidebar_workzo_home_button", use_container_width=True, on_click=queue_navigation, args=("dashboard",))
        st.markdown("<div class='workzo-sidebar-muted'>Beta V9.4 · Cleaner workspace</div>", unsafe_allow_html=True)

        st.markdown("<div class='workzo-sidebar-section-label'>Profile</div>", unsafe_allow_html=True)
        st.markdown(f"<span class='workzo-sidebar-chip'>🌍 {html.escape(str(st.session_state.get('country', '')))}</span>", unsafe_allow_html=True)
        st.markdown(f"<span class='workzo-sidebar-chip'>👤 {html.escape(str(st.session_state.get('user_status', txt('not_specified'))))}</span>", unsafe_allow_html=True)
        if st.session_state.get("migration_country") and st.session_state.get("migration_country") != st.session_state.country:
            st.markdown(f"<span class='workzo-sidebar-chip'>🎯 {html.escape(str(st.session_state.migration_country))}</span>", unsafe_allow_html=True)

        st.markdown("<div class='workzo-sidebar-section-label'>Settings</div>", unsafe_allow_html=True)
        language_list = language_options if language_options else ["English", "German", "Dutch"]
        current_language = st.session_state.get("preferred_language", "English")
        if current_language not in language_list:
            current_language = "English" if "English" in language_list else language_list[0]
        preferred = st.selectbox(
            txt("preferred_language"),
            language_list,
            index=language_list.index(current_language),
            key="sidebar_preferred_language",
            help=txt("preferred_language_help")
        )
        set_single_preferred_language(preferred)

        st.markdown("<div class='workzo-sidebar-section-label'>Navigation</div>", unsafe_allow_html=True)
        nav_items = get_nav_items()
        nav_labels = {key: label for key, label, _ in nav_items}
        nav_descriptions = {key: desc for key, _, desc in nav_items}
        nav_keys = [key for key, _, _ in nav_items]
        if st.session_state.get("founder_unlocked") and "founder_dashboard" not in nav_keys:
            nav_keys.append("founder_dashboard")
            nav_labels["founder_dashboard"] = "🔐 " + txt("founder_dashboard")
            nav_descriptions["founder_dashboard"] = "Private founder analytics and feedback."
        current_nav = st.session_state.get("nav_page", "dashboard")
        if current_nav not in nav_keys:
            current_nav = "dashboard"
        page_key = st.radio(
            txt("navigation"),
            nav_keys,
            index=nav_keys.index(current_nav),
            format_func=lambda k: nav_labels.get(k, k),
            label_visibility="collapsed",
            key=f"sidebar_main_navigation_radio_{st.session_state.get('nav_change_nonce', 0)}"
        )
        if page_key != current_nav:
            st.session_state.page = page_key
            st.session_state.nav_page = page_key
            request_scroll_to_top()
            st.rerun()
        st.session_state.page = page_key
        st.session_state.nav_page = page_key
        track_feature_view(nav_labels.get(page_key, page_key))
        st.markdown(f"<div class='workzo-sidebar-muted'>{html.escape(nav_descriptions.get(page_key, ''))}</div>", unsafe_allow_html=True)

        st.markdown("<div class='workzo-sidebar-section-label'>Progress</div>", unsafe_allow_html=True)
        progress_pct = get_workflow_progress()
        st.progress(progress_pct / 100)
        st.markdown(f"<div class='sidebar-progress-label'>{txt('workflow_progress')}: {progress_pct}%</div>", unsafe_allow_html=True)

        founder_pin = os.getenv("FOUNDER_PIN") or get_streamlit_secret("FOUNDER_PIN")
        with st.expander(txt("founder_access"), expanded=False):
            if founder_pin:
                entered_pin = st.text_input(txt("founder_pin"), type="password", key="founder_pin_input_sidebar")
                if entered_pin == founder_pin:
                    st.session_state.founder_unlocked = True
                    st.success("Founder mode unlocked.")
                    sync_navigation_state("founder_dashboard")
            else:
                st.caption("FOUNDER_PIN is not configured. For local testing, create a temporary PIN below. Before public testing, add FOUNDER_PIN in Streamlit Secrets.")
                temp_pin = st.text_input("Temporary founder PIN", type="password", key="founder_temp_pin_sidebar")
                if temp_pin and len(temp_pin) >= 4:
                    st.session_state.founder_unlocked = True
                    st.success("Founder mode unlocked for this session.")
                    sync_navigation_state("founder_dashboard")

        st.markdown("---")
        if st.button(txt("edit_setup"), use_container_width=True, key="sidebar_edit_onboarding_button"):
            reset_onboarding()
            request_scroll_to_top()
            st.rerun()

    if page_key == "dashboard":
        st.markdown(f"""
        <div class="hero-card">
            <div style="display:flex; justify-content:space-between; gap:14px; flex-wrap:wrap; align-items:center;">
                <div>
                    <div style="font-size:0.95rem; color:#cbd5e1; margin-bottom:6px;">{txt('career_command_center')}</div>
                    <div style="font-size:1.65rem; font-weight:700; color:white; margin-bottom:6px;">{txt('career_move_organized')}</div>
                    <div style="color:#dbeafe; max-width:760px;">{txt('start_here_intro')}</div>
                </div>
                <div>
                    <div class="pill">{txt('country_label')}: {st.session_state.country}</div>
                    <div class="pill">{txt('status_label')}: {st.session_state.get('user_status', txt('not_specified'))}</div>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)

        # Scores first: users immediately see the main CV result before navigation help.
        s1, s2 = st.columns(2)
        with s1:
            render_metric_card(txt("resume_score"), str(st.session_state.cv_score_value or "—"), txt("metric_resume_quality"))
        with s2:
            render_metric_card(txt("ats_score"), str(st.session_state.ats_score_value or "—"), txt("metric_ats_friendly"))

        render_recommended_next_action()

        st.markdown(f"<div class='workzo-section-title'>{txt('navigation_help')}</div>", unsafe_allow_html=True)
        h1, h2, h3 = st.columns(3)
        with h1:
            st.markdown(f"<div class='nav-help-card'><b>🔎 {txt('job_assist')}</b><br>{txt('job_assist_desc_short')}</div>", unsafe_allow_html=True)
            st.button(txt("go_job_assist"), use_container_width=True, key="dash_go_job_assist", on_click=queue_navigation, args=("job_assist",))
        with h2:
            st.markdown(f"<div class='nav-help-card'><b>📄 {txt('cv_documents')}</b><br>{txt('cv_documents_desc_short')}</div>", unsafe_allow_html=True)
            st.button(txt("go_cv_documents"), use_container_width=True, key="dash_go_cv_docs", on_click=queue_navigation, args=("cv_documents",))
        with h3:
            st.markdown(f"<div class='nav-help-card'><b>🤖 {txt('workobot')}</b><br>{txt('workobot_desc_short')}</div>", unsafe_allow_html=True)
            st.button(txt("go_workobot"), use_container_width=True, key="dash_go_workobot", on_click=queue_navigation, args=("workobot",))

        st.markdown(f"<div class='workzo-section-title'>{txt('what_next')}</div>", unsafe_allow_html=True)
        n1, n2 = st.columns([1.1, 0.9])
        with n1:
            st.markdown("<div class='glass-card'>", unsafe_allow_html=True)
            status = st.session_state.get("user_status", "")
            if st.session_state.next_actions_detected:
                st.markdown(st.session_state.next_actions_detected)
            target_country = st.session_state.get("migration_country") or st.session_state.get("country")
            if "migrate" in status.lower() or "abroad" in status.lower():
                st.markdown(f"- {txt('next_migrate_1')}\n- {txt('next_migrate_2')}\n- {txt('next_migrate_3')}")
            elif "graduate" in status.lower():
                st.markdown(f"- {txt('next_graduate_1')}\n- {txt('next_graduate_2')}\n- {txt('next_graduate_3')}")
            elif "career changer" in status.lower():
                st.markdown(f"- {txt('next_changer_1')}\n- {txt('next_changer_2')}\n- {txt('next_changer_3')}")
            else:
                st.markdown(f"- {txt('next_default_1')}\n- {txt('next_default_2')}\n- {txt('next_default_3')}")
            st.markdown("</div>", unsafe_allow_html=True)
        with n2:
            st.markdown("<div class='glass-card'>", unsafe_allow_html=True)
            if st.session_state.country_cv_readiness:
                st.markdown(f"**{txt('country_cv_readiness')}:** {st.session_state.country_cv_readiness}")
            if st.session_state.status_guidance:
                st.markdown(f"**{txt('status_guidance')}:** {st.session_state.status_guidance}")
            if not st.session_state.country_cv_readiness and not st.session_state.status_guidance:
                target_country = st.session_state.get("migration_country") or st.session_state.get("country")
            st.markdown(f"- {txt('target_country_label')}: **{target_country}**\n- {txt('judge_market')}\n- {txt('use_country_template')}")
            st.markdown("</div>", unsafe_allow_html=True)

        st.markdown(f"<div class='workzo-section-title'>{txt('resume_insights')}</div>", unsafe_allow_html=True)
        st.markdown("<div class='resume-insights-line'>", unsafe_allow_html=True)

        insight_col, strength_col, improve_col = st.columns([1.45, 1, 1], gap="large")
        with insight_col:
            st.markdown(f"<div class='workzo-mini-title'>{txt('detected_summary')}</div>", unsafe_allow_html=True)
            st.markdown(f"<div class='workzo-text'>{html.escape(st.session_state.profile_summary or txt('not_analyzed'))}</div>", unsafe_allow_html=True)
            if st.session_state.key_skills_detected:
                st.markdown(f"<div class='workzo-muted-title'>{txt('detected_skills')}</div>", unsafe_allow_html=True)
                st.markdown(st.session_state.key_skills_detected)
        with strength_col:
            st.markdown(f"<div class='workzo-mini-title'>{txt('strengths')}</div>", unsafe_allow_html=True)
            st.markdown(st.session_state.resume_strengths if st.session_state.resume_strengths else f"<span class='small-muted'>{txt('not_analyzed')}</span>", unsafe_allow_html=True)
        with improve_col:
            st.markdown(f"<div class='workzo-mini-title'>{txt('improvements')}</div>", unsafe_allow_html=True)
            st.markdown(st.session_state.resume_improvements if st.session_state.resume_improvements else f"<span class='small-muted'>{txt('not_analyzed')}</span>", unsafe_allow_html=True)

        st.markdown("</div>", unsafe_allow_html=True)
        st.markdown("---")
        # Removed country-fit cards from dashboard to keep it focused.


    elif page_key == "cv_documents":
        show_document_tools()

    elif page_key == "job_assist":
        st.subheader(txt("job_assist"))
        assist_tab1, assist_tab2 = st.tabs([txt("find_jobs"), txt("understand_job")])

        with assist_tab1:
            st.markdown(f"### {txt('find_jobs')}")
            st.caption("Live jobs are matched globally by selected country, career status, experience level, and remote/online options. Use “Anywhere” for more results.")

            target_country_options = country_options
            country_index = target_country_options.index(st.session_state.country) if st.session_state.country in target_country_options else 0
            search_country = st.selectbox("Country for job search", target_country_options, index=country_index, key="job_assist_search_country")

            job_search_focus = st.selectbox(
                "Job search focus",
                [
                    "Based on my career status",
                    "Student / Thesis / Internship",
                    "Freshers / entry-level",
                    "Apply online / remote",
                    "Career changer friendly",
                    "Experienced roles"
                ],
                key="job_search_focus",
                help="This helps WorkZo search with the right seniority and job keywords."
            )

            location_options = [f"Anywhere in {search_country}"] + fetch_country_cities(search_country)
            final_location = st.selectbox(
                txt("preferred_location"),
                options=location_options,
                index=0,
                help="This box is searchable. Start typing the city name inside this dropdown.",
                key="job_assist_location"
            )
            if final_location.startswith("Anywhere in "):
                final_location = search_country

            cv_for_jobs = st.text_area(txt("your_cv"), value=organize_cv_for_display(st.session_state.cv_text), height=240, key="job_assist_cv")
            target_titles = st.text_input(
                "Optional target job titles",
                placeholder="Example: Data Analyst, Junior IT Support, Customer Success",
                key="job_assist_target_titles",
                help="Leave this empty and WorkZo will use roles detected from your CV. Add 2–4 titles for better matching."
            )
            job_status_for_search = st.session_state.get('user_status', 'Not specified')
            if job_search_focus == "Student / Thesis / Internship":
                job_status_for_search = STUDENT_STATUS_INTERNAL
            elif job_search_focus == "Freshers / entry-level":
                job_status_for_search = "Fresh graduate / entry level"
            elif job_search_focus == "Apply online / remote":
                job_status_for_search = "Apply online / remote"
            elif job_search_focus == "Career changer friendly":
                job_status_for_search = "Career changer"
            elif job_search_focus == "Experienced roles":
                job_status_for_search = "Experienced professional"
            st.caption(f"Matching for: {job_status_for_search} • {search_country}")
            if is_student_thesis_status(job_status_for_search):
                render_student_opportunity_guidance(search_country)

            if st.button(txt("generate"), key="btn_find_jobs_v42"):
                track_button_click("Find Jobs", "Job Assist")
                if not cv_for_jobs.strip():
                    st.warning("Please provide your CV.")
                else:
                    with st.spinner("Searching worldwide live jobs..."):
                        roles_from_input = [x.strip() for x in target_titles.split(",") if x.strip()]
                        roles_from_input = build_role_suggestions(
                            roles_from_input,
                            st.session_state.suggested_roles_detected,
                            st.session_state.current_role_detected
                        )

                        live_jobs = fetch_live_jobs_global(search_country, roles_from_input, final_location, job_status_for_search)

                        render_job_board_search_cards(search_country, final_location, roles_from_input, job_status_for_search)

                        render_live_jobs(live_jobs, search_country, max_visible=30)

                        plan = generate_job_search_plan(
                            search_country,
                            final_location,
                            roles_from_input,
                            cv_for_jobs,
                            live_jobs
                        )
                        render_job_plan(plan)

        with assist_tab2:
            st.markdown(f"### {txt('understand_job')}")
            st.caption("Paste the job description to see fit, hidden expectations, salary estimate, likely interview focus, and how to tailor your application.")

            job_desc = st.text_area("Paste the job description", key="job_desc_v42", height=240)

            if st.button(txt("analyze"), key="btn_understand_job_v42"):
                track_button_click("Understand Job", "Job Assist")
                job_desc_combined = (job_desc or "").strip()
                if not job_desc_combined.strip():
                    st.warning("Please paste the job description.")
                else:
                    with st.spinner("Analyzing the job..."):
                        prompt = f"""
{txt('country_label')}: {st.session_state.country}
Candidate CV:
{st.session_state.cv_text}

Analyze this job description for the candidate.

Return in this exact structure:

1. Role Snapshot
2. Main Responsibilities
3. Must-Have Requirements
4. Nice-to-Have Requirements
5. Hidden Expectations
6. Salary Range Estimate for {st.session_state.country}
7. Fit Verdict
Give a realistic fit score from 0 to 100 and explain.
8. Match Breakdown
- Strong matches
- Partial matches
- Clear gaps
9. Risks / Concerns
10. How to Tailor the CV
11. Likely Interview Focus
12. Should You Apply?
Give a short recommendation: Yes / Yes, but tailor / Maybe later

Job description:
{job_desc_combined}
"""
                        result = run_ai_prompt(prompt)
                        if render_error_or_success(result):
                            st.session_state.latest_job_analysis = result
                            st.session_state.job_fit_score_value = parse_score(result)
                            if st.session_state.job_fit_score_value is not None:
                                show_gauge(st.session_state.job_fit_score_value, txt("job_fit_score"))
                            st.markdown(f"### {txt('job_fit_analysis')}")
                            render_section_cards(result, default_expand=True)

    elif page_key == "founder_dashboard":
        if st.session_state.get("founder_unlocked"):
            render_founder_dashboard()
        else:
            st.warning(txt("no_founder_pin"))

    elif page_key in ["workobot", "interview_practice", "career_insights"]:
        # Single Work-O-Bot view. Removed duplicate tabs/buttons to keep the page simple.
        show_workobot()
    render_feedback_collector(nav_labels.get(page_key, page_key) if isinstance(page_key, str) else "General")
    render_issue_reporter(nav_labels.get(page_key, page_key) if isinstance(page_key, str) else "General")

    st.divider()
    st.caption("WORKZO AI V8.0 • Beta • Downloads • Issue reporting • Cleaner UX")
    st.caption("⚠️ WorkZo AI is currently in Beta. Some results may not be perfect yet. Your feedback helps improve the tool. CV text and personal documents are not stored in analytics.")

# =========================================================
# ROUTER
# =========================================================
if not st.session_state.onboarding_complete or st.session_state.page == "onboarding":
    show_onboarding()
else:
    show_dashboard()

